# -*- coding: utf-8 -*-
"""
Estrattore Partite - App base
Stack: Python + Streamlit + Supabase

Funzionalità v1:
  1. ESTRATTORE  -> incolla i blocchi "ULTIMI INCONTRI" di 2 squadre,
                    parsing automatico, risultati modificabili, copia, salvataggio a DB.
  2. DATABASE    -> elenco partite dalla più recente alla meno recente,
                    modifica/inserimento risultati, dettaglio e copia, export Excel.
  3. CONFIGURAZIONE -> gestione utenti + backup Excel dell'intera App.
"""

import html as _html
import hashlib
import io
import json
import re
import uuid
from datetime import date, datetime
from difflib import SequenceMatcher

import analisi
import evidenze
import segnali
import racconto
import statistico
import snapshot as snapmod
import backtest as bt

import pandas as pd
import streamlit as st
from streamlit.components.v1 import html as components_html

# --- Supabase / bcrypt sono opzionali all'avvio, così l'app parte anche senza secrets ---
try:
    from supabase import create_client
except Exception:  # pragma: no cover
    create_client = None

try:
    import bcrypt
except Exception:  # pragma: no cover
    bcrypt = None


st.set_page_config(page_title="Estrattore Partite", page_icon="⚽", layout="wide")

ESITI = {"V", "N", "P"}


# =============================================================================
#  SUPABASE
# =============================================================================
@st.cache_resource
def get_client():
    if create_client is None:
        return None
    try:
        url = st.secrets["supabase"]["url"]
        key = st.secrets["supabase"]["key"]
    except Exception:
        return None
    return create_client(url, key)


def supabase_pronto():
    return get_client() is not None


# =============================================================================
#  PARSING
# =============================================================================
def _num(s):
    """Pulisce un numero: rimuove i (rigori) e converte la virgola in punto."""
    if s is None:
        return None
    s = re.sub(r"\(.*?\)", "", str(s)).replace(",", ".").strip()
    return s if s != "" else None


def _to_int(s):
    s = _num(s)
    if s is None:
        return None
    try:
        return int(float(s))
    except ValueError:
        return None


def _is_data(r):
    return re.fullmatch(r"\d{2}\.\d{2}\.\d{2}", r.strip()) is not None


def _data_iso(r):
    return datetime.strptime(r.strip(), "%d.%m.%y").date()


def _is_header(r):
    return re.search(r"INCONTR", r, re.I) is not None


def _team_da_header(r):
    return re.sub(r".*INCONTR\w*\s*:?\s*", "", r, flags=re.I).strip()


def _parse_sezione(righe):
    """Estrae le partite da una singola sezione squadra.
    Ritorna (partite, avanzo) dove avanzo = righe non consumate (di solito le quote)."""
    partite, avanzo = [], []
    i, n = 0, len(righe)
    while i < n:
        r = righe[i]
        if _is_data(r):
            data = _data_iso(r)
            i += 1
            buf, esito = [], None
            while i < n:
                x = righe[i]
                if _is_data(x):          # nuova partita: quella corrente è incompleta
                    break
                if x in ESITI:           # l'esito chiude la partita
                    esito = x
                    i += 1
                    break
                buf.append(x)
                i += 1
            # struttura attesa in coda: casa, casa, trasf, trasf, gol_casa, gol_trasf
            if esito and len(buf) >= 7:
                partite.append({
                    "data_iso": data,
                    "data_str": data.strftime("%d.%m.%y"),
                    "competizione": buf[-7],
                    "casa": buf[-6],
                    "trasferta": buf[-4],
                    "gol_casa": _to_int(buf[-2]),
                    "gol_trasferta": _to_int(buf[-1]),
                    "qualificatore": " ".join(buf[:-7]).strip() or None,
                })
            # se non valida, la ignoriamo silenziosamente
        else:
            avanzo.append(r)
            i += 1
    return partite, avanzo


def _parse_quote(avanzo):
    """Interpreta le righe numeriche di coda.
    Ordine atteso (ogni riga può mancare):
      1X2 (3 numeri) · O/U 2.5 · Goal/NoGoal · forma (casa, trasf) · valore rose (casa, trasf)

    Distinzione forma / valore rose:
      - la FORMA è sempre in formato decimale (X.XX o X.X, anche con virgola) -> es. 7.21, 8.9
      - il VALORE ROSE è sempre un numero intero -> es. 287 325 (può stare anche sotto i 10)
    Quando compare una sola delle due, la riconosciamo dal formato (decimale=forma,
    intero=rose). Con entrambe presenti la forma è comunque la prima riga."""
    num = []  # (valori, è_decimale)
    for r in avanzo:
        toks = r.split()
        norm = [t.replace(",", ".") for t in toks]
        if toks and all(re.fullmatch(r"-?\d+(\.\d+)?", t) for t in norm):
            is_dec = any(("." in t or "," in t) for t in toks)
            num.append(([float(t) for t in norm], is_dec))

    q = {}
    tre = next((v for v, _ in num if len(v) == 3), None)
    due = [(v, d) for v, d in num if len(v) == 2]

    if tre:
        q["1"], q["X"], q["2"] = tre
        if len(due) >= 1:
            q["over25"], q["under25"] = due[0][0]
        if len(due) >= 2:
            q["goal"], q["nogoal"] = due[1][0]
        extra = due[2:]
    else:
        # nessuna quota 1X2 -> le righe a 2 numeri sono forma/rose
        extra = due

    for vals, is_dec in extra:
        if is_dec and "forma_casa" not in q:            # decimale -> forma
            q["forma_casa"], q["forma_trasferta"] = vals
        elif not is_dec and "val_casa" not in q:        # intero -> valore rose
            q["val_casa"], q["val_trasferta"] = vals
        else:                                           # fallback posizionale
            if "forma_casa" not in q:
                q["forma_casa"], q["forma_trasferta"] = vals
            elif "val_casa" not in q:
                q["val_casa"], q["val_trasferta"] = vals
    return q


def parse_incontri(testo):
    """Parsing completo del testo incollato.
    Ritorna: (df_partite, team1, team2, quote)."""
    righe = [r.strip() for r in testo.splitlines()]
    righe = [r for r in righe if r and not re.search(r"mostra\s+pi", r, re.I)]

    # La sezione "Testa a Testa" (H2H) sta tra le ultime partite e le quote finali.
    # La sua ultima partita non ha la lettera esito, quindi "inghiottirebbe" le righe
    # delle quote facendole perdere. La rimuoviamo conservando le righe-quota in coda.
    idx_h2h = next((i for i, r in enumerate(righe)
                    if re.search(r"testa\s*a\s*testa|head\s*to\s*head|precedenti", r, re.I)),
                   None)
    if idx_h2h is not None:
        coda_num = []
        for r in reversed(righe[idx_h2h + 1:]):
            toks = [t.replace(",", ".") for t in r.split()]
            if toks and all(re.fullmatch(r"-?\d+(\.\d+)?", t) for t in toks):
                coda_num.insert(0, r)
            else:
                break
        righe = righe[:idx_h2h] + coda_num

    headers = [(i, _team_da_header(r)) for i, r in enumerate(righe) if _is_header(r)]
    if not headers:
        return pd.DataFrame(), None, None, {}

    sezioni = []
    for k, (idx, team) in enumerate(headers):
        start = idx + 1
        end = headers[k + 1][0] if k + 1 < len(headers) else len(righe)
        sezioni.append((team, righe[start:end]))

    tutte, ultimo_avanzo = [], []
    for _, corpo in sezioni:
        partite, avanzo = _parse_sezione(corpo)
        tutte.extend(partite)
        ultimo_avanzo = avanzo  # l'avanzo utile è quello dell'ultima sezione (le quote)

    quote = _parse_quote(ultimo_avanzo)

    # dedup su (data, casa, trasferta): la stessa partita compare in entrambe le squadre
    viste, uniche = set(), []
    for p in tutte:
        chiave = (p["data_iso"], p["casa"], p["trasferta"])
        if chiave in viste:
            continue
        viste.add(chiave)
        uniche.append(p)

    df = pd.DataFrame(uniche)
    if not df.empty:
        df = df.sort_values("data_iso", ascending=False).reset_index(drop=True)

    team1 = headers[0][1] if len(headers) >= 1 else None
    team2 = headers[1][1] if len(headers) >= 2 else None

    # L'intestazione può avere una forma diversa (es. TUTTA MAIUSCOLA) rispetto ai nomi
    # nelle partite: riporto team1/team2 esattamente alla forma usata nello storico,
    # altrimenti il motore non trova le partite della squadra ("storico insufficiente").
    if not df.empty:
        nomi = set(df["casa"]) | set(df["trasferta"])
        def _risolvi(nome):
            if not nome:
                return nome
            for n in nomi:
                if _key(n) == _key(nome):
                    return n
            return nome
        team1, team2 = _risolvi(team1), _risolvi(team2)

    return df, team1, team2, quote


# =============================================================================
#  EXPORT / COPIA
# =============================================================================
def testo_export(df, team1, team2, quote):
    out = ["PARTITE:"]
    for _, r in df.iterrows():
        gc = r.get("Gol Casa")
        gt = r.get("Gol Trasferta")
        gc = "" if pd.isna(gc) else int(gc)
        gt = "" if pd.isna(gt) else int(gt)
        ris = f"{gc}-{gt}" if gc != "" and gt != "" else "?-?"
        out.append(f"{r['Casa']} - {r['Trasferta']} | {ris}")

    if team1 and team2:
        out += ["", f"PARTITA DA PRONOSTICARE: {team1} (casa) - {team2} (trasferta)"]

    if quote:
        out.append("")
        out.append("QUOTE (iniziale → variazione):")

        def _riga(label, p1, p2, m1, m2):
            prima = f"{quote[p1]} / {quote[p2]}" if p1 in quote else None
            mod = f"{quote[m1]} / {quote[m2]}" if m1 in quote else None
            if prima and mod:
                return f"{label}: {prima}  →  {mod}"
            if prima:
                return f"{label}: {prima}"
            if mod:
                return f"{label}: (variazione) {mod}"
            return None

        if "1" in quote or "1_mod" in quote:
            prima = f"{quote.get('1')} / {quote.get('X')} / {quote.get('2')}" if "1" in quote else None
            mod = f"{quote.get('1_mod')} / {quote.get('X_mod')} / {quote.get('2_mod')}" if "1_mod" in quote else None
            if prima and mod:
                out.append(f"1X2: {prima}  →  {mod}")
            elif prima:
                out.append(f"1X2: {prima}")
            elif mod:
                out.append(f"1X2: (variazione) {mod}")
        for riga in [_riga("O2.5/U2.5", "over25", "under25", "over25_mod", "under25_mod"),
                     _riga("Goal/NoGoal", "goal", "nogoal", "goal_mod", "nogoal_mod")]:
            if riga:
                out.append(riga)
        if "forma_casa" in quote:
            out += ["", "FORMA:",
                    f"{team1 or 'Casa'} (casa): {quote['forma_casa']}",
                    f"{team2 or 'Trasferta'} (trasferta): {quote['forma_trasferta']}"]
        if "val_casa" in quote:
            out += ["", "VALORI ROSE:",
                    f"{team1 or 'Casa'} (casa): {quote['val_casa']}",
                    f"{team2 or 'Trasferta'} (trasferta): {quote['val_trasferta']}"]
    return "\n".join(out)


def to_excel(dfs: dict) -> bytes:
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as w:
        for nome, d in dfs.items():
            (d if not d.empty else pd.DataFrame({"vuoto": []})).to_excel(
                w, sheet_name=nome[:31], index=False)
    return buf.getvalue()


# Colonne partite -> intestazioni leggibili (ordine incluso) per l'export Excel
EXPORT_COLONNE = [
    ("data", "Data"), ("ora", "Ora"), ("competizione", "Competizione"),
    ("squadra_casa", "Casa"), ("squadra_trasferta", "Trasferta"),
    ("gol_casa", "Gol Casa"), ("gol_trasferta", "Gol Trasferta"),
    ("qualificatore", "Note"), ("tipo_partita", "Tipo partita"),
    ("da_compilare", "Da compilare"),
    ("quota_iniziale_1", "Quota iniziale 1"), ("quota_iniziale_x", "Quota iniziale X"),
    ("quota_iniziale_2", "Quota iniziale 2"),
    ("variazione_quota_1", "Variazione di quota 1"), ("variazione_quota_x", "Variazione di quota X"),
    ("variazione_quota_2", "Variazione di quota 2"),
    ("quota_iniziale_over", "Quota iniziale Over 2.5"),
    ("quota_iniziale_under", "Quota iniziale Under 2.5"),
    ("variazione_quota_over", "Variazione di quota Over 2.5"),
    ("variazione_quota_under", "Variazione di quota Under 2.5"),
    ("quota_iniziale_goal", "Quota iniziale Goal"), ("quota_iniziale_nogoal", "Quota iniziale NoGoal"),
    ("variazione_quota_goal", "Variazione di quota Goal"),
    ("variazione_quota_nogoal", "Variazione di quota NoGoal"),
    ("forma_casa", "Forma Casa"), ("forma_trasferta", "Forma Trasferta"),
    ("val_casa", "Valore Rosa Casa"), ("val_trasferta", "Valore Rosa Trasferta"),
    ("is_target", "Partita da pronosticare"),
    ("inserito_da", "Inserito da"), ("creato_il", "Creato il"), ("aggiornato_il", "Aggiornato il"),
]


def partite_per_export(df):
    """Riordina e rinomina le colonne partite in intestazioni leggibili per l'Excel."""
    if df.empty:
        return df
    cols = [(k, v) for k, v in EXPORT_COLONNE if k in df.columns]
    out = df[[k for k, _ in cols]].copy()
    out.columns = [v for _, v in cols]
    return out


# =============================================================================
#  DB PARTITE
# =============================================================================
def _fetch_tutte(_cli, tabella, order_col="data", desc=True):
    """Scarica TUTTE le righe di una tabella superando il limite di 1000 di Supabase,
    con paginazione a blocchi. Senza questo, le righe più vecchie sparivano.
    NON va messa in cache: riceve il client Supabase (non hashabile)."""
    righe = []
    step = 1000
    start = 0
    while True:
        q = _cli.table(tabella).select("*")
        if order_col:
            q = q.order(order_col, desc=desc)
        res = q.range(start, start + step - 1).execute()
        batch = res.data or []
        righe.extend(batch)
        if len(batch) < step:
            break
        start += step
    return righe


@st.cache_data(ttl=600, show_spinner=False)
def carica_partite():
    cli = get_client()
    if not cli:
        return pd.DataFrame()
    df = pd.DataFrame(_fetch_tutte(cli, "partite", "data", True))
    if not df.empty and "data" in df:
        df["data"] = pd.to_datetime(df["data"]).dt.date
    return df


def salva_partite(records):
    """records: lista di dict pronti per upsert (con chiave unica data+casa+trasferta)."""
    cli = get_client()
    if not cli:
        raise RuntimeError("Supabase non configurato.")
    cli.table("partite").upsert(
        records, on_conflict="data,squadra_casa,squadra_trasferta"
    ).execute()
    st.cache_data.clear()


def aggiorna_partite(records):
    """Aggiorna partite ESISTENTI per id (UPDATE reale, non upsert: tocca solo le
    colonne indicate, senza rischiare inserimenti di righe vuote)."""
    cli = get_client()
    if not cli:
        raise RuntimeError("Supabase non configurato.")
    for rec in records:
        rid = rec.get("id")
        if not rid:
            continue
        payload = {k: v for k, v in rec.items() if k != "id"}
        if payload:
            cli.table("partite").update(payload).eq("id", rid).execute()
    st.cache_data.clear()


def elimina_partite(ids):
    """Elimina le partite indicate e i dati agganciati (pronostici collegati)."""
    cli = get_client()
    if not cli:
        raise RuntimeError("Supabase non configurato.")
    for mid in ids:
        if not mid:
            continue
        try:
            cli.table("pronostici").delete().eq("partita_id", str(mid)).execute()
        except Exception:
            pass
        cli.table("partite").delete().eq("id", str(mid)).execute()
    st.cache_data.clear()


# =============================================================================
#  COMPETIZIONI (categorizzazione)
# =============================================================================
CATEGORIE = ["Non assegnata", "Campionato", "Playoff", "Coppa nazionale",
             "Coppa internazionale", "Coppa/torneo secondario", "Amichevole", "Altro"]


@st.cache_data(ttl=600, show_spinner=False)
def carica_competizioni():
    cli = get_client()
    if not cli:
        return pd.DataFrame()
    res = cli.table("competizioni").select("*").order("nome_corto").execute()
    return pd.DataFrame(res.data or [])


def upsert_competizioni(records):
    cli = get_client()
    if not cli:
        raise RuntimeError("Supabase non configurato.")
    # genera l'id lato app per le righe nuove (robusto anche se il DB non ha il default)
    for rec in records:
        if not rec.get("id"):
            rec["id"] = str(uuid.uuid4())
    cli.table("competizioni").upsert(records).execute()
    st.cache_data.clear()


def elimina_competizione(cid):
    cli = get_client()
    cli.table("competizioni").delete().eq("id", cid).execute()
    st.cache_data.clear()


# --- Calibrazione & pronostici ---
def carica_calibrazione():
    """Ritorna {'over25': {'xs':[...], 'ys':[...]}, 'goal': {...}} se presenti."""
    cli = get_client()
    if not cli:
        return {}
    try:
        rows = cli.table("calibrazione").select("*").execute().data or []
    except Exception:
        return {}
    out = {}
    for r in rows:
        try:
            out[r["mercato"]] = {"xs": json.loads(r["xs"]), "ys": json.loads(r["ys"])}
        except Exception:
            pass
    return out


def salva_calibrazione(mercato, iso, n, brier_pre, brier_post):
    cli = get_client()
    if not cli:
        raise RuntimeError("Supabase non configurato.")
    cli.table("calibrazione").upsert({
        "mercato": mercato, "xs": json.dumps(iso["xs"]), "ys": json.dumps(iso["ys"]),
        "n": n, "brier_pre": brier_pre, "brier_post": brier_post,
        "aggiornato_il": datetime.utcnow().isoformat(),
    }).execute()


def analisi_discrepanze():
    """Valida la tesi confidence/quota: per i pronostici salvati con risultato, confronta
    la probabilità statistica con l'implicita grezza (1/quota) e misura il tasso di
    successo per fascia di scarto e per livello di alert."""
    cli = get_client()
    if not cli:
        return None
    rows = cli.table("pronostici").select("*").not_.is_("gol_casa", "null").execute().data or []
    dati = []
    soglie = analisi.PESI_DEFAULT["alert_soglie"]
    for r in rows:
        q, stat = r.get("quota"), r.get("prob")
        if not q or not stat or float(q) <= 0:
            continue
        implied = 1 / float(q)
        won = _pronostico_vinto(r.get("mercato") or "", r.get("gol_casa"), r.get("gol_trasferta"))
        if won is None:
            continue
        dati.append({"delta": implied - float(stat), "stat": float(stat),
                     "implied": implied, "won": 1 if won else 0,
                     "alert": analisi.livello_alert(float(stat), implied, soglie)})
    if len(dati) < 5:
        return {"n": len(dati), "insufficiente": True}

    def aggrega(sel, nome):
        if not sel:
            return None
        n = len(sel)
        return {"gruppo": nome, "n": n,
                "tasso": sum(d["won"] for d in sel) / n,
                "stat_media": sum(d["stat"] for d in sel) / n,
                "implicita_media": sum(d["implied"] for d in sel) / n}

    fasce = [
        ("quota molto più bassa (mercato svaluta)", lambda d: d < -0.12),
        ("quota più bassa", lambda d: -0.12 <= d < -0.04),
        ("allineate", lambda d: -0.04 <= d <= 0.04),
        ("quota più alta (mercato conferma)", lambda d: 0.04 < d <= 0.12),
        ("quota molto più alta (discrepanza)", lambda d: d > 0.12),
    ]
    per_scarto = [a for nome, cond in fasce
                  if (a := aggrega([d for d in dati if cond(d["delta"])], nome))]
    per_alert = [a for liv in (None, "basso", "medio", "alto")
                 if (a := aggrega([d for d in dati if d["alert"] == liv],
                                  liv or "nessun alert"))]
    return {"n": len(dati), "per_scarto": per_scarto, "per_alert": per_alert}


def salva_pronostico(record):
    cli = get_client()
    if not cli:
        raise RuntimeError("Supabase non configurato.")
    cli.table("pronostici").insert(record).execute()
    st.cache_data.clear()


def upsert_pronostico(record):
    """Salva/aggiorna il pronostico pre-partita: uno solo per partita. Se la partita
    ha già un risultato salvato, il pronostico è 'congelato' e non viene sovrascritto.
    Resiliente: se la colonna 'riepilogo' non esiste ancora nel DB, salva senza."""
    cli = get_client()
    if not cli:
        return

    def _do(rec):
        pid = rec.get("partita_id")
        esistente = None
        if pid:
            r = cli.table("pronostici").select("id,gol_casa").eq("partita_id", pid).execute()
            esistente = (r.data or [None])[0]
        if esistente:
            if esistente.get("gol_casa") is not None:
                return
            cli.table("pronostici").update(rec).eq("id", esistente["id"]).execute()
        else:
            cli.table("pronostici").insert(rec).execute()

    _opzionali = ("scheda_json", "riepilogo", "mercato_ragionato", "score_ragionato",
                  "merc_motore", "conf_motore", "merc_statistico", "conf_statistico",
                  "merc_fusione", "conf_fusione", "merc_solo_stat", "conf_solo_stat",
                  "merc_ev", "val_ev", "pron_cristiano")
    try:
        _do(record)
    except Exception:
        # fallback progressivo: rimuovi le colonne opzionali che il DB potrebbe non avere
        rec2 = {k: v for k, v in record.items()
                if k not in ("scheda_json", "mercato_ragionato", "score_ragionato",
                             "merc_motore", "conf_motore", "merc_statistico",
                             "conf_statistico", "merc_fusione", "conf_fusione",
                             "merc_solo_stat", "conf_solo_stat", "merc_ev", "val_ev",
                             "pron_cristiano")}
        try:
            _do(rec2)
        except Exception:
            _do({k: v for k, v in record.items() if k not in _opzionali})
    st.cache_data.clear()


@st.cache_data(ttl=600, show_spinner=False)
def carica_pronostici():
    cli = get_client()
    if not cli:
        return pd.DataFrame()
    return pd.DataFrame(_fetch_tutte(cli, "pronostici", "creato_il", True))


def completa_risultati_pronostici():
    """Riempie il risultato reale nei pronostici salvati, abbinandoli alle partite
    del database (per squadre + data) che ora hanno un risultato. Ritorna n aggiornati."""
    cli = get_client()
    if not cli:
        return 0
    pron = cli.table("pronostici").select("*").is_("gol_casa", "null").execute().data or []
    if not pron:
        return 0
    part = carica_partite()
    if part.empty:
        return 0
    part = part[part["gol_casa"].notna()]
    agg = 0
    for pr in pron:
        m = part[(part["squadra_casa"] == pr["squadra_casa"]) &
                 (part["squadra_trasferta"] == pr["squadra_trasferta"])]
        if pr.get("data"):
            m2 = m[m["data"].astype(str) == str(pr["data"])]
            m = m2 if not m2.empty else m
        if m.empty:
            continue
        r = m.sort_values("data").iloc[-1]
        cli.table("pronostici").update({
            "gol_casa": int(r["gol_casa"]), "gol_trasferta": int(r["gol_trasferta"]),
        }).eq("id", pr["id"]).execute()
        agg += 1
    return agg


def analisi_quote_bookmaker():
    """Studia come ragiona il bookmaker: raggruppa i pronostici salvati (con risultato)
    per fascia di quota e mostra quanto spesso l'esito si verifica davvero."""
    cli = get_client()
    if not cli:
        return None
    rows = cli.table("pronostici").select("*").not_.is_("gol_casa", "null").execute().data or []
    dati = []
    for r in rows:
        q = r.get("quota")
        if not q:
            continue
        gc, gt = r.get("gol_casa"), r.get("gol_trasferta")
        merc = (r.get("mercato") or "")
        vinto = _pronostico_vinto(merc, gc, gt)
        if vinto is None:
            continue
        dati.append((float(q), 1 if vinto else 0))
    if len(dati) < 5:
        return None
    fasce = [("bassa (<1.5)", lambda q: q < 1.5), ("media (1.5-2.5)", lambda q: 1.5 <= q < 2.5),
             ("alta (2.5-4)", lambda q: 2.5 <= q < 4), ("altissima (>=4)", lambda q: q >= 4)]
    out = []
    for nome, cond in fasce:
        sel = [(q, v) for q, v in dati if cond(q)]
        if not sel:
            continue
        out.append({"fascia": nome, "n": len(sel),
                    "vinte": sum(v for _, v in sel),
                    "tasso": sum(v for _, v in sel) / len(sel),
                    "implicita": sum(1 / q for q, _ in sel) / len(sel)})
    return out


def _pronostico_vinto(mercato, gc, gt):
    if gc is None or gt is None:
        return None
    gc, gt = int(gc), int(gt)
    tot = gc + gt
    m = mercato.strip()
    if m.startswith("Over 2.5"):
        return tot >= 3
    if m.startswith("Under 2.5"):
        return tot < 3
    if m.startswith("Over 1.5"):
        return tot >= 2
    if m.startswith("Under 1.5"):
        return tot < 2
    if m.startswith("Over 3.5"):
        return tot >= 4
    if m.startswith("Under 3.5"):
        return tot < 4
    if m == "Goal":
        return gc > 0 and gt > 0
    if m == "No Goal":
        return not (gc > 0 and gt > 0)
    if m.startswith("1X"):
        return gc >= gt
    if m.startswith("X2"):
        return gc <= gt
    if m.startswith("12"):
        return gc != gt
    if m.startswith("1"):
        return gc > gt
    if m.startswith("X"):
        return gc == gt
    if m.startswith("2"):
        return gc < gt
    return None


def _pezzo_vinto(pezzo, gc, gt):
    """Valuta un singolo mercato, inclusi gli over/under di SQUADRA
    ('Over 0.5 casa' = casa segna >=1; 'Over 1.5 trasferta' = ospite segna >=2)."""
    if gc is None or gt is None:
        return None
    gc, gt = int(gc), int(gt)
    m = pezzo.strip().lower()
    # over/under di squadra (casa / trasferta)
    import re as _re
    mm = _re.match(r"(over|under)\s+([0-9]+(?:\.[0-9]+)?)\s+(casa|trasferta|ospite|away|home)", m)
    if mm:
        verso, soglia, chi = mm.group(1), float(mm.group(2)), mm.group(3)
        gol = gc if chi in ("casa", "home") else gt
        return gol > soglia if verso == "over" else gol < soglia
    # altrimenti mercato standard: riuso _pronostico_vinto (con la capitalizzazione originale)
    return _pronostico_vinto(pezzo.strip(), gc, gt)


def _combo_vinta(testo, gc, gt):
    """Valuta un pronostico combo (pezzi separati da '+', tutti da vincere insieme, AND).
    Ritorna True (tutti vinti), False (almeno uno perso), o None se un pezzo è ignoto."""
    if not testo or not testo.strip():
        return None
    pezzi = [p.strip() for p in testo.split("+") if p.strip()]
    if not pezzi:
        return None
    tutti_veri = True
    for p in pezzi:
        esito = _pezzo_vinto(p, gc, gt)
        if esito is None:
            return None          # un pezzo non riconosciuto -> non valutabile
        if esito is False:
            tutti_veri = False   # basta uno perso per perdere la combo
    return tutti_veri


def _txt(v):
    """Converte in stringa pulita gestendo None e NaN (mai crash su .strip())."""
    if v is None:
        return ""
    try:
        if pd.isna(v):
            return ""
    except (TypeError, ValueError):
        pass
    return str(v).strip()


def _codici_mancanti(df, comp_df):
    """Ritorna i codici competizione presenti nelle partite estratte ma NON ancora
    mappati nell'anagrafica competizioni."""
    if df is None or df.empty or "competizione" not in df.columns:
        return []
    presenti = set()
    if comp_df is not None and not comp_df.empty:
        for _, c in comp_df.iterrows():
            presenti |= _chiavi_competizione(c)
    mancanti = []
    visti = set()
    for cod in df["competizione"]:
        c = _txt(cod)
        if not c or _key(c) in visti:
            continue
        visti.add(_key(c))
        if _key(c) not in presenti:
            mancanti.append(c)
    return mancanti


def label_competizione(nome_lungo, nazione):
    """Etichetta leggibile: 'Nome | NAZIONE'."""
    nl = _txt(nome_lungo)
    na = _txt(nazione)
    if nl and na:
        return f"{nl} | {na}"
    return nl or na or ""


def _key(s):
    return re.sub(r"\s+", " ", _txt(s)).casefold()


def _chiavi_competizione(c):
    """Tutte le forme con cui una competizione può comparire nello storico:
    nome corto (SA), nome lungo da solo (Serie A) e nome lungo + nazione
    (Serie A | ITALIA). Serve ad agganciare la stessa competizione comunque sia scritta."""
    ks = {
        _key(c.get("nome_corto")),
        _key(c.get("nome_lungo")),
        _key(label_competizione(c.get("nome_lungo"), c.get("nazione"))),
    }
    ks.discard("")
    return ks


def categoria_di(codice_o_label, comp_df):
    """Data una competizione (codice breve, nome lungo o 'nome | nazione') restituisce
    la categoria assegnata, se presente e diversa da 'Non assegnata'."""
    if comp_df is None or comp_df.empty or not codice_o_label:
        return None
    k = _key(codice_o_label)
    for _, c in comp_df.iterrows():
        cat = c.get("categoria")
        if not cat or cat == "Non assegnata":
            continue
        if k in _chiavi_competizione(c):
            return cat
    return None


ND = "ND"  # categoria/tipo non determinabile


def categoria_o_nd(codice_o_label, comp_df):
    return categoria_di(codice_o_label, comp_df) or ND


def _label_da_comp(val, comp_df):
    """Traduce il valore salvato (es. codice '2L') nel nome leggibile dell'anagrafica
    ('Vtora Liga | Bulgaria'). Se non è in anagrafica, restituisce il valore così com'è."""
    v = _txt(val)
    if not v or comp_df is None or comp_df.empty:
        return v
    k = _key(v)
    for _, c in comp_df.iterrows():
        if k in _chiavi_competizione(c):
            return label_competizione(c.get("nome_lungo"), c.get("nazione")) or v
    return v


# =============================================================================
#  PARSING "ESTRATTORE RISULTATI" (risultati per competizione, senza data)
# =============================================================================
def _norm_squadra(s):
    """Toglie il codice paese '(Kaz)' e normalizza per il confronto."""
    s = re.sub(r"\(.*?\)", "", _txt(s))
    return re.sub(r"\s+", " ", s).strip()


def parse_risultati(testo):
    """Parsa il formato a blocchi per competizione.
    Ritorna lista di dict: competizione(label), nazione, casa, trasferta, qualificatore, gol_casa, gol_trasferta.
    Ogni match = squadra_casa (x2), squadra_trasferta (x2), [qualificatore], gol_casa, gol_trasferta.
    Le righe non-match (2 consecutive) sono l'intestazione competizione: nome + nazione."""
    righe = [r.strip() for r in testo.splitlines()]
    righe = [r for r in righe if r]  # via le righe vuote
    n = len(righe)

    def is_int(x):
        return re.fullmatch(r"\d+", re.sub(r"\(.*?\)", "", x).strip()) is not None

    risultati = []
    comp_corr, naz_corr = None, None
    i = 0
    while i < n:
        # blocco partita? richiede casa,casa,trasf,trasf ripetute
        if i + 5 < n and righe[i] == righe[i + 1] and righe[i + 2] == righe[i + 3]:
            casa = righe[i]
            trasf = righe[i + 2]
            j = i + 4
            qualif = None
            if j < n and not is_int(righe[j]):   # qualificatore opzionale (Dopo Suppl./Rigori)
                qualif = righe[j]
                j += 1
            if j + 1 < n and is_int(righe[j]) and is_int(righe[j + 1]):
                gc = int(re.sub(r"\(.*?\)", "", righe[j]))
                gt = int(re.sub(r"\(.*?\)", "", righe[j + 1]))
                risultati.append({
                    "competizione": label_competizione(comp_corr, naz_corr) or None,
                    "nome_lungo": comp_corr, "nazione": naz_corr,
                    "casa": casa, "trasferta": trasf,
                    "qualificatore": qualif,
                    "gol_casa": gc, "gol_trasferta": gt,
                })
                i = j + 2
                continue
        # altrimenti è un'intestazione competizione: nome + nazione
        comp_corr = righe[i]
        naz_corr = righe[i + 1] if i + 1 < n else None
        i += 2
    return risultati


def parse_pianificazione(testo):
    """Come parse_risultati ma i blocchi finiscono con l'orario (es. '18:00') invece
    dei gol. Ritorna: competizione(label), nome_lungo, nazione, casa, trasferta, ora."""
    righe = [r.strip() for r in testo.splitlines()]
    righe = [r for r in righe if r]
    n = len(righe)

    def is_time(x):
        return re.fullmatch(r"\d{1,2}:\d{2}", x.strip()) is not None

    out = []
    comp_corr, naz_corr = None, None
    i = 0
    while i < n:
        if i + 3 < n and righe[i] == righe[i + 1] and righe[i + 2] == righe[i + 3]:
            casa, trasf = righe[i], righe[i + 2]
            j = i + 4
            ora = None
            if j < n and is_time(righe[j]):
                ora = righe[j]
                j += 1
            out.append({
                "competizione": label_competizione(comp_corr, naz_corr) or None,
                "nome_lungo": comp_corr, "nazione": naz_corr,
                "casa": casa, "trasferta": trasf, "ora": ora,
            })
            i = j
            continue
        comp_corr = righe[i]
        naz_corr = righe[i + 1] if i + 1 < n else None
        i += 2
    return out


def _similarita(a, b):
    return SequenceMatcher(None, _key(a), _key(b)).ratio()


# =============================================================================
#  AUTENTICAZIONE
# =============================================================================
def hash_pw(pw):
    return bcrypt.hashpw(pw.encode(), bcrypt.gensalt()).decode()


def check_pw(pw, h):
    try:
        return bcrypt.checkpw(pw.encode(), h.encode())
    except Exception:
        return False


@st.cache_data(ttl=600, show_spinner=False)
def carica_utenti():
    cli = get_client()
    if not cli:
        return []
    return cli.table("utenti").select("*").execute().data or []


def crea_utente(username, password, ruolo="user"):
    cli = get_client()
    cli.table("utenti").insert({
        "username": username.strip(),
        "password_hash": hash_pw(password),
        "ruolo": ruolo,
    }).execute()
    st.cache_data.clear()


def _autofill_hint(coppie):
    """Marca i campi (per aria-label) con l'attributo autocomplete giusto, così il
    gestore password di iPhone/Android riempie username e password nei punti corretti."""
    js = (
        "const doc=window.parent.document;"
        "function s(){let ok=true;const m=" + json.dumps(coppie) + ";"
        "m.forEach(function(p){var e=doc.querySelector('input[aria-label=\"'+p[0]+'\"]');"
        "if(e){e.setAttribute('autocomplete',p[1]);}else{ok=false;}});return ok;}"
        "let n=0;const iv=setInterval(function(){if(s()||n++>25)clearInterval(iv);},150);"
    )
    components_html("<script>" + js + "</script>", height=0)


def login_gate():
    """Ritorna l'utente loggato oppure blocca l'app mostrando login/bootstrap."""
    if "user" in st.session_state:
        return st.session_state["user"]

    if not supabase_pronto():
        st.warning("⚙️ Supabase non è ancora configurato. "
                   "Aggiungi le credenziali in `.streamlit/secrets.toml` "
                   "(vedi `secrets.toml.example`).")
        st.stop()

    if bcrypt is None:
        st.error("Manca la libreria `bcrypt`. Aggiungila a requirements.txt.")
        st.stop()

    utenti = carica_utenti()

    # Primo avvio: nessun utente -> crea il primo admin
    if not utenti:
        st.title("👤 Primo avvio — crea l'amministratore")
        with st.form("bootstrap"):
            u = st.text_input("Username")
            p1 = st.text_input("Password", type="password")
            p2 = st.text_input("Ripeti password", type="password")
            if st.form_submit_button("Crea admin"):
                if not u or not p1:
                    st.error("Username e password obbligatori.")
                elif p1 != p2:
                    st.error("Le password non coincidono.")
                else:
                    crea_utente(u, p1, ruolo="admin")
                    st.success("Admin creato! Ora effettua il login.")
                    st.rerun()
        _autofill_hint([["Username", "username"], ["Password", "new-password"],
                        ["Ripeti password", "new-password"]])
        st.stop()

    # Login normale
    st.title("🔐 Accesso")
    with st.form("login"):
        u = st.text_input("Username")
        p = st.text_input("Password", type="password")
        if st.form_submit_button("Entra"):
            match = next((x for x in utenti if x["username"] == u.strip()), None)
            if match and check_pw(p, match["password_hash"]):
                st.session_state["user"] = {"username": match["username"], "ruolo": match["ruolo"]}
                st.rerun()
            else:
                st.error("Credenziali non valide.")
    _autofill_hint([["Username", "username"], ["Password", "current-password"]])
    st.stop()


# =============================================================================
#  PAGINA: ESTRATTORE
# =============================================================================
def _reset_estrattore():
    """Pulisce i campi della maschera estrattore per inserire una nuova partita,
    restando sulla pagina. Chiamata come callback (on_click), così azzera i widget
    prima che vengano ricreati. Tiene la competizione selezionata (spesso è la stessa)."""
    for k in ("editor_estrattore",
              "q1", "q1m", "q2", "q2m", "qx", "qxm",
              "qo", "qom", "qu", "qum", "qg", "qgm", "qn", "qnm",
              "fc", "ft", "vc", "vt"):
        st.session_state.pop(k, None)
    # il campo testo usa una chiave dinamica: cambiando il nonce nasce un campo VUOTO
    st.session_state["_estr_nonce"] = st.session_state.get("_estr_nonce", 0) + 1
    st.session_state.pop("_estr_quote_sig", None)
    st.session_state.pop("_estr_salvato", None)


def pagina_estrattore(user):
    st.header("📥 Ultimi risultati e quote")
    st.caption("Incolla i blocchi 'ULTIMI INCONTRI' di 2 squadre alla volta, poi rivedi e salva.")

    # Prima informazione (opzionale): competizione/tipo della partita da pronosticare
    comp_df_estr = carica_competizioni()
    opz_estr = {"— non specificata —": None}
    if not comp_df_estr.empty:
        for _, cr in comp_df_estr.sort_values("nome_lungo", na_position="last").iterrows():
            lab = label_competizione(cr.get("nome_lungo"), cr.get("nazione")) or _txt(cr.get("nome_corto"))
            if lab:
                opz_estr[lab] = _txt(cr.get("nome_corto")) or lab
    comp_target_lab = st.selectbox(
        "Competizione della partita da pronosticare (opzionale)", list(opz_estr.keys()),
        key="estr_comp_target",
        help="Campionato, coppa o amichevole della partita che vuoi pronosticare. "
             "Determina categoria e livello usati dal motore. Puoi lasciarla non specificata.")
    comp_target = opz_estr.get(comp_target_lab)

    _nonce = st.session_state.get("_estr_nonce", 0)
    testo = st.text_area("Incolla qui il testo", height=260, key=f"testo_estrattore_{_nonce}")

    if not testo.strip():
        st.info("In attesa del testo…")
        return

    df, team1, team2, quote = parse_incontri(testo)

    # Se il testo incollato cambia (es. incolli prima i risultati e pochi secondi dopo
    # aggiungi le quote), riallinea i campi quota/forma/rose ai NUOVI valori letti.
    # Altrimenti Streamlit terrebbe "congelato" il valore comparso la prima volta.
    _sig = json.dumps(quote, sort_keys=True, default=str)
    if st.session_state.get("_estr_quote_sig") != _sig:
        st.session_state["_estr_quote_sig"] = _sig
        _mappa = {"q1": "1", "qx": "X", "q2": "2", "qo": "over25", "qu": "under25",
                  "qg": "goal", "qn": "nogoal", "fc": "forma_casa", "ft": "forma_trasferta",
                  "vc": "val_casa", "vt": "val_trasferta"}
        for _wk, _qk in _mappa.items():
            val = quote.get(_qk)
            st.session_state[_wk] = "" if val in (None, "") else str(val)

    if df.empty:
        st.warning("Nessuna partita riconosciuta. Controlla che ci sia la riga 'ULTIMI INCONTRI:'.")
        return

    st.success(f"Riconosciute **{len(df)}** partite uniche"
               + (f" · {team1} (casa) vs {team2} (trasferta)" if team1 and team2 else ""))

    # --- Campionati mancanti dall'anagrafica: segnalali e falli compilare al volo ---
    comp_df_estr = carica_competizioni()
    mancanti = _codici_mancanti(df, comp_df_estr)
    if mancanti:
        st.warning(f"⚠️ {len(mancanti)} campionato/i non ancora in anagrafica: "
                   f"{', '.join(mancanti)}. Compilali qui sotto: verranno salvati in "
                   "Configurazione e usati subito dal motore.")
        tab_nuovi = pd.DataFrame({
            "Nome corto": mancanti,                       # autocompilato dal codice mancante
            "Nome lungo": ["" for _ in mancanti],
            "Nazione": ["" for _ in mancanti],
            "Categoria": ["Campionato" for _ in mancanti],
            "Livello": [pd.NA for _ in mancanti],
        })
        ed_nuovi = st.data_editor(
            tab_nuovi, use_container_width=True, hide_index=True, key="editor_comp_mancanti",
            column_config={
                "Nome corto": st.column_config.TextColumn(
                    "Nome corto", help="Codice riconosciuto dal testo (già compilato)."),
                "Categoria": st.column_config.SelectboxColumn("Categoria", options=CATEGORIE),
                "Livello": st.column_config.NumberColumn(
                    "Livello", min_value=1, step=1,
                    help="1 = prima divisione, 2 = seconda… (per amichevoli/coppe lascia vuoto)."),
            })
        if st.button("💾 Salva campionati mancanti", type="primary", key="salva_comp_mancanti"):
            recs = []
            for _, rr in ed_nuovi.iterrows():
                corto = _txt(rr["Nome corto"])
                if not corto:
                    continue
                liv = rr["Livello"]
                rec = {
                    "nome_corto": corto,
                    "nome_lungo": _txt(rr["Nome lungo"]) or None,
                    "nazione": _txt(rr["Nazione"]) or None,
                    "categoria": _txt(rr["Categoria"]) or "Non assegnata",
                }
                try:
                    if liv is not None and not pd.isna(liv):
                        rec["livello"] = int(liv)
                except (TypeError, ValueError):
                    pass
                recs.append(rec)
            if recs:
                try:
                    upsert_competizioni(recs)
                    st.success(f"Salvati {len(recs)} campionati. Ora compaiono in Configurazione "
                               "e sono usati dal motore.")
                    st.rerun()
                except Exception as e:
                    st.error(f"Errore nel salvataggio: {e}")
            else:
                st.info("Nessun campionato da salvare.")

    # tabella modificabile
    vista = pd.DataFrame({
        "Data": df["data_iso"],
        "Competizione": df["competizione"],
        "Casa": df["casa"],
        "Trasferta": df["trasferta"],
        "Gol Casa": df["gol_casa"].astype("Int64"),
        "Gol Trasferta": df["gol_trasferta"].astype("Int64"),
        "Note": df["qualificatore"],
    })
    edit = st.data_editor(
        vista, use_container_width=True, num_rows="dynamic", key="editor_estrattore",
        column_config={
            "Data": st.column_config.DateColumn(format="DD.MM.YY"),
            "Gol Casa": st.column_config.NumberColumn(min_value=0, step=1),
            "Gol Trasferta": st.column_config.NumberColumn(min_value=0, step=1),
        },
    )

    # --- quote / valori rose (modificabili) ---
    with st.expander("Quote e valori rose (partita da pronosticare)", expanded=bool(quote)):
        st.markdown("**1X2**")
        st.caption("Quota iniziale (dallo scraper, modificabile)")
        c = st.columns(3)
        q1 = c[0].text_input("1", key="q1")
        qx = c[1].text_input("X", key="qx")
        q2 = c[2].text_input("2", key="q2")
        st.caption("Variazione di quota (manuale)")
        c = st.columns(3)
        q1m = c[0].text_input("1 ", "", key="q1m")
        qxm = c[1].text_input("X ", "", key="qxm")
        q2m = c[2].text_input("2 ", "", key="q2m")

        st.markdown("**Over / Under 2.5**")
        st.caption("Quota iniziale")
        c = st.columns(2)
        qo = c[0].text_input("Over 2.5", key="qo")
        qu = c[1].text_input("Under 2.5", key="qu")
        st.caption("Variazione di quota")
        c = st.columns(2)
        qom = c[0].text_input("Over 2.5 ", "", key="qom")
        qum = c[1].text_input("Under 2.5 ", "", key="qum")

        st.markdown("**Goal / NoGoal**")
        st.caption("Quota iniziale")
        c = st.columns(2)
        qg = c[0].text_input("Goal", key="qg")
        qn = c[1].text_input("NoGoal", key="qn")
        st.caption("Variazione di quota")
        c = st.columns(2)
        qgm = c[0].text_input("Goal ", "", key="qgm")
        qnm = c[1].text_input("NoGoal ", "", key="qnm")

        st.markdown("**Forma e valore rose**")
        c = st.columns(2)
        fc = c[0].text_input(f"Forma {team1 or 'casa'}", key="fc")
        ft = c[1].text_input(f"Forma {team2 or 'trasferta'}", key="ft")
        c = st.columns(2)
        vc = c[0].text_input(f"Valore rosa {team1 or 'casa'}", key="vc")
        vt = c[1].text_input(f"Valore rosa {team2 or 'trasferta'}", key="vt")

    quote_edit = {}
    for k, v in [("1", q1), ("X", qx), ("2", q2),
                 ("1_mod", q1m), ("X_mod", qxm), ("2_mod", q2m),
                 ("over25", qo), ("under25", qu),
                 ("over25_mod", qom), ("under25_mod", qum),
                 ("goal", qg), ("nogoal", qn),
                 ("goal_mod", qgm), ("nogoal_mod", qnm),
                 ("forma_casa", fc), ("forma_trasferta", ft),
                 ("val_casa", vc), ("val_trasferta", vt)]:
        vv = _num(v)
        if vv is not None:
            quote_edit[k] = float(vv)

    # --- copia ---
    st.subheader("📋 Copia (partite + quote + forma + valori rose)")
    st.caption("Usa l'icona in alto a destra del riquadro per copiare tutto.")
    st.code(testo_export(edit, team1, team2, quote_edit), language="text")

    # --- salvataggio ---
    salva_target = st.checkbox(
        f"Crea anche la partita da pronosticare: {team1} - {team2} (con le quote)",
        value=True, disabled=not (team1 and team2),
        help="Oltre allo storico delle due squadre, crea la partita vera e propria che vuoi "
             "pronosticare, con le quote agganciate. Serve per ritrovarla in 🔮 Analisi e "
             "tra le partite che segui nel Database. Togli la spunta solo se vuoi salvare "
             "unicamente lo storico.")

    if st.button("💾 Salva nel database", type="primary"):
        if not supabase_pronto():
            st.error("Supabase non configurato.")
            return
        records = []
        for _, r in edit.iterrows():
            if not r["Casa"] or not r["Trasferta"] or pd.isna(r["Data"]):
                continue
            records.append({
                "data": str(r["Data"]),
                "competizione": r["Competizione"],
                "squadra_casa": r["Casa"],
                "squadra_trasferta": r["Trasferta"],
                "gol_casa": None if pd.isna(r["Gol Casa"]) else int(r["Gol Casa"]),
                "gol_trasferta": None if pd.isna(r["Gol Trasferta"]) else int(r["Gol Trasferta"]),
                "qualificatore": None if pd.isna(r["Note"]) else r["Note"],
                "inserito_da": user["username"],
                "aggiornato_il": datetime.utcnow().isoformat(),
            })

        # salva subito lo storico
        try:
            if records:
                salva_partite(records)
        except Exception as e:
            st.error(f"Errore nel salvataggio: {e}")
            return

        n_salvate = len(records)

        if salva_target and team1 and team2:
            payload = {
                "is_target": True, "da_compilare": False,
                "quota_iniziale_1": quote_edit.get("1"), "quota_iniziale_x": quote_edit.get("X"),
                "quota_iniziale_2": quote_edit.get("2"),
                "variazione_quota_1": quote_edit.get("1_mod"),
                "variazione_quota_x": quote_edit.get("X_mod"),
                "variazione_quota_2": quote_edit.get("2_mod"),
                "quota_iniziale_over": quote_edit.get("over25"),
                "quota_iniziale_under": quote_edit.get("under25"),
                "variazione_quota_over": quote_edit.get("over25_mod"),
                "variazione_quota_under": quote_edit.get("under25_mod"),
                "quota_iniziale_goal": quote_edit.get("goal"),
                "quota_iniziale_nogoal": quote_edit.get("nogoal"),
                "variazione_quota_goal": quote_edit.get("goal_mod"),
                "variazione_quota_nogoal": quote_edit.get("nogoal_mod"),
                "forma_casa": quote_edit.get("forma_casa"),
                "forma_trasferta": quote_edit.get("forma_trasferta"),
                "val_casa": quote_edit.get("val_casa"), "val_trasferta": quote_edit.get("val_trasferta"),
                "inserito_da": user["username"],
                "aggiornato_il": datetime.utcnow().isoformat(),
            }
            # competizione scelta in alto (opzionale): imposta anche la categoria
            if comp_target:
                payload["competizione"] = comp_target
                payload["tipo_partita"] = categoria_o_nd(comp_target, comp_df_estr)
            # cerca una riga pianificata/pending con le stesse squadre (per non duplicare)
            part_now = carica_partite()
            esistente = None
            if not part_now.empty:
                nc, nt = _key(_norm_squadra(team1)), _key(_norm_squadra(team2))
                cand = part_now[
                    (part_now["squadra_casa"].map(lambda x: _key(_norm_squadra(x))) == nc) &
                    (part_now["squadra_trasferta"].map(lambda x: _key(_norm_squadra(x))) == nt)]
                if not cand.empty:
                    # riusa SOLO una partita pianificata o una fixture-target ancora
                    # SENZA risultato — mai una partita storica già giocata (altrimenti
                    # la nuova sfida erediterebbe data e risultato di quella vecchia).
                    dac = (cand["da_compilare"] == True) if "da_compilare" in cand \
                        else pd.Series(False, index=cand.index)
                    tgt = (cand["is_target"] == True) if "is_target" in cand \
                        else pd.Series(False, index=cand.index)
                    senza_ris = cand["gol_casa"].isna() | cand["gol_trasferta"].isna()
                    riusa = cand[(dac | tgt) & senza_ris]
                    esistente = riusa.iloc[0] if not riusa.empty else None
            try:
                if esistente is not None:
                    upd = dict(payload)
                    upd["id"] = esistente["id"]   # mantiene data pianificata, aggancia le quote
                    aggiorna_partite([upd])
                else:
                    rec = dict(payload)
                    rec["data"] = str(date.today())
                    rec["squadra_casa"], rec["squadra_trasferta"] = team1, team2
                    salva_partite([rec])
                n_salvate += 1
            except Exception as e:
                st.error(f"Errore nel salvataggio della partita da pronosticare: {e}")
                return

        st.success(f"Salvate {n_salvate} righe. Le vedi nella sezione Database.")
        st.cache_data.clear()
        st.session_state["_estr_salvato"] = True

    # dopo un salvataggio riuscito: pulsante per inserire subito una nuova partita
    if st.session_state.get("_estr_salvato"):
        st.divider()
        st.caption("Vuoi inserire un'altra partita? Pulisci la maschera e riparti da capo.")
        st.button("➕ Aggiungi nuova partita", type="primary",
                  on_click=_reset_estrattore, key="btn_nuova_partita")


# =============================================================================
#  PAGINA: DATABASE
# =============================================================================
def genera_docx_archivio(df, comp_df, con_analisi=True):
    """Crea un archivio Word: una pagina per fixture (partita seguita) con intestazione,
    ultime partite delle due squadre, (opzionale) analisi & pronostico, e risultato."""
    try:
        from docx import Document
    except Exception:
        raise RuntimeError("Libreria python-docx non disponibile. "
                           "Aggiungi 'python-docx' a requirements.txt.")
    doc = Document()
    fx = df[df["is_target"] == True] if "is_target" in df.columns else df.iloc[0:0]
    if "data" in fx.columns:
        fx = fx.sort_values("data", ascending=False)
    if fx.empty:
        doc.add_heading("Archivio partite", level=1)
        doc.add_paragraph("Nessuna partita seguita da esportare.")
        bio = io.BytesIO(); doc.save(bio); return bio.getvalue()

    pron = carica_pronostici()
    calibr = carica_calibrazione()
    first = True
    for _, row in fx.iterrows():
        if not first:
            doc.add_page_break()
        first = False
        home, away = row["squadra_casa"], row["squadra_trasferta"]
        doc.add_heading(f"{home} - {away}", level=1)
        data = row["data"].strftime("%d.%m.%Y") if hasattr(row["data"], "strftime") else _txt(row.get("data"))
        meta_parts = []
        if _txt(row.get("competizione")):
            meta_parts.append(f"Campionato: {_label_da_comp(row.get('competizione'), comp_df)}")
        if data:
            meta_parts.append(f"Data: {data}")
        if _txt(row.get("ora")):
            meta_parts.append(f"Ora: {_txt(row.get('ora'))}")
        if meta_parts:
            doc.add_paragraph().add_run("    ".join(meta_parts)).italic = True

        doc.add_heading(f"Ultime partite {home}", level=2)
        for r in _ultime_partite_testo(df, home, escludi_id=row.get("id")):
            doc.add_paragraph(r, style="List Bullet")
        doc.add_heading(f"Ultime partite {away}", level=2)
        for r in _ultime_partite_testo(df, away, escludi_id=row.get("id")):
            doc.add_paragraph(r, style="List Bullet")

        # snapshot analisi salvato (per quote + mercati completi)
        snap = None
        rec = None
        if not pron.empty and "partita_id" in pron.columns:
            m = pron[pron["partita_id"] == str(row["id"])]
            if not m.empty:
                rec = m.iloc[0]
                if "scheda_json" in rec.index and _txt(rec.get("scheda_json")):
                    try:
                        snap = json.loads(rec["scheda_json"])
                    except Exception:
                        snap = None
        finita = _num_ok(row.get("gol_casa")) and _num_ok(row.get("gol_trasferta"))
        # se non ho snapshot e la partita è futura, calcolo dal vivo
        live = None
        if snap is None and not finita:
            try:
                live, _h, _a, _o = _analizza_row(df, row, comp_df, calibratori=calibr)
            except Exception:
                live = None
        dati = snap or live

        # QUOTE (in entrambe le versioni del Word)
        odds = (dati.get("odds") if dati else None) or _eff_odds(row)
        if odds:
            doc.add_heading("Quote", level=2)
            etich = [("1", "1"), ("X", "X"), ("2", "2"), ("over25", "Over 2.5"),
                     ("under25", "Under 2.5"), ("goal", "Goal"), ("nogoal", "No Goal")]
            righe_q = [f"{lab}: {odds[k]}" for k, lab in etich if odds.get(k)]
            if righe_q:
                doc.add_paragraph("   ·   ".join(righe_q))

        if con_analisi:
            doc.add_heading("Analisi e pronostico", level=2)
            if dati:
                p = dati["prob"]
                b = dati["best"]
                conf = f" (confidence {b['confidence']:.0f}/100)" if b.get("confidence") is not None else ""
                doc.add_paragraph().add_run(f"Pronostico: {b['mercato']}{conf}").bold = True
                doc.add_paragraph(
                    f"1 {p['1']*100:.0f}%   X {p['X']*100:.0f}%   2 {p['2']*100:.0f}%"
                    f"      Over 2.5 {dati['over_prob']*100:.0f}%   Goal {dati['btts_prob']*100:.0f}%")
                if p.get("lambda_home") is not None:
                    doc.add_paragraph(
                        f"Gol attesi: {home} {p['lambda_home']:.2f} · {away} {p['lambda_away']:.2f}")
                if p.get("risultati"):
                    doc.add_paragraph("Risultati più probabili: " +
                                      " · ".join(f"{r['risultato']} ({r['p']*100:.0f}%)"
                                                 for r in p["risultati"][:6]))
                # TUTTI I MERCATI in tabella
                mercati = sorted(dati["mercati"], key=lambda x: -x["confidence"])
                if mercati:
                    doc.add_paragraph().add_run("Tutti i mercati:").bold = True
                    tab = doc.add_table(rows=1, cols=4)
                    tab.style = "Light Grid Accent 1"
                    hdr = tab.rows[0].cells
                    hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = \
                        "Mercato", "Confidence", "Statistica", "Quota / alert"
                    for mm in mercati:
                        cells = tab.add_row().cells
                        cells[0].text = str(mm["mercato"])
                        cells[1].text = f"{mm['confidence']:.0f}/100"
                        cells[2].text = f"{mm['prob']*100:.0f}%"
                        q = ""
                        if mm.get("market_prob") is not None:
                            q = f"{mm['market_prob']*100:.0f}%"
                            if mm.get("quota"):
                                q += f" @ {mm['quota']:.2f}"
                            if mm.get("alert"):
                                q += f"  ⚠{mm['alert']}"
                        cells[3].text = q
                if dati.get("alerts"):
                    doc.add_paragraph().add_run("Alert quota:").bold = True
                    for al in dati["alerts"]:
                        doc.add_paragraph(
                            f"{al['mercato']}: {al['livello']} "
                            f"(stat {al['prob']*100:.0f}% vs quota {al['market_prob']*100:.0f}%)",
                            style="List Bullet")
            elif finita:
                doc.add_paragraph("Nessuna scheda pre-partita salvata per questa partita.")
            else:
                doc.add_paragraph("Analisi non disponibile.")

        doc.add_heading("Risultato", level=2)
        if _num_ok(row.get("gol_casa")) and _num_ok(row.get("gol_trasferta")):
            doc.add_paragraph(f"{int(row['gol_casa'])} - {int(row['gol_trasferta'])}")
        else:
            doc.add_paragraph("In attesa")

    bio = io.BytesIO(); doc.save(bio); return bio.getvalue()


def genera_docx_nuova_analisi(df, comp_df):
    """Archivio Word con la NUOVA analisi ragionata (evidenze + convergenze + signal),
    una pagina per partita seguita, più i dati soliti (ultime partite, quote, risultato)."""
    try:
        from docx import Document
    except Exception:
        raise RuntimeError("Libreria python-docx non disponibile. "
                           "Aggiungi 'python-docx' a requirements.txt.")
    doc = Document()
    fx = df[df["is_target"] == True] if "is_target" in df.columns else df.iloc[0:0]
    if "data" in fx.columns:
        fx = fx.sort_values("data", ascending=False)
    if fx.empty:
        doc.add_heading("Archivio partite — nuova analisi", level=1)
        doc.add_paragraph("Nessuna partita seguita da esportare.")
        bio = io.BytesIO(); doc.save(bio); return bio.getvalue()

    first = True
    for _, row in fx.iterrows():
        if not first:
            doc.add_page_break()
        first = False
        home, away = row["squadra_casa"], row["squadra_trasferta"]
        doc.add_heading(f"{home} - {away}", level=1)
        data = row["data"].strftime("%d.%m.%Y") if hasattr(row["data"], "strftime") else _txt(row.get("data"))
        meta = []
        if _txt(row.get("competizione")):
            meta.append(f"Campionato: {_label_da_comp(row.get('competizione'), comp_df)}")
        if data:
            meta.append(f"Data: {data}")
        if _txt(row.get("ora")):
            meta.append(f"Ora: {_txt(row.get('ora'))}")
        if meta:
            doc.add_paragraph().add_run("    ".join(meta)).italic = True

        doc.add_heading(f"Ultime partite {home}", level=2)
        for r in _ultime_partite_testo(df, home, escludi_id=row.get("id")):
            doc.add_paragraph(r, style="List Bullet")
        doc.add_heading(f"Ultime partite {away}", level=2)
        for r in _ultime_partite_testo(df, away, escludi_id=row.get("id")):
            doc.add_paragraph(r, style="List Bullet")

        odds = _eff_odds(row)
        if odds:
            doc.add_heading("Quote", level=2)
            etich = [("1", "1"), ("X", "X"), ("2", "2"), ("over25", "Over 2.5"),
                     ("under25", "Under 2.5"), ("goal", "Goal"), ("nogoal", "No Goal")]
            righe_q = [f"{lab}: {odds[k]}" for k, lab in etich if odds.get(k)]
            if righe_q:
                doc.add_paragraph("   ·   ".join(righe_q))

        # analisi ragionata (pre-partita, filtrata per data)
        doc.add_heading("Analisi ragionata", level=2)
        racc = analisi_ragionata(df, home, away, data_partita=row.get("data"),
                                 odds=odds, variazioni=_variazioni_da_row(row),
                                 escludi_id=row.get("id"),
                                 competizione=_label_da_comp(row.get("competizione"), comp_df))
        if not racc:
            doc.add_paragraph("Storico insufficiente per l'analisi.")
        else:
            doc.add_paragraph().add_run(f"Pronostico: {racc['pronostico']['testo']}").bold = True
            for sez in racc["sezioni"]:
                doc.add_heading(sez["titolo"], level=3)
                for r in sez["righe"]:
                    doc.add_paragraph(r, style="List Bullet")

        doc.add_heading("Risultato", level=2)
        if _num_ok(row.get("gol_casa")) and _num_ok(row.get("gol_trasferta")):
            doc.add_paragraph(f"{int(row['gol_casa'])} - {int(row['gol_trasferta'])}")
        else:
            doc.add_paragraph("In attesa")

    bio = io.BytesIO(); doc.save(bio); return bio.getvalue()


def genera_docx_mercati(df, comp_df):
    """Word SNELLO: solo 'squadra A - squadra B', le sezioni mercati (Over/Under,
    Goal/No Goal, 1X2 con signal/prob/quota/EV/edge) e il risultato finale."""
    try:
        from docx import Document
    except Exception:
        raise RuntimeError("Libreria python-docx non disponibile.")
    doc = Document()
    fx = df[df["is_target"] == True] if "is_target" in df.columns else df.iloc[0:0]
    if "data" in fx.columns:
        fx = fx.sort_values("data", ascending=False)
    if fx.empty:
        doc.add_heading("Mercati per partita", level=1)
        doc.add_paragraph("Nessuna partita seguita da esportare.")
        bio = io.BytesIO(); doc.save(bio); return bio.getvalue()

    SEZIONI_MERCATI = {"Over / Under", "Goal / No Goal", "1X2"}
    first = True
    for _, row in fx.iterrows():
        if not first:
            doc.add_page_break()
        first = False
        home, away = row["squadra_casa"], row["squadra_trasferta"]
        doc.add_heading(f"{home} - {away}", level=1)

        odds = _eff_odds(row)
        racc = analisi_ragionata(df, home, away, data_partita=row.get("data"),
                                 odds=odds, variazioni=_variazioni_da_row(row),
                                 escludi_id=row.get("id"),
                                 competizione=_label_da_comp(row.get("competizione"), comp_df))
        if not racc:
            doc.add_paragraph("Storico insufficiente per l'analisi.")
        else:
            for sez in racc["sezioni"]:
                if sez["titolo"] in SEZIONI_MERCATI:
                    doc.add_heading(sez["titolo"], level=2)
                    for r in sez["righe"]:
                        doc.add_paragraph(r, style="List Bullet")

        doc.add_heading("Risultato", level=2)
        if _num_ok(row.get("gol_casa")) and _num_ok(row.get("gol_trasferta")):
            doc.add_paragraph(f"{int(row['gol_casa'])} - {int(row['gol_trasferta'])}")
        else:
            doc.add_paragraph("In attesa")

    bio = io.BytesIO(); doc.save(bio); return bio.getvalue()


def _snapshot_analisi(a, odds):
    """Cattura l'analisi COMPLETA in formato dati (JSON-safe) per riproporla identica
    dopo la partita, senza ricalcolare."""
    p = a["prob"]
    return {
        "prob": {"1": p["1"], "X": p["X"], "2": p["2"],
                 "lambda_home": p.get("lambda_home"), "lambda_away": p.get("lambda_away"),
                 "risultati": [{"risultato": r["risultato"], "p": r["p"]}
                               for r in p.get("risultati", [])[:8]]},
        "over_prob": a["over_prob"], "btts_prob": a["btts_prob"],
        "over_prob_raw": a.get("over_prob_raw"), "btts_prob_raw": a.get("btts_prob_raw"),
        "elo": a.get("elo", {}), "risultati_per_esito": a.get("risultati_per_esito", {}),
        "griglia_coerente": a.get("griglia_coerente"),
        "blended_mercato": a.get("blended_mercato"), "calibrato": a.get("calibrato"),
        "reasons": list(a.get("reasons", [])), "risks": list(a.get("risks", [])),
        "alerts": a.get("alerts", []),
        "best": {k: a["best"].get(k) for k in
                 ("mercato", "confidence", "prob", "market_prob", "alert", "signal_ratio")},
        "mercati": [{k: m.get(k) for k in ("mercato", "gruppo", "prob", "confidence",
                                           "market_prob", "quota", "alert", "var_quota")}
                    for m in a["mercati"]],
        "odds": odds or {},
    }


def render_scheda_st(a, home, away):
    """Disegna la scheda COMPLETA (miglior pronostico, probabilità, alert, tutti i mercati).
    Accetta sia l'analisi dal vivo sia uno snapshot salvato: stesse chiavi."""
    p = a["prob"]
    best = a["best"]
    col = _col_conf(best["confidence"])
    reasons = "".join(f'<div style="color:{COL["hi"]};font-size:13px;margin-top:3px;">✓ {_esc(x)}</div>'
                      for x in a.get("reasons", [])) or \
        f'<div style="color:{COL["lo"]};font-size:13px;">—</div>'
    risks = "".join(f'<div style="color:{COL["loss"]};font-size:13px;margin-top:3px;">⚠ {_esc(x)}</div>'
                    for x in a.get("risks", [])) or \
        f'<div style="color:{COL["lo"]};font-size:13px;">nessun rischio rilevante</div>'
    consenso = ""
    if best.get("market_prob") is not None:
        acol = {"basso": COL["draw"], "medio": "#e08a2b", "alto": COL["loss"]}.get(best.get("alert"))
        base = (f'<div style="color:{COL["lo"]};font-size:12px;margin-top:10px;">'
                f'Statistiche {best["prob"]*100:.0f}% · quota (grezza) {best["market_prob"]*100:.0f}%')
        if best.get("alert") and acol:
            base += f' — <span style="color:{acol};font-weight:600;">alert {best["alert"]}</span>'
        consenso = base + "</div>"
    sr = best.get("signal_ratio")
    barra_sr = _barra(sr * 100, col) if sr is not None else ""
    st.html(
        f'<div style="{FONT}max-width:520px;background:{COL["panel"]};border:1px solid {COL["line"]};'
        f'border-left:4px solid {col};border-radius:14px;padding:16px;">'
        f'<div style="color:{COL["lo"]};font-size:11px;text-transform:uppercase;letter-spacing:.15em;">'
        f'Miglior pronostico</div>'
        f'<div style="display:flex;justify-content:space-between;align-items:baseline;margin-top:4px;">'
        f'<span style="color:{COL["hi"]};font-size:26px;font-weight:700;">{_esc(best["mercato"])}</span>'
        f'<span style="color:{col};font-size:22px;font-weight:700;">{best["confidence"]:.0f}<span '
        f'style="font-size:13px;color:{COL["lo"]};">/100</span></span></div>'
        f'<div style="margin-top:12px;">{reasons}</div>'
        f'<div style="margin-top:10px;">{risks}</div>{consenso}</div>')

    st.subheader("📈 Probabilità del modello")
    cc = st.columns(3)
    cc[0].metric(f"1 {home}", f'{p["1"]*100:.0f}%')
    cc[1].metric("X", f'{p["X"]*100:.0f}%')
    cc[2].metric(f"2 {away}", f'{p["2"]*100:.0f}%')
    cc = st.columns(2)
    cc[0].metric("Over 2.5", f'{a["over_prob"]*100:.0f}%')
    cc[1].metric("Goal", f'{a["btts_prob"]*100:.0f}%')
    if p.get("lambda_home") is not None:
        elo = a.get("elo", {})
        extra = (f' · Elo {elo["home"]:.0f} vs {elo["away"]:.0f}'
                 if elo.get("home") is not None else "")
        st.caption(f'Gol attesi: {home} {p["lambda_home"]:.2f} · {away} {p["lambda_away"]:.2f}{extra}')
    if p.get("risultati"):
        st.markdown("**Risultati più probabili:** " +
                    " · ".join(f'{r["risultato"]} ({r["p"]*100:.0f}%)' for r in p["risultati"][:6]))

    if a.get("alerts"):
        st.subheader("🚨 Alert quota")
        acol = {"basso": COL["draw"], "medio": "#e08a2b", "alto": COL["loss"]}
        rows = ""
        for al in a["alerts"]:
            colr = acol.get(al["livello"], COL["lo"])
            verso = "quota più alta" if (al.get("delta") or 0) > 0 else "quota più bassa"
            rows += (
                f'<div style="{FONT}display:flex;justify-content:space-between;align-items:center;'
                f'background:{COL["panel"]};border:1px solid {COL["line"]};border-left:4px solid {colr};'
                f'border-radius:10px;padding:8px 12px;margin-bottom:6px;max-width:520px;">'
                f'<span style="color:{COL["hi"]};font-weight:600;">{_esc(al["mercato"])}</span>'
                f'<span style="color:{COL["lo"]};font-size:12px;">stat {al["prob"]*100:.0f}% · '
                f'quota {al["market_prob"]*100:.0f}% · '
                f'<b style="color:{colr};text-transform:uppercase;">{al["livello"]}</b> ({verso})</span></div>')
        st.html(rows)

    st.subheader("🎯 Tutti i mercati")
    acol = {"basso": COL["draw"], "medio": "#e08a2b", "alto": COL["loss"]}
    cards = ""
    for m in sorted(a["mercati"], key=lambda x: -x["confidence"]):
        mk = ""
        if m.get("market_prob") is not None:
            q = f' @ {m["quota"]:.2f}' if m.get("quota") else ""
            mk = f' · quota {m["market_prob"]*100:.0f}%{q}'
            if m.get("alert"):
                colm = acol.get(m["alert"], COL["lo"])
                mk += f' · <b style="color:{colm};">⚠{m["alert"]}</b>'
            if m.get("var_quota") and abs(m["var_quota"]) >= 0.05:
                mk += f' · quota {"↓" if m["var_quota"]>0 else "↑"}'
        elif m.get("quota"):
            mk = f' · quota {m["quota"]:.2f}'
        cards += (
            f'<div style="{FONT}background:{COL["panel"]};border:1px solid {COL["line"]};'
            f'border-radius:12px;padding:10px 14px;margin-bottom:7px;max-width:520px;">'
            f'<div style="display:flex;justify-content:space-between;align-items:center;">'
            f'<span style="color:{COL["hi"]};font-weight:600;">{_esc(m["mercato"])}</span>'
            f'<span style="color:{COL["lo"]};font-size:12px;">conf '
            f'<b style="color:{_col_conf(m["confidence"])};">{m["confidence"]:.0f}</b>/100 · '
            f'stat {m["prob"]*100:.0f}%{mk}</span></div>'
            f'<div style="margin-top:7px;">{_barra(m["confidence"], _col_conf(m["confidence"]))}</div></div>')
    st.html(cards)


def _dettaglio_partita(df, row, comp_df):
    """Scheda COMPLETA di una partita: analisi dal vivo se non ha risultato, scheda
    pre-partita salvata (integrale) se è già finita (niente ricalcolo col senno di poi)."""
    home, away = row["squadra_casa"], row["squadra_trasferta"]
    finita = _num_ok(row.get("gol_casa")) and _num_ok(row.get("gol_trasferta"))
    st.divider()
    st.subheader(f"🔎 {home} - {away}")
    meta = []
    if _txt(row.get("competizione")):
        meta.append(_label_da_comp(row.get("competizione"), comp_df))
    if _txt(row.get("ora")):
        meta.append(f"ore {_txt(row.get('ora'))}")
    if meta:
        st.caption(" · ".join(meta))

    if finita:
        gc, gt = int(row["gol_casa"]), int(row["gol_trasferta"])
        st.markdown(f"**Risultato finale: {gc} - {gt}**")
        pron = carica_pronostici()
        rec = None
        if not pron.empty and "partita_id" in pron.columns:
            m = pron[pron["partita_id"] == str(row["id"])]
            if not m.empty:
                rec = m.iloc[0]
        snap = None
        if rec is not None and "scheda_json" in (rec.index if hasattr(rec, "index") else []):
            try:
                snap = json.loads(rec["scheda_json"]) if _txt(rec.get("scheda_json")) else None
            except Exception:
                snap = None
        if snap:
            won = _pronostico_vinto(snap["best"]["mercato"], gc, gt)
            esito = "✅ vinto" if won is True else ("❌ perso" if won is False else "—")
            st.markdown(f"Esito pronostico: **{esito}**")
            render_scheda_st(snap, home, away)
        else:
            st.info("Per questa partita non è stata salvata la scheda pre-partita completa "
                    "(verrà salvata per i pronostici generati d'ora in avanti). Non la ricalcolo "
                    "a posteriori perché userebbe dati successivi alla partita.")
        return

    # partita non ancora giocata: analisi dal vivo, scheda completa
    a, home, away, odds = _analizza_row(df, row, comp_df, calibratori=carica_calibrazione())
    if a.get("errore"):
        st.warning("Dati storici insufficienti per l'analisi.")
        return
    render_scheda_st(a, home, away)
    st.caption("Per regolare i pesi del modello usa 🔮 Analisi & Pronostico.")


def pagina_database(user):
    st.header("🗄️ Database partite")

    if not supabase_pronto():
        st.warning("Supabase non configurato.")
        return

    if st.button("🔄 Ricarica"):
        st.cache_data.clear()

    df = carica_partite()
    if df.empty:
        st.info("Nessuna partita salvata.")
        return

    # --- Partite da compilare ---
    if "da_compilare" in df.columns:
        dac = df[df["da_compilare"] == True]
        if not dac.empty:
            with st.expander(f"🧩 Partite da compilare ({len(dac)})", expanded=True):
                st.caption("Partite pianificate in attesa di ultimi risultati / quote. "
                           "Escono da qui quando aggiungi risultato o quote. "
                           "Spunta 🗑️ per eliminarne una.")
                vdac = pd.DataFrame({
                    "id": dac["id"],
                    "Data": dac["data"],
                    "Ora": dac.get("ora"),
                    "Competizione": dac.get("competizione"),
                    "Casa": dac["squadra_casa"],
                    "Trasferta": dac["squadra_trasferta"],
                    "🗑️": [False] * len(dac),
                }).reset_index(drop=True)
                edac = st.data_editor(
                    vdac, use_container_width=True, hide_index=True, key="editor_dac",
                    disabled=["id", "Data", "Ora", "Competizione", "Casa", "Trasferta"],
                    column_config={"id": None,
                                   "Data": st.column_config.DateColumn(format="DD.MM.YY"),
                                   "🗑️": st.column_config.CheckboxColumn("🗑️")})
                el_dac = edac[edac["🗑️"] == True]
                if len(el_dac):
                    st.warning(f"{len(el_dac)} partite selezionate per l'eliminazione.")
                    conf_dac = st.checkbox("Confermo l'eliminazione", key="conf_dac")
                    if st.button("🗑️ Elimina selezionate", key="del_dac"):
                        if not conf_dac:
                            st.warning("Spunta la conferma prima di eliminare.")
                        else:
                            try:
                                elimina_partite([r["id"] for _, r in el_dac.iterrows()])
                                st.success(f"Eliminate {len(el_dac)} partite.")
                                st.rerun()
                            except Exception as e:
                                st.error(f"Errore: {e}")

    # filtro rapido
    q = st.text_input("Cerca squadra")
    cc = st.columns([1, 1])
    solo_seguite = cc[0].checkbox(
        "Mostra solo le partite che seguo", value=True,
        help="Le fixture per cui hai incollato ultimi risultati + quote. "
             "Disattiva per vedere anche lo storico usato dal motore.")
    base = df
    if solo_seguite and "is_target" in df.columns:
        base = df[df["is_target"] == True]
    vis = base
    if q:
        m = base["squadra_casa"].str.contains(q, case=False, na=False) | \
            base["squadra_trasferta"].str.contains(q, case=False, na=False)
        vis = base[m]

    st.caption(f"{len(vis)} partite (senza risultato in cima, poi dalla più recente).")

    # ordina: prima le partite SENZA risultato (da aggiornare), poi le altre per data desc.
    # Aggiungo l'id come ultimo criterio: ordine DETERMINISTICO (niente riordini casuali
    # su parità di data, che confondevano l'editor).
    vis = vis.copy()
    vis["_senza_ris"] = vis["gol_casa"].isna() | vis["gol_trasferta"].isna()
    ordina = ["_senza_ris"] + (["data"] if "data" in vis.columns else [])
    ordina += (["id"] if "id" in vis.columns else [])
    vis = vis.sort_values(ordina, ascending=[False] + [False] * (len(ordina) - 1))

    # opzioni competizione dal menu (competizioni a sistema) + valori già presenti
    comp_df = carica_competizioni()
    opzioni_comp = []
    for _, cr in comp_df.sort_values("nome_lungo", na_position="last").iterrows() \
            if not comp_df.empty else []:
        lab = label_competizione(cr.get("nome_lungo"), cr.get("nazione")) or _txt(cr.get("nome_corto"))
        if lab:
            opzioni_comp.append(lab)
    esistenti = [_label_da_comp(x, comp_df)
                 for x in vis.get("competizione", pd.Series(dtype=object)) if _txt(x)]
    opzioni_comp = sorted(set(opzioni_comp) | set(esistenti))

    vista = pd.DataFrame({
        "id": vis["id"],
        "Data": vis["data"],
        "Competizione": [_label_da_comp(x, comp_df) for x in vis.get("competizione",
                         pd.Series([None] * len(vis)))],
        "Casa": vis["squadra_casa"],
        "Trasferta": vis["squadra_trasferta"],
        "Gol Casa": vis["gol_casa"].astype("Int64"),
        "Gol Trasferta": vis["gol_trasferta"].astype("Int64"),
        "🔍": [False] * len(vis),
        "🗑️": [False] * len(vis),
    }).reset_index(drop=True)
    orig = vista.copy()   # per confrontare cosa è cambiato

    # Chiave dell'editor legata all'ORDINE degli id: se la tabella si riordina (es. dopo
    # un salvataggio, quando le partite con risultato scendono), la chiave cambia e
    # l'editor riparte pulito, senza riapplicare modifiche vecchie a righe diverse.
    _firma_righe = ",".join(str(x) for x in vista["id"].tolist())
    _key_editor = "editor_db_" + hashlib.md5(_firma_righe.encode()).hexdigest()[:10]

    edit = st.data_editor(
        vista, use_container_width=True, hide_index=True, key=_key_editor,
        disabled=["id", "Casa", "Trasferta"],
        column_config={
            "id": None,  # nascosta
            "Data": st.column_config.DateColumn(format="DD.MM.YY"),
            "Competizione": st.column_config.SelectboxColumn(
                options=opzioni_comp,
                help="Scegli tra le competizioni a sistema (Configurazione). "
                     "Cambiarla aggiorna anche la categoria della partita."),
            "Gol Casa": st.column_config.NumberColumn(min_value=0, step=1),
            "Gol Trasferta": st.column_config.NumberColumn(min_value=0, step=1),
            "🔍": st.column_config.CheckboxColumn(
                "🔍", help="Spunta una partita per aprirne sotto il dettaglio con l'analisi."),
            "🗑️": st.column_config.CheckboxColumn(
                "🗑️", help="Spunta le partite da eliminare, poi usa il pulsante 'Elimina'."),
        },
    )

    # --- dettaglio/analisi della partita spuntata con 🔍 ---
    da_vedere = edit[edit["🔍"] == True] if "🔍" in edit else edit.iloc[0:0]
    if len(da_vedere) >= 1:
        mid = da_vedere.iloc[0]["id"]
        rr = df[df["id"] == mid].iloc[0]
        _dettaglio_partita(df, rr, carica_competizioni())

    col1, col2 = st.columns(2)
    if col1.button("💾 Salva modifiche risultati", type="primary"):
        def _val(x):
            return None if pd.isna(x) else int(x)
        records = []
        for i in range(len(edit)):
            gc, gt = _val(edit.iloc[i]["Gol Casa"]), _val(edit.iloc[i]["Gol Trasferta"])
            gc0, gt0 = _val(orig.iloc[i]["Gol Casa"]), _val(orig.iloc[i]["Gol Trasferta"])
            comp_new = _txt(edit.iloc[i]["Competizione"]) or None
            comp_old = _txt(orig.iloc[i]["Competizione"]) or None
            data_new = edit.iloc[i]["Data"]
            data_old = orig.iloc[i]["Data"]
            data_cambiata = str(data_new) != str(data_old)
            if gc == gc0 and gt == gt0 and comp_new == comp_old and not data_cambiata:
                continue   # riga non modificata: non la tocco (salvataggio istantaneo)
            rec = {"id": edit.iloc[i]["id"], "gol_casa": gc, "gol_trasferta": gt,
                   "aggiornato_il": datetime.utcnow().isoformat()}
            if gc is not None and gt is not None:
                rec["da_compilare"] = False
            if comp_new != comp_old:
                rec["competizione"] = comp_new
                # riallinea la categoria alla nuova competizione
                rec["tipo_partita"] = categoria_o_nd(comp_new, comp_df)
            if data_cambiata and data_new is not None and not pd.isna(data_new):
                rec["data"] = str(data_new)
            records.append(rec)
        try:
            if records:
                aggiorna_partite(records)
                st.success(f"Aggiornate {len(records)} partite.")
                st.rerun()   # ricarica pulito: la tabella si riordina e l'editor riparte da zero
            else:
                st.info("Nessuna modifica da salvare.")
        except Exception as e:
            st.error(f"Errore: {e}")

    # --- eliminazione partite selezionate (col 🗑️) ---
    da_elim = edit[edit["🗑️"] == True] if "🗑️" in edit else edit.iloc[0:0]
    if len(da_elim):
        st.divider()
        st.warning(f"{len(da_elim)} partite selezionate per l'eliminazione:")
        for _, rr in da_elim.iterrows():
            st.markdown(f"- {rr['Casa']} - {rr['Trasferta']} ({rr['Data']})")
        conf = st.checkbox("Confermo: elimina definitivamente queste partite e i dati agganciati "
                           "(pronostici collegati)")
        if st.button("🗑️ Elimina selezionate", type="secondary"):
            if not conf:
                st.warning("Spunta la conferma prima di eliminare.")
            else:
                try:
                    elimina_partite([rr["id"] for _, rr in da_elim.iterrows()])
                    st.success(f"Eliminate {len(da_elim)} partite.")
                    st.rerun()
                except Exception as e:
                    st.error(f"Errore: {e}")

    col2.download_button(
        "⬇️ Esporta partite (Excel)",
        data=to_excel({"partite": partite_per_export(df)}),
        file_name="partite.xlsx",
        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    )

    # --- editor quote / valori rose di una partita ---
    st.divider()
    st.subheader("✏️ Quote e valori di una partita")
    st.caption("Inserisci o correggi le quote (anche variazioni) e i valori rosa. "
               "Aggiungere le quote rende la partita pronosticabile in 🔮 Analisi.")
    fx = df.copy()
    cond = pd.Series(False, index=fx.index)
    if "is_target" in fx:
        cond = cond | (fx["is_target"] == True)
    if "da_compilare" in fx:
        cond = cond | (fx["da_compilare"] == True)
    fx = fx[cond] if cond.any() else fx
    if "data" in fx:
        fx = fx.sort_values("data", ascending=False)
    opz_q = {}
    for _, r in fx.iterrows():
        d = r["data"].strftime("%d.%m.%y") if hasattr(r["data"], "strftime") else r["data"]
        opz_q[f'{d} · {r["squadra_casa"]} - {r["squadra_trasferta"]}'] = r["id"]
    if not opz_q:
        st.caption("Nessuna partita da pronosticare o pianificata. Aggiungine una dall'estrattore "
                   "o dalla pianificazione.")
    else:
        sel_q = st.selectbox("Partita", ["—"] + list(opz_q.keys()), key="quote_sel")
        if sel_q != "—":
            mid = opz_q[sel_q]
            rr = df[df["id"] == mid].iloc[0]

            def _n(col):
                v = rr.get(col)
                return None if (v is None or pd.isna(v)) else float(v)

            def _in(label, col, key):
                return st.number_input(label, value=_n(col), step=0.01, format="%.2f",
                                       key=f"qe_{key}")

            st.markdown("**1X2** — quota iniziale e variazione (la variazione può essere negativa)")
            c = st.columns(3)
            with c[0]:
                q1 = _in("1", "quota_iniziale_1", "1")
                v1 = _in("Var. 1", "variazione_quota_1", "v1")
            with c[1]:
                qx = _in("X", "quota_iniziale_x", "x")
                vx = _in("Var. X", "variazione_quota_x", "vx")
            with c[2]:
                q2 = _in("2", "quota_iniziale_2", "2")
                v2 = _in("Var. 2", "variazione_quota_2", "v2")

            st.markdown("**Over/Under 2.5** e **Goal/No Goal**")
            c = st.columns(4)
            with c[0]:
                qov = _in("Over 2.5", "quota_iniziale_over", "ov")
                vov = _in("Var. Over", "variazione_quota_over", "vov")
            with c[1]:
                qun = _in("Under 2.5", "quota_iniziale_under", "un")
                vun = _in("Var. Under", "variazione_quota_under", "vun")
            with c[2]:
                qgo = _in("Goal", "quota_iniziale_goal", "go")
                vgo = _in("Var. Goal", "variazione_quota_goal", "vgo")
            with c[3]:
                qng = _in("No Goal", "quota_iniziale_nogoal", "ng")
                vng = _in("Var. NoGoal", "variazione_quota_nogoal", "vng")

            st.markdown("**Forma** (indice del sito, es. 7.21) e **valori rosa** "
                        "(opzionali: i valori rosa mitigano il livello di lega)")
            c = st.columns(2)
            with c[0]:
                fcasa = _in("Forma casa", "forma_casa", "fc")
                vcasa = _in("Valore rosa casa", "val_casa", "vc")
            with c[1]:
                ftras = _in("Forma trasferta", "forma_trasferta", "ft")
                vtras = _in("Valore rosa trasferta", "val_trasferta", "vt")

            if st.button("💾 Salva quote/valori", type="primary", key="save_quote"):
                payload = {
                    "quota_iniziale_1": q1, "quota_iniziale_x": qx, "quota_iniziale_2": q2,
                    "variazione_quota_1": v1, "variazione_quota_x": vx, "variazione_quota_2": v2,
                    "quota_iniziale_over": qov, "quota_iniziale_under": qun,
                    "variazione_quota_over": vov, "variazione_quota_under": vun,
                    "quota_iniziale_goal": qgo, "quota_iniziale_nogoal": qng,
                    "variazione_quota_goal": vgo, "variazione_quota_nogoal": vng,
                    "forma_casa": fcasa, "forma_trasferta": ftras,
                    "val_casa": vcasa, "val_trasferta": vtras,
                    "aggiornato_il": datetime.utcnow().isoformat(),
                }
                # se ci sono quote, la partita diventa pronosticabile ed esce da "da compilare"
                if any(x is not None for x in (q1, qx, q2, qov, qun, qgo, qng)):
                    payload["is_target"] = True
                    payload["da_compilare"] = False
                try:
                    aggiorna_partite([{"id": mid, **payload}])
                    st.success("Quote/valori salvati. La partita è ora pronosticabile in 🔮 Analisi.")
                    st.rerun()
                except Exception as e:
                    st.error(f"Errore: {e}")

    # --- dettaglio grafico ---
    st.subheader("🔎 Dettaglio partita")
    opzioni = {
        f'{r["data"].strftime("%d.%m.%y") if hasattr(r["data"], "strftime") else r["data"]}'
        f' · {r["squadra_casa"]} - {r["squadra_trasferta"]}': r
        for _, r in vis.iterrows()
    }
    scelta = st.selectbox("Seleziona una partita", list(opzioni.keys()))
    if not scelta:
        return

    row = opzioni[scelta]
    home, away = row["squadra_casa"], row["squadra_trasferta"]

    # tutte le partite di ciascuna squadra (dal DB completo), dalla più recente
    home_matches = df[(df["squadra_casa"] == home) | (df["squadra_trasferta"] == home)] \
        .sort_values("data", ascending=False)
    away_matches = df[(df["squadra_casa"] == away) | (df["squadra_trasferta"] == away)] \
        .sort_values("data", ascending=False)

    st.html(render_hero(row, home_matches, away_matches))

    st.markdown("###### Tutti i risultati")
    t_home, t_away = st.tabs([f"🟠 {home} ({len(home_matches)})",
                              f"🔵 {away} ({len(away_matches)})"])
    with t_home:
        st.caption("Evidenziate le partite giocate in **casa**.")
        st.html(render_lista(home, home_matches, "casa"))
    with t_away:
        st.caption("Evidenziate le partite giocate in **trasferta**.")
        st.html(render_lista(away, away_matches, "trasf"))

    with st.expander("📋 Copia testo (per export su altra AI)"):
        gc = "" if pd.isna(row["gol_casa"]) else int(row["gol_casa"])
        gt = "" if pd.isna(row["gol_trasferta"]) else int(row["gol_trasferta"])
        ris = f"{gc}-{gt}" if gc != "" and gt != "" else "?-?"
        blocco = [f"{home} - {away} | {ris}",
                  f"Data: {row['data']}  Competizione: {row.get('competizione', '')}"]
        if _num_ok(row.get("quota_iniziale_1")):
            blocco.append(f"1X2: {row['quota_iniziale_1']} / {row['quota_iniziale_x']} / {row['quota_iniziale_2']}")
        if _num_ok(row.get("variazione_quota_1")):
            blocco.append(f"1X2 (variazione): {row['variazione_quota_1']} / {row['variazione_quota_x']} / {row['variazione_quota_2']}")
        if _num_ok(row.get("quota_iniziale_over")):
            blocco.append(f"O/U 2.5: {row['quota_iniziale_over']} / {row['quota_iniziale_under']}")
        if _num_ok(row.get("variazione_quota_over")):
            blocco.append(f"O/U 2.5 (variazione): {row['variazione_quota_over']} / {row['variazione_quota_under']}")
        if _num_ok(row.get("quota_iniziale_goal")):
            blocco.append(f"Goal/NoGoal: {row['quota_iniziale_goal']} / {row['quota_iniziale_nogoal']}")
        if _num_ok(row.get("variazione_quota_goal")):
            blocco.append(f"Goal/NoGoal (variazione): {row['variazione_quota_goal']} / {row['variazione_quota_nogoal']}")
        if _num_ok(row.get("forma_casa")) or _num_ok(row.get("forma_trasferta")):
            blocco.append(f"Forma: {row.get('forma_casa')} / {row.get('forma_trasferta')}")
        if _num_ok(row.get("val_casa")):
            blocco.append(f"Valori rose: {row['val_casa']} / {row['val_trasferta']}")
        st.code("\n".join(blocco), language="text")


# =============================================================================
#  PAGINA: ESTRATTORE RISULTATI
# =============================================================================
def pagina_estrattore_risultati(user):
    st.header("📊 Estrattore risultati")
    st.caption("Incolla i risultati per competizione: il punteggio viene agganciato "
               "alle partite già presenti nel database (match per nome squadra).")

    if not supabase_pronto():
        st.warning("Supabase non configurato.")
        return

    testo = st.text_area("Incolla qui i risultati", height=260, key="testo_risultati")
    if not testo.strip():
        st.info("In attesa dei risultati…")
        return

    risultati = parse_risultati(testo)
    if not risultati:
        st.warning("Nessun risultato riconosciuto.")
        return

    part = carica_partite()
    comp_df = carica_competizioni()

    # competizioni presenti nel testo
    comp_viste = []
    for r in risultati:
        if r["competizione"] and r["competizione"] not in [c[0] for c in comp_viste]:
            comp_viste.append((r["competizione"], r["nome_lungo"], r["nazione"]))

    # normalizzati DB per il match
    if not part.empty:
        part = part.copy()
        part["_c"] = part["squadra_casa"].map(lambda x: _key(_norm_squadra(x)))
        part["_t"] = part["squadra_trasferta"].map(lambda x: _key(_norm_squadra(x)))

    righe, meta = [], []
    for r in risultati:
        nc, nt = _key(_norm_squadra(r["casa"])), _key(_norm_squadra(r["trasferta"]))
        cand = part[(part["_c"] == nc) & (part["_t"] == nt)] if not part.empty else part
        pending = cand[cand["gol_casa"].isna()] if not cand.empty else cand

        mid, applica = None, False
        if not cand.empty and not pending.empty:
            mid = pending.iloc[0]["id"]
            stato = "✅ da compilare"
            applica = True
        elif not cand.empty:
            mid = cand.iloc[0]["id"]
            stato = "⚠️ già con risultato"
        else:
            stato = "❌ non trovata"

        cat = categoria_di(r["competizione"], comp_df)
        righe.append({
            "Applica": applica,
            "Competizione": r["competizione"] or "—",
            "Categoria": cat or "—",
            "Casa": _norm_squadra(r["casa"]),
            "Trasferta": _norm_squadra(r["trasferta"]),
            "Risultato": f'{r["gol_casa"]}-{r["gol_trasferta"]}',
            "Stato": stato,
        })
        meta.append({"id": mid, "gc": r["gol_casa"], "gt": r["gol_trasferta"],
                     "qualif": r["qualificatore"], "competizione": r["competizione"],
                     "tipo": cat})

    st.markdown(f"**{len(righe)}** risultati letti · "
                f"{sum(1 for x in righe if x['Stato'].startswith('✅'))} agganciabili")

    edit = st.data_editor(
        pd.DataFrame(righe), use_container_width=True, hide_index=True, key="editor_risultati",
        disabled=["Competizione", "Categoria", "Casa", "Trasferta", "Risultato", "Stato"],
        column_config={"Applica": st.column_config.CheckboxColumn(width="small")},
    )

    # --- aggancio manuale per somiglianza (righe non trovate) ---
    pending_all = part[part["gol_casa"].isna()] if not part.empty else part
    scelte_manuali = {}  # idx risultato -> id partita scelta
    non_trovate = [i for i, x in enumerate(righe) if x["Stato"].startswith("❌")]
    if non_trovate and not pending_all.empty:
        with st.expander(f"🔗 Aggancio manuale per somiglianza ({len(non_trovate)} non trovate)"):
            opzioni_lbl = {
                f'{p["squadra_casa"]} - {p["squadra_trasferta"]}'
                f'  ({p["data"]})': p["id"] for _, p in pending_all.iterrows()
            }
            for i in non_trovate:
                r = risultati[i]
                # ordina i candidati per somiglianza media casa+trasferta
                punteggi = []
                for _, p in pending_all.iterrows():
                    s = (_similarita(_norm_squadra(r["casa"]), p["squadra_casa"])
                         + _similarita(_norm_squadra(r["trasferta"]), p["squadra_trasferta"])) / 2
                    punteggi.append((s, f'{p["squadra_casa"]} - {p["squadra_trasferta"]}  ({p["data"]})'))
                punteggi.sort(reverse=True)
                suggeriti = [lbl for s, lbl in punteggi[:6] if s > 0.4]
                scelta = st.selectbox(
                    f'{_norm_squadra(r["casa"])} - {_norm_squadra(r["trasferta"])} '
                    f'({r["gol_casa"]}-{r["gol_trasferta"]})',
                    ["(nessuno)"] + suggeriti, key=f"manu_{i}")
                if scelta != "(nessuno)":
                    scelte_manuali[i] = opzioni_lbl[scelta]

    # competizioni nuove (non ancora in anagrafica)
    esistenti = set()
    if not comp_df.empty:
        esistenti = {_key(label_competizione(c.get("nome_lungo"), c.get("nazione")))
                     for _, c in comp_df.iterrows()}
    nuove = [(lbl, nl, na) for (lbl, nl, na) in comp_viste if _key(lbl) not in esistenti]
    if nuove:
        st.caption("Nuove competizioni che verranno aggiunte in Configurazione: "
                   + ", ".join(lbl for lbl, _, _ in nuove))

    if st.button("💾 Aggancia risultati", type="primary"):
        # 1) registra le competizioni nuove (categoria da assegnare in Config)
        if nuove:
            upsert_competizioni([
                {"nome_lungo": nl, "nazione": na, "categoria": "Non assegnata"}
                for _, nl, na in nuove
            ])

        # 2) costruisce gli aggiornamenti (match automatici + manuali)
        def _prepara_update(m, mid):
            u = {"id": mid, "gol_casa": m["gc"], "gol_trasferta": m["gt"],
                 "da_compilare": False, "aggiornato_il": datetime.utcnow().isoformat()}
            if m["qualif"]:
                u["qualificatore"] = m["qualif"]
            match_row = part[part["id"] == mid]
            comp_attuale = match_row.iloc[0].get("competizione") if not match_row.empty else None
            tipo_attuale = match_row.iloc[0].get("tipo_partita") if not match_row.empty else None
            if m["tipo"]:
                u["tipo_partita"] = m["tipo"]
            elif not (tipo_attuale and str(tipo_attuale).strip()):
                u["tipo_partita"] = ND  # categoria non determinabile
            if m["competizione"] and (comp_attuale is None or str(comp_attuale).strip() == ""):
                u["competizione"] = m["competizione"]
            return u

        updates, visti = [], set()
        for idx, rrow in edit.iterrows():
            m = meta[idx]
            if rrow["Applica"] and m["id"]:
                updates.append(_prepara_update(m, m["id"]))
                visti.add(m["id"])
        for idx, mid in scelte_manuali.items():
            if mid not in visti:
                updates.append(_prepara_update(meta[idx], mid))
                visti.add(mid)

        try:
            if updates:
                aggiorna_partite(updates)
            st.success(f"Agganciati {len(updates)} risultati."
                       + (f" Aggiunte {len(nuove)} competizioni in Configurazione." if nuove else ""))
            st.cache_data.clear()
        except Exception as e:
            st.error(f"Errore: {e}")


# =============================================================================
#  PAGINA: ESTRATTORE PIANIFICAZIONE
# =============================================================================
def pagina_estrattore_pianificazione(user):
    st.header("🗓️ Estrattore pianificazione")
    st.caption("Incolla le partite in programma: vengono create nel database come "
               "**partite da compilare** (poi aggiungi ultimi risultati, quote, ecc.).")

    if not supabase_pronto():
        st.warning("Supabase non configurato.")
        return

    # --- inserimento manuale rapido (una partita alla volta) ---
    with st.expander("➕ Aggiungi una partita manualmente"):
        cc = st.columns(2)
        m_data = cc[0].date_input("Data", value=None, format="DD/MM/YYYY", key="man_data")
        m_ora = cc[1].text_input("Ora (opzionale)", key="man_ora", placeholder="20:45")
        cc = st.columns(2)
        m_casa = cc[0].text_input("Squadra casa", key="man_casa")
        m_trasf = cc[1].text_input("Squadra trasferta", key="man_trasf")
        # menu a tendina delle competizioni a sistema
        comp_df = carica_competizioni()
        opzioni = {"— nessuna —": None}
        if not comp_df.empty:
            for _, cr in comp_df.sort_values("nome_lungo", na_position="last").iterrows():
                lab = label_competizione(cr.get("nome_lungo"), cr.get("nazione")) \
                    or _txt(cr.get("nome_corto"))
                codice = _txt(cr.get("nome_corto")) or lab
                if lab:
                    opzioni[lab] = codice
        m_comp_lab = st.selectbox(
            "Competizione", list(opzioni.keys()), key="man_comp",
            help="Scegli tra le competizioni già a sistema (Configurazione). "
                 "Se manca, aggiungila prima in Configurazione.")
        m_comp = opzioni.get(m_comp_lab)
        if st.button("Aggiungi partita", key="man_btn"):
            if not m_data or not m_casa.strip() or not m_trasf.strip():
                st.warning("Servono almeno data, squadra casa e squadra trasferta.")
            else:
                try:
                    salva_partite([{
                        "data": str(m_data), "ora": m_ora.strip() or None,
                        "squadra_casa": m_casa.strip(), "squadra_trasferta": m_trasf.strip(),
                        "competizione": m_comp,
                        "tipo_partita": "ND", "da_compilare": True,
                        "inserito_da": user["username"],
                        "aggiornato_il": datetime.utcnow().isoformat(),
                    }])
                    st.success(f"Partita aggiunta: {m_casa.strip()} - {m_trasf.strip()}")
                except Exception as e:
                    st.error(f"Errore: {e}")

    data_batch = st.date_input(
        "Data delle partite (obbligatoria, vale per tutta la tranche)",
        value=None, format="DD/MM/YYYY")
    if not data_batch:
        st.info("Seleziona la data della tranche per continuare. "
                "Tutte le partite caricate insieme avranno questa data.")
        return

    testo = st.text_area("Incolla qui la pianificazione", height=260, key="testo_pian")
    if not testo.strip():
        st.info("In attesa della pianificazione…")
        return

    fixtures = parse_pianificazione(testo)
    if not fixtures:
        st.warning("Nessuna partita riconosciuta.")
        return

    part = carica_partite()
    comp_df = carica_competizioni()

    # per rilevare i duplicati già presenti (stessa data + squadre)
    esistenti_key = set()
    if not part.empty:
        for _, p in part.iterrows():
            esistenti_key.add((str(p["data"]), _key(_norm_squadra(p["squadra_casa"])),
                               _key(_norm_squadra(p["squadra_trasferta"]))))

    comp_viste, righe, meta = [], [], []
    for f in fixtures:
        if f["competizione"] and f["competizione"] not in [c[0] for c in comp_viste]:
            comp_viste.append((f["competizione"], f["nome_lungo"], f["nazione"]))
        cat = categoria_o_nd(f["competizione"], comp_df)  # 'ND' se non determinabile
        casa, trasf = _norm_squadra(f["casa"]), _norm_squadra(f["trasferta"])
        gia = (str(data_batch), _key(casa), _key(trasf)) in esistenti_key
        righe.append({
            "Crea": not gia,
            "Ora": f["ora"] or "",
            "Competizione": f["competizione"] or "—",
            "Categoria": cat,
            "Casa": casa,
            "Trasferta": trasf,
            "Stato": "già presente" if gia else "nuova",
        })
        meta.append({"competizione": f["competizione"], "cat": cat})

    st.markdown(f"**{len(righe)}** partite lette · data **{data_batch.strftime('%d.%m.%Y')}** · "
                f"{sum(1 for x in righe if x['Stato'] == 'nuova')} nuove")

    edit = st.data_editor(
        pd.DataFrame(righe), use_container_width=True, hide_index=True, num_rows="fixed",
        key="editor_pian",
        disabled=["Ora", "Competizione", "Categoria", "Casa", "Trasferta", "Stato"],
        column_config={"Crea": st.column_config.CheckboxColumn(width="small")},
    )

    esistenti = set()
    if not comp_df.empty:
        esistenti = {_key(label_competizione(c.get("nome_lungo"), c.get("nazione")))
                     for _, c in comp_df.iterrows()}
    nuove = [(lbl, nl, na) for (lbl, nl, na) in comp_viste if _key(lbl) not in esistenti]
    if nuove:
        st.caption("Nuove competizioni che verranno aggiunte in Configurazione: "
                   + ", ".join(lbl for lbl, _, _ in nuove))

    if st.button("💾 Crea partite da compilare", type="primary"):
        if nuove:
            upsert_competizioni([
                {"nome_lungo": nl, "nazione": na, "categoria": "Non assegnata"}
                for _, nl, na in nuove
            ])
        records = []
        for idx, rrow in edit.iterrows():
            if not rrow["Crea"] or not rrow["Casa"] or not rrow["Trasferta"]:
                continue
            records.append({
                "data": str(data_batch),
                "ora": rrow["Ora"] or None,
                "squadra_casa": rrow["Casa"],
                "squadra_trasferta": rrow["Trasferta"],
                "competizione": meta[idx]["competizione"],
                "tipo_partita": meta[idx]["cat"],   # categoria o 'ND'
                "da_compilare": True,
                "inserito_da": user["username"],
                "aggiornato_il": datetime.utcnow().isoformat(),
            })
        try:
            if records:
                salva_partite(records)  # upsert su (data, casa, trasferta)
            st.success(f"Create {len(records)} partite da compilare per il "
                       f"{data_batch.strftime('%d.%m.%Y')}."
                       + (f" Aggiunte {len(nuove)} competizioni." if nuove else ""))
            st.cache_data.clear()
        except Exception as e:
            st.error(f"Errore: {e}")


# =============================================================================
#  DETTAGLIO PARTITA (vista grafica in-app)
# =============================================================================
COL = {
    "ink": "#0E1622", "panel": "#17223A", "panel2": "#1E2B47",
    "line": "rgba(255,255,255,0.08)",
    "home": "#FF8A5B", "away": "#38BDF8",
    "hi": "#EEF2F8", "lo": "#93A1B5",
    "win": "#34D399", "draw": "#94A3B8", "loss": "#FB7185",
}
FONT = "font-family:Inter,-apple-system,system-ui,sans-serif;"


def _esc(x):
    return _html.escape("" if x is None else str(x))


def _num_ok(v):
    return v is not None and not pd.isna(v)


def _esito(row, team):
    """V/N/P dal punto di vista della squadra 'team'. None se non giocata."""
    gc, gt = row.get("gol_casa"), row.get("gol_trasferta")
    if not _num_ok(gc) or not _num_ok(gt):
        return None
    gc, gt = int(gc), int(gt)
    if gc == gt:
        return "N"
    casa_vince = gc > gt
    in_casa = row["squadra_casa"] == team
    if in_casa:
        return "V" if casa_vince else "P"
    return "V" if not casa_vince else "P"


def _col_esito(r):
    return {"V": COL["win"], "N": COL["draw"], "P": COL["loss"]}.get(r, COL["lo"])


def _dots(team, matches, n=5):
    out, cnt = "", 0
    for _, m in matches.iterrows():
        r = _esito(m, team)
        if r is None:
            continue
        c = _col_esito(r)
        out += (f'<span style="width:22px;height:22px;border-radius:7px;display:inline-flex;'
                f'align-items:center;justify-content:center;font-weight:700;font-size:12px;'
                f'background:{c}22;color:{c};border:1px solid {c}55;margin-left:6px;">{r}</span>')
        cnt += 1
        if cnt >= n:
            break
    return out or f'<span style="color:{COL["lo"]};font-size:12px;">—</span>'


def _odd_cell(label, prima, mod, fav, accent):
    """Mostra 'prima' quota e, se presente, la 'modificata' (in evidenza, con la
    prima barrata sopra)."""
    if not _num_ok(prima) and not _num_ok(mod):
        return ""
    ha_mod = _num_ok(mod)
    principale = float(mod) if ha_mod else float(prima)
    bg = f"{accent}1A" if fav else COL["panel2"]
    bd = f"{accent}88" if fav else COL["line"]
    lc = accent if fav else COL["lo"]
    barrata = ""
    if ha_mod and _num_ok(prima):
        barrata = (f'<div style="font-size:11px;color:{COL["lo"]};text-decoration:line-through;'
                   f'margin-top:2px;">{float(prima):.2f}</div>')
    tag_mod = (f'<div style="font-size:9px;font-weight:700;letter-spacing:.05em;color:{accent};'
               f'margin-top:2px;">VAR</div>') if ha_mod else ""
    return (f'<div style="flex:1;border-radius:12px;padding:12px 6px;text-align:center;'
            f'background:{bg};border:1px solid {bd};">'
            f'<div style="font-size:11px;font-weight:600;text-transform:uppercase;'
            f'letter-spacing:.04em;color:{lc};">{label}</div>'
            f'{barrata}'
            f'<div style="font-weight:700;font-size:19px;color:{COL["hi"]};margin-top:2px;">'
            f'{principale:.2f}</div>{tag_mod}</div>')


def _odds_row(cells):
    cells = [c for c in cells if c]
    if not cells:
        return ""
    return f'<div style="display:flex;gap:8px;margin-top:8px;">{"".join(cells)}</div>'


def render_hero(row, home_matches, away_matches):
    home, away = row["squadra_casa"], row["squadra_trasferta"]
    o = {
        "q1": row.get("quota_iniziale_1"), "qx": row.get("quota_iniziale_x"),
        "q2": row.get("quota_iniziale_2"),
        "q1_mod": row.get("variazione_quota_1"), "qx_mod": row.get("variazione_quota_x"),
        "q2_mod": row.get("variazione_quota_2"),
        "q_over25": row.get("quota_iniziale_over"), "q_under25": row.get("quota_iniziale_under"),
        "q_over25_mod": row.get("variazione_quota_over"),
        "q_under25_mod": row.get("variazione_quota_under"),
        "q_goal": row.get("quota_iniziale_goal"), "q_nogoal": row.get("quota_iniziale_nogoal"),
        "q_goal_mod": row.get("variazione_quota_goal"),
        "q_nogoal_mod": row.get("variazione_quota_nogoal"),
        "forma_casa": row.get("forma_casa"), "forma_trasferta": row.get("forma_trasferta"),
        "val_casa": row.get("val_casa"), "val_trasferta": row.get("val_trasferta"),
    }

    def _eff(prima, mod):
        """Valore effettivo per calcolare il favorito: modificata se presente."""
        if _num_ok(mod):
            return float(mod)
        return float(prima) if _num_ok(prima) else None

    # --- intestazione duello ---
    hero = (
        f'<div style="{FONT}max-width:460px;margin:0 auto;">'
        f'<div style="border-radius:18px;padding:18px;background:{COL["panel"]};'
        f'border:1px solid {COL["line"]};">'
        f'<div style="display:flex;align-items:center;justify-content:space-between;gap:10px;">'
        f'<div style="flex:1;min-width:0;">'
        f'<div style="font-weight:700;font-size:24px;color:{COL["hi"]};line-height:1.1;">{_esc(home)}</div>'
        f'<div style="font-size:12px;font-weight:600;color:{COL["home"]};margin-top:2px;">casa</div></div>'
        f'<div style="font-weight:700;font-size:13px;color:{COL["lo"]};background:{COL["ink"]};'
        f'padding:4px 12px;border-radius:9px;border:1px solid {COL["line"]};">VS</div>'
        f'<div style="flex:1;min-width:0;text-align:right;">'
        f'<div style="font-weight:700;font-size:24px;color:{COL["hi"]};line-height:1.1;">{_esc(away)}</div>'
        f'<div style="font-size:12px;font-weight:600;color:{COL["away"]};margin-top:2px;">trasferta</div></div>'
        f'</div>'
        f'<div style="display:flex;align-items:center;justify-content:space-between;'
        f'margin-top:16px;padding-top:14px;border-top:1px solid {COL["line"]};">'
        f'<div>{_dots(home, home_matches)}</div>'
        f'<span style="font-size:10px;text-transform:uppercase;letter-spacing:.15em;color:{COL["lo"]};">forma</span>'
        f'<div>{_dots(away, away_matches)}</div>'
        f'</div></div>'
    )

    # --- quote (prima + modificata) ---
    quote_html = ""
    e1, ex, e2 = _eff(o["q1"], o["q1_mod"]), _eff(o["qx"], o["qx_mod"]), _eff(o["q2"], o["q2_mod"])
    min_1x2 = min([v for v in (e1, ex, e2) if v is not None], default=None)
    r1 = _odds_row([
        _odd_cell("1", o["q1"], o["q1_mod"], e1 is not None and e1 == min_1x2, COL["home"]),
        _odd_cell("X", o["qx"], o["qx_mod"], ex is not None and ex == min_1x2, COL["draw"]),
        _odd_cell("2", o["q2"], o["q2_mod"], e2 is not None and e2 == min_1x2, COL["away"]),
    ])

    eo, eu = _eff(o["q_over25"], o["q_over25_mod"]), _eff(o["q_under25"], o["q_under25_mod"])
    r2 = _odds_row([
        _odd_cell("Over 2.5", o["q_over25"], o["q_over25_mod"],
                  eo is not None and eu is not None and eo < eu, "#A78BFA"),
        _odd_cell("Under 2.5", o["q_under25"], o["q_under25_mod"],
                  eo is not None and eu is not None and eu < eo, "#A78BFA"),
    ])

    eg, en = _eff(o["q_goal"], o["q_goal_mod"]), _eff(o["q_nogoal"], o["q_nogoal_mod"])
    r3 = _odds_row([
        _odd_cell("Goal", o["q_goal"], o["q_goal_mod"],
                  eg is not None and en is not None and eg < en, "#F472B6"),
        _odd_cell("No Goal", o["q_nogoal"], o["q_nogoal_mod"],
                  eg is not None and en is not None and en < eg, "#F472B6"),
    ])
    if r1 or r2 or r3:
        quote_html = (f'<div style="{FONT}max-width:460px;margin:18px auto 0;">'
                      f'<div style="font-size:12px;font-weight:600;text-transform:uppercase;'
                      f'letter-spacing:.15em;color:{COL["lo"]};margin-bottom:6px;">'
                      f'Quote <span style="font-weight:400;text-transform:none;letter-spacing:0;">'
                      f'· barrata = iniziale, grande = variazione</span></div>'
                      f'{r1}{r2}{r3}</div>')

    # --- forma squadre (voto ~0-10) ---
    forma_html = ""
    fc, ft = o["forma_casa"], o["forma_trasferta"]
    if _num_ok(fc) or _num_ok(ft):
        fc_v = float(fc) if _num_ok(fc) else None
        ft_v = float(ft) if _num_ok(ft) else None

        def _forma_chip(val, accent, alto):
            if val is None:
                return (f'<div style="flex:1;text-align:center;padding:12px 8px;border-radius:12px;'
                        f'background:{COL["panel2"]};border:1px solid {COL["line"]};color:{COL["lo"]};">—</div>')
            hot = alto
            bg = f"{accent}1A" if hot else COL["panel2"]
            bd = f"{accent}88" if hot else COL["line"]
            return (f'<div style="flex:1;text-align:center;padding:12px 8px;border-radius:12px;'
                    f'background:{bg};border:1px solid {bd};">'
                    f'<div style="font-weight:700;font-size:20px;color:{COL["hi"]};">{val:g}</div>'
                    f'<div style="height:5px;border-radius:99px;margin-top:8px;background:{COL["ink"]};overflow:hidden;">'
                    f'<div style="height:100%;width:{min(max(val,0),10)*10:.0f}%;background:{accent};"></div></div>'
                    f'</div>')

        casa_alto = _num_ok(fc) and (not _num_ok(ft) or fc_v >= ft_v)
        trasf_alto = _num_ok(ft) and (not _num_ok(fc) or ft_v >= fc_v)
        forma_html = (
            f'<div style="{FONT}max-width:460px;margin:18px auto 0;">'
            f'<div style="font-size:12px;font-weight:600;text-transform:uppercase;'
            f'letter-spacing:.15em;color:{COL["lo"]};margin-bottom:6px;">Forma</div>'
            f'<div style="display:flex;gap:8px;">'
            f'{_forma_chip(fc_v, COL["home"], casa_alto)}{_forma_chip(ft_v, COL["away"], trasf_alto)}'
            f'</div></div>'
        )

    # --- valore rosa (tiro alla fune) ---
    rosa_html = ""
    vc, vt = o["val_casa"], o["val_trasferta"]
    if _num_ok(vc) and _num_ok(vt) and (float(vc) + float(vt)) > 0:
        vc, vt = float(vc), float(vt)
        hp = vc / (vc + vt) * 100
        rosa_html = (
            f'<div style="{FONT}max-width:460px;margin:18px auto 0;">'
            f'<div style="font-size:12px;font-weight:600;text-transform:uppercase;'
            f'letter-spacing:.15em;color:{COL["lo"]};margin-bottom:6px;">Valore rosa</div>'
            f'<div style="border-radius:16px;padding:14px;background:{COL["panel"]};border:1px solid {COL["line"]};">'
            f'<div style="display:flex;justify-content:space-between;font-weight:700;font-size:19px;margin-bottom:8px;">'
            f'<span style="color:{COL["home"]};">{vc:g}</span><span style="color:{COL["away"]};">{vt:g}</span></div>'
            f'<div style="position:relative;height:14px;border-radius:99px;overflow:hidden;background:{COL["ink"]};display:flex;">'
            f'<div style="width:{hp:.1f}%;background:linear-gradient(90deg,{COL["home"]},{COL["home"]}cc);"></div>'
            f'<div style="width:{100-hp:.1f}%;background:linear-gradient(90deg,{COL["away"]}cc,{COL["away"]});"></div>'
            f'</div>'
            f'<div style="display:flex;justify-content:space-between;margin-top:8px;font-size:12px;color:{COL["lo"]};">'
            f'<span>{hp:.0f}%</span><span>rapporto rose</span><span>{100-hp:.0f}%</span></div>'
            f'</div></div>'
        )

    return hero + quote_html + forma_html + rosa_html


def render_lista(team, matches, venue):
    """venue='casa' evidenzia le gare in casa; venue='trasf' quelle in trasferta."""
    accent = COL["home"] if venue == "casa" else COL["away"]
    etichetta = "CASA" if venue == "casa" else "TRASF"
    righe = ""
    for _, m in matches.iterrows():
        casa, trasf = m["squadra_casa"], m["squadra_trasferta"]
        gc, gt = m.get("gol_casa"), m.get("gol_trasferta")
        sc = "?" if not _num_ok(gc) else int(gc)
        st_ = "?" if not _num_ok(gt) else int(gt)
        r = _esito(m, team)
        c = _col_esito(r)
        lettera = r or "–"
        evid = (venue == "casa" and casa == team) or (venue == "trasf" and trasf == team)
        bg = f"{accent}14" if evid else COL["panel"]
        left = f"border-left:3px solid {accent};" if evid else f"border-left:3px solid transparent;"
        badge = (f'<span style="font-size:9px;font-weight:700;letter-spacing:.06em;color:{accent};'
                 f'background:{accent}22;padding:2px 6px;border-radius:5px;margin-left:6px;">{etichetta}</span>'
                 ) if evid else ""
        wc = 600 if casa == team else 400
        wt = 600 if trasf == team else 400
        qualif = m.get("qualificatore")
        tipo = m.get("tipo_partita")
        extra = f' · {_esc(qualif)}' if qualif else ""
        if tipo and str(tipo) not in ("", "Non assegnata", "None"):
            extra += f' · {_esc(tipo)}'
        data_str = m["data"].strftime("%d.%m.%y") if hasattr(m["data"], "strftime") else _esc(m["data"])
        righe += (
            f'<div style="display:flex;align-items:center;gap:12px;border-radius:12px;padding:10px 12px;'
            f'{left}background:{bg};border:1px solid {COL["line"]};margin-bottom:8px;">'
            f'<span style="width:26px;height:26px;border-radius:8px;display:flex;align-items:center;'
            f'justify-content:center;font-weight:700;font-size:13px;color:{c};background:{c}22;'
            f'border:1px solid {c}55;flex-shrink:0;">{lettera}</span>'
            f'<div style="flex:1;min-width:0;">'
            f'<div style="font-size:14px;color:{COL["hi"]};white-space:nowrap;overflow:hidden;text-overflow:ellipsis;">'
            f'<span style="font-weight:{wc};">{_esc(casa)}</span>'
            f'<span style="color:{COL["lo"]};"> {sc}–{st_} </span>'
            f'<span style="font-weight:{wt};">{_esc(trasf)}</span>{badge}</div>'
            f'<div style="font-size:11px;color:{COL["lo"]};margin-top:2px;">{data_str} · {_esc(m.get("competizione"))}{extra}</div>'
            f'</div></div>'
        )
    return f'<div style="{FONT}max-width:460px;margin:0 auto;">{righe}</div>'


# =============================================================================
#  PAGINA: CONFIGURAZIONE
# =============================================================================
def salva_backtest_snapshot(n_valutate, min_storico, metriche, nota=None):
    cli = get_client()
    if not cli:
        return
    cli.table("backtest_snapshot").insert({
        "n_valutate": int(n_valutate),
        "min_storico": int(min_storico),
        "metriche_json": json.dumps(metriche),
        "nota": nota or None,
    }).execute()
    st.cache_data.clear()


@st.cache_data(ttl=300, show_spinner=False)
def carica_backtest_snapshot():
    cli = get_client()
    if not cli:
        return pd.DataFrame()
    try:
        res = cli.table("backtest_snapshot").select("*").order("creato_il", desc=True).execute()
        return pd.DataFrame(res.data or [])
    except Exception:
        return pd.DataFrame()


def _backtest_una_partita(df, comp_df, riga, recency_decay=None, home_adv=None):
    """Ricostruisce la probabilità pre-partita per UNA partita conclusa, usando solo i
    dati precedenti (walk-forward). Ritorna (prob, (nh,na), tre_pick) o None.
    tre_pick = {'motore':merc, 'statistico':merc, 'fusione':merc}."""
    home = riga["squadra_casa"]
    away = riga["squadra_trasferta"]
    data_p = riga["data"]
    pid = riga.get("id")
    t_liv = _livello_di(riga.get("competizione"), comp_df)
    t_cat = categoria_di(riga.get("competizione"), comp_df)
    t_key = _key(riga.get("competizione")) if riga.get("competizione") else None
    ph = _partite_squadra_evidenze(df, home, data_p, pid, comp_df, t_liv, t_cat, t_key, recency_decay)
    pa = _partite_squadra_evidenze(df, away, data_p, pid, comp_df, t_liv, t_cat, t_key, recency_decay)
    if not ph or not pa:
        return None
    hcap_h = _handicap_livello(ph, t_liv)
    hcap_a = _handicap_livello(pa, t_liv)
    ev = evidenze.costruisci_evidenze(ph, pa, odds=None, hcap_home=hcap_h, hcap_away=hcap_a,
                                      home_adv=home_adv)
    sig = segnali.calcola_signal(ev)
    stat = statistico.analizza(ph, pa)
    racc = racconto.racconta(home, away, ev, sig, statistico=stat)
    # stessi tre motori dello Storico: motore, fusione a media, solo statistico
    tre = {
        "motore": (racc.get("pronostico") or {}).get("mercato"),
        "fusione": (racc.get("fusione_media") or {}).get("mercato"),
        "statistico": (racc.get("solo_statistico") or {}).get("mercato"),
    }
    return ev["prob"], (len(ph), len(pa)), tre


def _snapshot_prematch_una(df, comp_df, riga):
    """Costruisce lo snapshot pre-match (feature + target) per UNA partita conclusa,
    usando solo i dati precedenti (walk-forward). Ritorna (features, target, nh, na) o None."""
    home = riga["squadra_casa"]
    away = riga["squadra_trasferta"]
    data_p = riga["data"]
    pid = riga.get("id")
    gc, gt = riga.get("gol_casa"), riga.get("gol_trasferta")
    if not (_num_ok(gc) and _num_ok(gt)):
        return None
    t_liv = _livello_di(riga.get("competizione"), comp_df)
    t_cat = categoria_di(riga.get("competizione"), comp_df)
    t_key = _key(riga.get("competizione")) if riga.get("competizione") else None
    ph = _partite_squadra_evidenze(df, home, data_p, pid, comp_df, t_liv, t_cat, t_key)
    pa = _partite_squadra_evidenze(df, away, data_p, pid, comp_df, t_liv, t_cat, t_key)
    if not ph or not pa:
        return None
    hcap_h = _handicap_livello(ph, t_liv)
    hcap_a = _handicap_livello(pa, t_liv)
    ev = evidenze.costruisci_evidenze(ph, pa, odds=None, hcap_home=hcap_h, hcap_away=hcap_a)
    sig = segnali.calcola_signal(ev)
    feat = snapmod.costruisci_snapshot(ph, pa, ev, sig)
    tgt = snapmod.costruisci_target(gc, gt)
    return feat, tgt, len(ph), len(pa)


def salva_snapshot_prematch(partita_id, data, home, away, competizione, feat, tgt, nh, na):
    """Salva (upsert) uno snapshot pre-match. Ritorna (True, None) o (False, messaggio_errore)."""
    cli = get_client()
    if not cli:
        return False, "nessun client Supabase"
    rec = {
        "partita_id": str(partita_id), "data": str(data) if data is not None else None,
        "squadra_casa": home, "squadra_trasferta": away,
        "competizione": competizione,
        "features_json": json.dumps(feat), "target_json": json.dumps(tgt),
        "n_home": nh, "n_away": na,
    }
    try:
        esiste = cli.table("snapshot_prematch").select("id").eq(
            "partita_id", str(partita_id)).execute()
        if esiste.data:
            cli.table("snapshot_prematch").update(rec).eq(
                "partita_id", str(partita_id)).execute()
        else:
            cli.table("snapshot_prematch").insert(rec).execute()
        return True, None
    except Exception as e:
        return False, str(e)
        return False


def genera_snapshot_prematch(df, comp_df, progress=None):
    """Genera e salva gli snapshot pre-match per le fixture pronosticate concluse.
    Ritorna (creati, saltati, primo_errore)."""
    cli = get_client()
    if not cli or df.empty:
        return 0, 0, "nessun client o database vuoto"
    gia = set()
    try:
        r = cli.table("snapshot_prematch").select("partita_id").execute()
        gia = {str(x["partita_id"]) for x in (r.data or [])}
    except Exception as e:
        return 0, 0, f"lettura tabella fallita: {e}"
    _mask = df["gol_casa"].notna() & df["gol_trasferta"].notna()
    if "is_target" in df.columns:
        concl = df[_mask & (df["is_target"] == True)]
    else:
        concl = df[_mask] if "gol_casa" in df.columns else df.iloc[0:0]
    creati = saltati = 0
    primo_errore = None
    righe = list(concl.iterrows())
    for i, (_, riga) in enumerate(righe):
        if progress and i % 5 == 0:
            progress.progress(i / max(1, len(righe)), text=f"Snapshot {i+1}/{len(righe)}…")
        pid = str(riga.get("id"))
        if pid in gia:
            saltati += 1
            continue
        try:
            res = _snapshot_prematch_una(df, comp_df, riga)
            if res is None:
                saltati += 1
                if primo_errore is None:
                    primo_errore = "storico insufficiente per alcune partite"
                continue
            feat, tgt, nh, na = res
            ok, err = salva_snapshot_prematch(
                pid, riga.get("data"), riga["squadra_casa"], riga["squadra_trasferta"],
                _label_da_comp(riga.get("competizione"), comp_df), feat, tgt, nh, na)
            if ok:
                creati += 1
            else:
                saltati += 1
                if primo_errore is None and err:
                    primo_errore = err
        except Exception as e:
            saltati += 1
            if primo_errore is None:
                primo_errore = str(e)
    return creati, saltati, primo_errore


def _confronto_book(df, comp_df, concluse, recency_decay=None, home_adv=None):
    """Confronta il motore col favorito del bookmaker sull'1X2, sulle partite dove ci sono
    tutte e tre le quote (1, X, 2). Ritorna un dict con: n, accuratezze, e i casi di
    DISSENSO (motore != favorito book) con chi ha ragione. None se nessuna partita utile."""
    n = 0
    mot_ok = book_ok = 0
    diss_tot = diss_mot_ok = diss_book_ok = 0
    for _, r in concluse.iterrows():
        gc, gt = r.get("gol_casa"), r.get("gol_trasferta")
        if not (_num_ok(gc) and _num_ok(gt)):
            continue
        q1 = _quota_storica(r, "1"); qx = _quota_storica(r, "X"); q2 = _quota_storica(r, "2")
        if not (q1 and qx and q2):
            continue
        esito = "1" if gc > gt else ("X" if gc == gt else "2")
        fav = min((("1", q1), ("X", qx), ("2", q2)), key=lambda t: t[1])[0]
        res = _backtest_una_partita(df, comp_df, r, recency_decay, home_adv)
        if res is None:
            continue
        prob = res[0]
        mot = max((("1", prob.get("1", 0)), ("X", prob.get("X", 0)), ("2", prob.get("2", 0))),
                  key=lambda t: t[1])[0]
        n += 1
        if mot == esito:
            mot_ok += 1
        if fav == esito:
            book_ok += 1
        if mot != fav:
            diss_tot += 1
            if mot == esito:
                diss_mot_ok += 1
            if fav == esito:
                diss_book_ok += 1
    if n == 0:
        return None
    return {
        "n": n, "mot_ok": mot_ok, "book_ok": book_ok,
        "acc_mot": mot_ok / n * 100, "acc_book": book_ok / n * 100,
        "diss_tot": diss_tot, "diss_mot_ok": diss_mot_ok, "diss_book_ok": diss_book_ok,
    }


def _report_backtest_testo(valutate, saltate, min_storico, tab, tab_tre, cal_sel, merc_sel, tab_roi, cal_tutti=None):
    """Costruisce un riepilogo TESTUALE del backtest (per copia-incolla rapido)."""
    L = []
    L.append(f"BACKTEST — {datetime.now():%Y-%m-%d %H:%M}")
    L.append(f"Valutate {valutate} partite · saltate {saltate} · storico minimo {min_storico}")
    L.append("")
    L.append("QUALITÀ PROBABILITÀ (Brier più basso = meglio; batte baseline = utile):")
    for r in tab:
        L.append(f"  {r['Mercato']:10s} N={r['N']:4d}  Brier {r['Brier']:.4f}  "
                 f"baseline {r['Baseline']:.4f}  {r['Batte baseline']}  "
                 f"LogLoss {r['Log Loss']:.4f}  calibErr {r['Errore calib.']}")
    L.append("")
    L.append("CONFRONTO TRE MOTORI (hit-rate pronostico di punta):")
    for r in tab_tre:
        L.append(f"  {r['Motore']:14s} valutati {r['Pronostici valutati']:4d}  "
                 f"azzeccati {r['Azzeccati']:4d}  riuscita {r['Riuscita %']}%")
    L.append("")
    if cal_tutti:
        L.append("CALIBRAZIONE (tutti i mercati):")
        for merc, cal in cal_tutti.items():
            L.append(f"  [{merc}]")
            for r in cal:
                L.append(f"    fascia {r['fascia']:8s} N={r['n']:4d}  prob media {r['prob_media']}%  "
                         f"reale {r['reale']}%  scarto {r['scarto']}pt")
    else:
        L.append(f"CALIBRAZIONE ({merc_sel}):")
        for r in cal_sel:
            L.append(f"  fascia {r['fascia']:8s} N={r['n']:4d}  prob media {r['prob_media']}%  "
                     f"reale {r['reale']}%  scarto {r['scarto']}pt")
    if tab_roi:
        L.append("")
        L.append("REDDITIVITÀ (dove ci sono quote):")
        for r in tab_roi:
            L.append(f"  {r['Mercato']:10s} giocate {r['Giocate (value)']:4d}  "
                     f"profitto {r['Profitto (u)']}u  ROI {r['ROI %']}%  "
                     f"maxDD {r['Max drawdown (u)']}u")
    return "\n".join(L)


def _report_backtest_pdf(valutate, saltate, min_storico, tab, tab_tre, cal_sel, merc_sel, tab_roi):
    """Genera il PDF del backtest in memoria (BytesIO)."""
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from io import BytesIO
    buf = BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=A4, topMargin=36, bottomMargin=36,
                            leftMargin=36, rightMargin=36)
    ss = getSampleStyleSheet()
    story = [Paragraph("Backtest — Estrattore Partite", ss["Title"]),
             Paragraph(f"{datetime.now():%d/%m/%Y %H:%M} · Valutate {valutate} · "
                       f"saltate {saltate} · storico minimo {min_storico}", ss["Normal"]),
             Spacer(1, 12)]

    def _tabella(titolo, intest, righe):
        story.append(Paragraph(titolo, ss["Heading2"]))
        dati = [intest] + righe
        t = Table(dati, hAlign="LEFT")
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#2c3e50")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.4, colors.grey),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f2f2f2")]),
        ]))
        story.append(t)
        story.append(Spacer(1, 12))

    _tabella("Qualità delle probabilità",
             ["Mercato", "N", "Brier", "Baseline", "Batte", "LogLoss", "CalibErr"],
             [[r["Mercato"], r["N"], f"{r['Brier']:.4f}", f"{r['Baseline']:.4f}",
               "Si" if "✅" in r["Batte baseline"] else "No",
               f"{r['Log Loss']:.4f}", r["Errore calib."]] for r in tab])
    _tabella("Confronto tra i tre motori",
             ["Motore", "Valutati", "Azzeccati", "Riuscita %"],
             [[r["Motore"].replace("🎯 ", "").replace("📊 ", "").replace("🏆 ", ""),
               r["Pronostici valutati"], r["Azzeccati"], r["Riuscita %"]] for r in tab_tre])
    _tabella(f"Calibrazione — {merc_sel}",
             ["Fascia", "N", "Prob media %", "Reale %", "Scarto pt"],
             [[r["fascia"], r["n"], r["prob_media"], r["reale"], r["scarto"]] for r in cal_sel])
    if tab_roi:
        _tabella("Redditività (dove ci sono quote)",
                 ["Mercato", "Giocate", "Profitto u", "ROI %", "MaxDD u"],
                 [[r["Mercato"], r["Giocate (value)"], r["Profitto (u)"],
                   r["ROI %"], r["Max drawdown (u)"]] for r in tab_roi])
    doc.build(story)
    buf.seek(0)
    return buf.getvalue()


def pagina_backtest(user):
    st.title("🧪 Backtest walk-forward")
    st.caption("Per ogni partita conclusa, il motore ricostruisce la probabilità pre-partita "
               "usando SOLO i dati precedenti (niente informazioni dal futuro) e la confronta "
               "con il risultato reale. Le metriche principali misurano la qualità delle "
               "probabilità, indipendentemente dalle quote.")

    df = carica_partite()
    if df.empty:
        st.info("Nessuna partita nel database.")
        return
    comp_df = carica_competizioni()
    # SOLO le fixture pronosticate (is_target) con risultato: sono le "vere" partite del
    # dataset. Le altre partite concluse sono lo STORICO incollato (materiale grezzo per
    # calcolare i pronostici), non vanno valutate né usate come campione.
    _conclmask = df["gol_casa"].notna() & df["gol_trasferta"].notna()
    if "is_target" in df.columns:
        concluse = df[_conclmask & (df["is_target"] == True)].copy()
    else:
        concluse = df[_conclmask].copy()
    st.markdown(f"Partite pronosticate concluse (valutabili): **{len(concluse)}**")
    st.caption("Sono le fixture che hai pronosticato e che poi hanno avuto un risultato. "
               "Lo storico incollato per calcolare i pronostici non entra nel campione.")

    # --- snapshot pre-match per il futuro Learning Engine (ML) ---
    with st.expander("🧠 Dataset Learning Engine (snapshot pre-match)"):
        st.caption("Salva la 'fotografia' pre-partita (feature calcolate solo con i dati "
                   "precedenti) più il risultato reale, per ogni partita conclusa. Accumula nel "
                   "tempo un dataset onesto (walk-forward) per un futuro modello ML. Non modifica "
                   "nulla dei motori attuali: raccoglie soltanto dati.")
        n_snap = 0
        try:
            _cli = get_client()
            if _cli:
                _r = _cli.table("snapshot_prematch").select("id", count="exact").execute()
                n_snap = _r.count if hasattr(_r, "count") and _r.count is not None else len(_r.data or [])
        except Exception:
            n_snap = None
        if n_snap is not None:
            st.markdown(f"Snapshot già salvati: **{n_snap}**")
        if st.button("📸 Genera snapshot mancanti"):
            _pr = st.progress(0.0, text="Costruzione snapshot…")
            _c, _s, _err = genera_snapshot_prematch(df, comp_df, _pr)
            _pr.progress(1.0, text="Completato.")
            st.success(f"Creati {_c} nuovi snapshot · {_s} saltati.")
            if _err:
                st.error(f"⚠️ Motivo dei saltati (primo errore): {_err}")
        st.caption("Suggerimento: rigenera ogni tanto (es. dopo aver inserito nuovi risultati) "
                   "per far crescere il dataset. Quando avrai molte più partite, da qui "
                   "costruiremo e valideremo il modello ML.")

    # --- diario di bordo: andamento dei backtest salvati nel tempo ---
    snap = carica_backtest_snapshot()
    if not snap.empty:
        with st.expander(f"📔 Diario di bordo ({len(snap)} snapshot salvati)", expanded=False):
            righe = []
            for _, s in snap.iterrows():
                try:
                    mj = json.loads(s.get("metriche_json") or "{}")
                except Exception:
                    mj = {}
                # nuova struttura {mercati:{}, tre_motori:{}} o vecchia piatta
                mercati = mj.get("mercati", mj) if isinstance(mj, dict) else {}
                tre = mj.get("tre_motori", {}) if isinstance(mj, dict) else {}
                battuti = sum(1 for v in mercati.values()
                              if isinstance(v, dict) and v.get("batte_baseline"))
                brier_medio = [v["brier"] for v in mercati.values()
                               if isinstance(v, dict) and v.get("brier") is not None]
                bm = round(sum(brier_medio) / len(brier_medio), 4) if brier_medio else None
                fus = tre.get("fusione", {}).get("riuscita") if tre else None
                data_s = str(s.get("creato_il", ""))[:16].replace("T", " ")
                righe.append({"Data": data_s, "N valutate": s.get("n_valutate"),
                              "Mercati che battono baseline": f"{battuti}/{len(mercati)}",
                              "Brier medio": bm,
                              "Fusione riuscita %": fus,
                              "Nota": s.get("nota") or ""})
            st.dataframe(pd.DataFrame(righe), use_container_width=True, hide_index=True)
            st.caption("Con più snapshot nel tempo vedrai se il motore migliora: Brier medio "
                       "in calo e più mercati che battono il baseline = progresso reale.")

    c = st.columns(2)
    min_storico = c[0].slider("Storico minimo per squadra", 5, 15, 8,
                              help="Valuta una partita solo se entrambe le squadre hanno "
                                   "almeno questo numero di partite precedenti.")
    max_part = c[1].slider("Max partite da valutare", 50, 1000, 300, step=50,
                           help="Limita il calcolo (le più recenti). Più alto = più lento.")
    c2 = st.columns(2)
    usa_recency = c2[0].checkbox("Attiva decadimento temporale (recency)", value=False,
                                 help="Pesa di più le partite recenti. Provalo e confronta "
                                      "il risultato con/senza per vedere se aiuta.")
    recency_gg = c2[1].slider("Decadimento (giorni)", 20, 180, 60, step=10,
                              disabled=not usa_recency,
                              help="Più basso = più peso alle partite recentissime. "
                                   "peso = exp(-giorni/decadimento).")
    recency_decay = recency_gg if usa_recency else None

    c3 = st.columns(2)
    usa_hadv = c3[0].checkbox("Testa vantaggio campo diverso", value=False,
                              help="Il default è 1.06 (casa segna +6%, ospite -6%). "
                                   "Alza per dare più peso al giocare in casa.")
    hadv_val = c3[1].slider("Vantaggio campo", 1.00, 1.20, 1.06, step=0.02,
                            disabled=not usa_hadv,
                            help="1.06 = default. 1.10-1.14 = campo più forte "
                                 "(riduce le vittorie ospite previste).")
    home_adv = hadv_val if usa_hadv else None

    if not st.button("▶️ Esegui backtest", type="primary"):
        return

    # ordina dalla più recente e limita
    if "data" in concluse.columns:
        concluse = concluse.sort_values("data", ascending=False)
    concluse = concluse.head(max_part)

    # raccogli (prob, esito) per mercato
    dati = {m: [] for m in bt.MERCATI_BINARI}
    scommesse = {m: [] for m in bt.MERCATI_BINARI}
    valutate = saltate = 0
    prog = st.progress(0.0, text="Ricostruzione pre-partita…")
    righe = list(concluse.iterrows())
    hit = {"motore": [0, 0], "statistico": [0, 0], "fusione": [0, 0]}  # [azzeccati, sbagliati]
    for i, (_, r) in enumerate(righe):
        if i % 10 == 0:
            prog.progress(i / max(1, len(righe)), text=f"Partita {i+1}/{len(righe)}…")
        try:
            gc, gt = int(r["gol_casa"]), int(r["gol_trasferta"])
            res = _backtest_una_partita(df, comp_df, r, recency_decay, home_adv)
            if not res:
                saltate += 1
                continue
            prob, (nh, na), tre = res
            if min(nh, na) < min_storico:
                saltate += 1
                continue
            valutate += 1
            for m, (chiave, fesito) in bt.MERCATI_BINARI.items():
                p = prob.get(chiave)
                if p is None:
                    continue
                y = fesito(gc, gt)
                dati[m].append((p / 100.0, y))
                q = _quota_storica(r, chiave)
                if q and p / 100.0 * q - 1 > 0:
                    scommesse[m].append((y, q))
            # hit-rate dei tre motori sul loro pronostico di punta
            for eng in ("motore", "statistico", "fusione"):
                merc = tre.get(eng)
                if not merc:
                    continue
                base = merc.split(" (")[0].replace(" totali", "")
                won = _pronostico_vinto(base, gc, gt)
                if won is True:
                    hit[eng][0] += 1
                elif won is False:
                    hit[eng][1] += 1
        except Exception:
            saltate += 1
    prog.progress(1.0, text="Completato.")

    st.success(f"Valutate {valutate} partite · saltate {saltate} (storico insufficiente).")
    if valutate == 0:
        st.warning("Nessuna partita con storico sufficiente. Abbassa lo storico minimo.")
        return

    # ---- CONFRONTO TRE MOTORI (hit-rate del pronostico di punta) ----
    st.subheader("Confronto tra i tre motori (walk-forward)")
    st.caption("Percentuale di volte in cui il pronostico DI PUNTA di ogni motore si è "
               "verificato, ricostruito pre-partita. È il confronto diretto di affidabilità.")
    tab_tre = []
    for eng, nome in (("motore", "🎯 Motore"), ("statistico", "📊 Statistico"), ("fusione", "🏆 Fusione")):
        v, p = hit[eng]
        tot = v + p
        tab_tre.append({"Motore": nome, "Pronostici valutati": tot,
                        "Azzeccati": v, "Riuscita %": round(v / tot * 100, 1) if tot else None})
    st.dataframe(pd.DataFrame(tab_tre), use_container_width=True, hide_index=True)
    hit_snap = {eng: {"azzeccati": hit[eng][0], "totale": hit[eng][0] + hit[eng][1],
                      "riuscita": round(hit[eng][0] / (hit[eng][0] + hit[eng][1]) * 100, 1)
                      if (hit[eng][0] + hit[eng][1]) else None}
                for eng in ("motore", "statistico", "fusione")}

    # ---- CONFRONTO COL BOOKMAKER (1X2) ----
    st.subheader("🆚 Motore vs Bookmaker (1X2)")
    cb = _confronto_book(df, comp_df, concluse, recency_decay, home_adv)
    if not cb:
        st.caption("Nessuna partita con tutte e tre le quote 1, X, 2 salvate: "
                   "il confronto col bookmaker non è disponibile.")
    else:
        st.caption(f"Confronto su {cb['n']} partite dove sono salvate tutte le quote 1/X/2. "
                   "Il 'favorito' del book è l'esito con la quota più bassa.")
        cc = st.columns(2)
        cc[0].metric("🎯 Motore azzecca 1X2", f"{cb['acc_mot']:.0f}%",
                     help=f"{cb['mot_ok']}/{cb['n']} esiti 1X2 corretti.")
        cc[1].metric("🏦 Favorito book azzecca", f"{cb['acc_book']:.0f}%",
                     help=f"{cb['book_ok']}/{cb['n']} esiti 1X2 corretti.")
        if cb["diss_tot"] > 0:
            st.markdown(f"**Quando il motore dissente dal book** ({cb['diss_tot']} partite "
                        "dove il motore indica un esito diverso dal favorito del book):")
            dd = st.columns(2)
            dd[0].metric("Motore ha ragione", f"{cb['diss_mot_ok']}/{cb['diss_tot']}",
                         help="Nei dissensi, quante volte l'esito del motore è quello reale.")
            dd[1].metric("Book ha ragione", f"{cb['diss_book_ok']}/{cb['diss_tot']}")
            if cb["diss_mot_ok"] > cb["diss_book_ok"]:
                st.success("Nei casi di dissenso, il motore batte il book: è lì che può "
                           "nascere il valore (il book sta sbagliando il favorito).")
            elif cb["diss_mot_ok"] < cb["diss_book_ok"]:
                st.info("Nei casi di dissenso vince più spesso il book. Normale: battere "
                        "l'accuratezza del book è difficilissimo. Il valore si cerca sul "
                        "PREZZO (ROI/EV), non sull'azzeccare più esiti.")
            else:
                st.info("Nei dissensi motore e book pari. Campione piccolo: leggi con cautela.")
        st.caption("⚠️ Atteso: il book è quasi imbattibile sull'ACCURATEZZA pura. Il vero "
                   "vantaggio si misura sul ROI (valore sul prezzo), non su chi azzecca di più.")

    # ---- METRICHE QUALITÀ MODELLO (protagoniste) ----
    st.subheader("Qualità delle probabilità (indipendente dalle quote)")
    st.caption("Brier e Log Loss: più bassi è meglio. 'Batte baseline' = il modello è "
               "meglio del predire sempre la frequenza media. Errore calib.: quanto le "
               "probabilità dichiarate corrispondono alla realtà (0 = perfetto).")
    tab = []
    metriche_snap = {}
    for m in bt.MERCATI_BINARI:
        c = dati[m]
        if not c:
            continue
        br = bt.brier(c)
        base = bt.baseline_brier(c)
        ll = bt.log_loss(c)
        ce = bt.calibration_error(c)
        tab.append({
            "Mercato": m, "N": len(c),
            "Brier": round(br, 4),
            "Baseline": round(base, 4),
            "Batte baseline": "✅" if br < base else "❌",
            "Log Loss": round(ll, 4),
            "Errore calib.": ce,
        })
        metriche_snap[m] = {"n": len(c), "brier": round(br, 4), "baseline": round(base, 4),
                            "log_loss": round(ll, 4), "calib_err": ce,
                            "batte_baseline": bool(br < base)}
    st.dataframe(pd.DataFrame(tab), use_container_width=True, hide_index=True)

    # --- salva snapshot nel diario di bordo ---
    with st.expander("💾 Salva questo risultato nel diario di bordo"):
        nota = st.text_input("Nota (facoltativa)", placeholder="es. stato attuale, prima di correggere i λ")
        if st.button("💾 Salva snapshot"):
            try:
                salva_backtest_snapshot(valutate, min_storico,
                                        {"mercati": metriche_snap, "tre_motori": hit_snap}, nota)
                st.success("Snapshot salvato: potrai confrontare i miglioramenti nel tempo.")
            except Exception as e:
                st.error(f"Salvataggio non riuscito (hai creato la tabella backtest_snapshot?): {e}")

    # ---- CALIBRAZIONE dettagliata per mercato ----
    st.subheader("Curva di calibrazione")
    merc_sel = st.selectbox("Mercato", list(bt.MERCATI_BINARI.keys()))
    cal = bt.calibration(dati[merc_sel], n_bin=5)
    if cal:
        cdf = pd.DataFrame(cal)
        cdf.columns = ["Fascia prob.", "N", "Prob. media %", "Reale %", "Scarto pt"]
        st.dataframe(cdf, use_container_width=True, hide_index=True)
        st.caption("Se 'Prob. media' ≈ 'Reale', il modello è ben calibrato in quella fascia.")

    # ---- ROI (secondario, solo dove ci sono quote) ----
    st.subheader("Redditività (solo dove esistono quote storiche)")
    tab_roi = []
    for m in bt.MERCATI_BINARI:
        r = bt.roi_yield(scommesse[m])
        if r and r["n"] >= 5:
            tab_roi.append({"Mercato": m, "Giocate (value)": r["n"],
                            "Profitto (u)": r["profitto"], "ROI %": r["roi"],
                            "Max drawdown (u)": r["max_drawdown"]})
    if tab_roi:
        st.dataframe(pd.DataFrame(tab_roi), use_container_width=True, hide_index=True)
        st.caption("Simulazione: 1 unità sulle giocate dove il modello vede value (EV>0). "
                   "Richiede quote storiche salvate: dove mancano, la partita conta solo per "
                   "le metriche del modello.")
    else:
        st.caption("Poche o nessuna quota storica disponibile: la redditività non è "
                   "calcolabile in modo affidabile. Le metriche del modello sopra restano valide.")

    # ---- ESPORTAZIONE: PDF + testo copiabile ----
    st.divider()
    st.subheader("📄 Esporta risultati")
    # calibrazione di TUTTI i mercati (per l'esportazione completa, senza toccare il menu)
    cal_tutti = {}
    for m in bt.MERCATI_BINARI:
        c = bt.calibration(dati[m], n_bin=5)
        if c:
            cal_tutti[m] = c
    testo = _report_backtest_testo(valutate, saltate, min_storico, tab, tab_tre, cal, merc_sel,
                                   tab_roi, cal_tutti)
    ce = st.columns(2)
    try:
        pdf_bytes = _report_backtest_pdf(valutate, saltate, min_storico, tab, tab_tre, cal, merc_sel, tab_roi)
        ce[0].download_button("⬇️ Scarica PDF", pdf_bytes,
                              file_name=f"backtest_{datetime.now():%Y%m%d_%H%M}.pdf",
                              mime="application/pdf")
    except Exception as e:
        ce[0].caption(f"PDF non disponibile: {e}")
    ce[1].download_button("⬇️ Scarica testo", testo.encode("utf-8"),
                          file_name=f"backtest_{datetime.now():%Y%m%d_%H%M}.txt",
                          mime="text/plain")
    st.caption("Oppure copia direttamente da qui (clic sull'icona in alto a destra del riquadro):")
    st.code(testo, language="text")


def _quota_storica(riga, mercato):
    """Quota salvata per un mercato, se presente (solo partite già analizzate come target)."""
    m = {"Over 2.5": "quota_iniziale_over", "1": "quota_iniziale_1",
         "X": "quota_iniziale_x", "2": "quota_iniziale_2", "Goal": "quota_iniziale_goal"}
    col = m.get(mercato)
    if not col or col not in riga:
        return None
    try:
        q = float(riga[col])
        return q if q > 1.0 else None
    except (TypeError, ValueError):
        return None


def pagina_configurazione(user):
    st.header("⚙️ Configurazione")

    if user["ruolo"] != "admin":
        st.info("Solo gli amministratori possono gestire utenti e backup.")
        return

    st.subheader("👥 Utenti")
    utenti = carica_utenti()
    if utenti:
        st.dataframe(
            pd.DataFrame(utenti)[["username", "ruolo", "creato_il"]],
            use_container_width=True, hide_index=True)

    with st.form("nuovo_utente"):
        st.markdown("**Nuovo utente**")
        u = st.text_input("Username")
        p = st.text_input("Password", type="password")
        ruolo = st.selectbox("Ruolo", ["user", "admin"])
        if st.form_submit_button("Crea utente"):
            if not u or not p:
                st.error("Compila username e password.")
            else:
                try:
                    crea_utente(u, p, ruolo)
                    st.success(f"Utente '{u}' creato.")
                    st.rerun()
                except Exception as e:
                    st.error(f"Errore: {e}")

    st.divider()
    st.subheader("📄 Archivio partite (Word)")
    st.caption("Scarica un documento Word con una pagina per partita seguita: intestazione "
               "(squadre, campionato, data, ora), ultime partite delle due squadre, e il "
               "risultato. I file vengono generati solo quando premi 'Prepara' (così la "
               "pagina resta veloce).")
    if st.button("🔄 Prepara i file Word", key="prepara_word"):
        try:
            dfp = carica_partite()
            cdf = carica_competizioni()
            with st.spinner("Genero i documenti Word…"):
                st.session_state["_docx_nuova"] = genera_docx_nuova_analisi(dfp, cdf)
                st.session_state["_docx_mercati"] = genera_docx_mercati(dfp, cdf)
                st.session_state["_docx_senza"] = genera_docx_archivio(dfp, cdf, con_analisi=False)
            st.success("File pronti: usa i pulsanti di download qui sotto.")
        except Exception as e:
            st.error(f"Impossibile generare il Word: {e}")
    if st.session_state.get("_docx_nuova"):
        ca = st.columns(3)
        ca[0].download_button(
            "⬇️ Partite con analisi", data=st.session_state["_docx_nuova"],
            file_name="archivio_analisi.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            help="Il nuovo racconto completo: motore, statistico, fusione e i pronostici, "
                 "più i dati delle partite.")
        ca[1].download_button(
            "⬇️ Solo mercati", data=st.session_state.get("_docx_mercati", b""),
            file_name="archivio_mercati.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            help="Snello: squadra A - squadra B, le sezioni Over/Under, Goal/No Goal, 1X2 "
                 "(signal/prob/quota/EV/edge) e il risultato.")
        ca[2].download_button(
            "⬇️ Senza analisi", data=st.session_state["_docx_senza"],
            file_name="archivio_partite.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            help="Solo intestazione, ultime partite, quote e risultato.")
        st.caption("Nota: i file riflettono i dati di quando hai premuto 'Prepara'. "
                   "Se aggiorni partite o quote, ripremi 'Prepara' per rigenerarli.")

    st.divider()
    st.subheader("🏆 Competizioni")
    st.caption("Assegna a ogni competizione la categoria e un nome corto. Il nome corto "
               "collega il nome lungo (estrattore risultati) al codice usato nello storico "
               "(estrattore ultimi risultati e quote).")

    comp_df = carica_competizioni()

    # rileva i codici brevi presenti nello storico ma non ancora in anagrafica
    part = carica_partite()
    if not part.empty and "competizione" in part:
        codici = sorted({str(c).strip() for c in part["competizione"].dropna()
                         if str(c).strip()})
        gia_noti = set()
        if not comp_df.empty:
            for _, c in comp_df.iterrows():
                gia_noti |= _chiavi_competizione(c)
        mancanti = [c for c in codici if _key(c) not in gia_noti]
        if mancanti:
            st.caption("Codici trovati nello storico non ancora in anagrafica: "
                       + ", ".join(mancanti))
            if st.button("➕ Aggiungi codici mancanti"):
                upsert_competizioni([{"nome_corto": c, "categoria": "Non assegnata"}
                                     for c in mancanti])
                st.rerun()

    if comp_df.empty:
        st.info("Nessuna competizione ancora registrata. Compaiono qui dopo un import "
                "dall'estrattore risultati o aggiungendo i codici dello storico.")
    else:
        vista = pd.DataFrame({
            "id": comp_df["id"],
            "Nome lungo": comp_df.get("nome_lungo"),
            "Nazione": comp_df.get("nazione"),
            "Nome corto": comp_df.get("nome_corto"),
            "Categoria": comp_df.get("categoria").fillna("Non assegnata"),
            "Livello": ([None if pd.isna(x) else int(x) for x in comp_df["livello"]]
                        if "livello" in comp_df else [None] * len(comp_df)),
        })
        edit = st.data_editor(
            vista, use_container_width=True, hide_index=True, num_rows="dynamic",
            key="editor_comp",
            column_config={
                "id": None,
                "Categoria": st.column_config.SelectboxColumn(options=CATEGORIE),
                "Livello": st.column_config.SelectboxColumn(
                    help="Solo per i campionati: 1 = prima divisione, 2 = seconda, 3 = terza… "
                         "Lascia vuoto per coppe e amichevoli.",
                    options=[1, 2, 3, 4, 5, 6, 7, 8, 9]),
            },
        )

        col1, col2 = st.columns(2)
        if col1.button("💾 Salva competizioni", type="primary"):
            records = []
            for _, r in edit.iterrows():
                rec = {
                    "nome_lungo": (r["Nome lungo"] or None) if not pd.isna(r["Nome lungo"]) else None,
                    "nazione": (r["Nazione"] or None) if not pd.isna(r["Nazione"]) else None,
                    "nome_corto": (r["Nome corto"] or None) if not pd.isna(r["Nome corto"]) else None,
                    "categoria": r["Categoria"] or "Non assegnata",
                    "livello": None if pd.isna(r["Livello"]) else int(r["Livello"]),
                }
                if not pd.isna(r["id"]):
                    rec["id"] = r["id"]
                records.append(rec)
            # eliminazioni: id presenti prima ma non più nell'editor
            ids_ora = {r["id"] for _, r in edit.iterrows() if not pd.isna(r["id"])}
            for cid in set(comp_df["id"]) - ids_ora:
                elimina_competizione(cid)
            try:
                if records:
                    upsert_competizioni(records)
                st.success("Competizioni salvate.")
                st.rerun()
            except Exception as e:
                st.error(f"Errore: {e}")

        if col2.button("🔗 Applica categorie alle partite"):
            comp_now = carica_competizioni()
            updates = []
            for _, m in part.iterrows():
                cat = categoria_di(m.get("competizione"), comp_now)
                if cat and cat != m.get("tipo_partita"):
                    updates.append({"id": m["id"], "tipo_partita": cat})
            try:
                if updates:
                    aggiorna_partite(updates)
                st.success(f"Categoria aggiornata su {len(updates)} partite.")
                st.cache_data.clear()
            except Exception as e:
                st.error(f"Errore: {e}")

    st.divider()
    st.subheader("💾 Backup Excel dell'App")
    if st.button("Genera backup"):
        partite = carica_partite()
        u_df = pd.DataFrame(utenti)
        if not u_df.empty:
            u_df = u_df.drop(columns=[c for c in ["password_hash"] if c in u_df.columns])
        xlsx = to_excel({"partite": partite_per_export(partite),
                         "competizioni": carica_competizioni(), "utenti": u_df})
        st.download_button(
            "⬇️ Scarica backup.xlsx", data=xlsx, file_name="backup_app.xlsx",
            mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        st.caption("Nota: le password (hash) non sono incluse nel backup per sicurezza.")


# =============================================================================
#  PAGINA: ANALISI & PRONOSTICO
# =============================================================================
def _eff_odds(row):
    """Quota effettiva (variazione se presente, altrimenti iniziale) per il modello."""
    def e(ini, var):
        if _num_ok(row.get(var)):
            return float(row.get(var))
        return float(row.get(ini)) if _num_ok(row.get(ini)) else None
    return {
        "1": e("quota_iniziale_1", "variazione_quota_1"),
        "X": e("quota_iniziale_x", "variazione_quota_x"),
        "2": e("quota_iniziale_2", "variazione_quota_2"),
        "over25": e("quota_iniziale_over", "variazione_quota_over"),
        "under25": e("quota_iniziale_under", "variazione_quota_under"),
        "goal": e("quota_iniziale_goal", "variazione_quota_goal"),
        "nogoal": e("quota_iniziale_nogoal", "variazione_quota_nogoal"),
    }


def _mappa_livelli(comp_df):
    """Mappa competizione->livello, con tutte le forme (corto, lungo, lungo+nazione)."""
    livelli = {}
    if comp_df is not None and not comp_df.empty and "livello" in comp_df:
        for _, cc in comp_df.iterrows():
            liv = cc.get("livello")
            try:
                if pd.isna(liv):
                    continue
            except (TypeError, ValueError):
                if liv is None:
                    continue
            for kk in _chiavi_competizione(cc):
                livelli[kk] = int(liv)
    return livelli


def _variazioni_da_row(row):
    """Estrae le variazioni di quota dalla riga partita, nel formato del motore."""
    variazioni = {}
    for kk in ("1", "x", "2", "over", "under", "goal", "nogoal"):
        v = row.get(f"variazione_quota_{kk}")
        try:
            if v is not None and not pd.isna(v):
                mapk = {"x": "X", "over": "over25", "under": "under25"}.get(kk, kk)
                variazioni[mapk] = float(v)
        except (TypeError, ValueError):
            pass
    return variazioni


def _analizza_row(df, row, comp_df, config=None, calibratori=None):
    """Prepara tutti gli input dalla riga partita e lancia il motore. Usato sia dalla
    pagina Analisi sia dal dettaglio nel Database (stessa logica, nessuna divergenza)."""
    home, away = row["squadra_casa"], row["squadra_trasferta"]
    odds = _eff_odds(row)
    rose = (row.get("val_casa"), row.get("val_trasferta"))
    tp = row.get("tipo_partita")
    tipo_target = tp if (tp and str(tp) not in ("ND", "Non assegnata", "None")) else None
    a = analisi.analizza_partita(
        home, away, df, odds=odds, data_partita=row.get("data"), config=config,
        calibratori=calibratori, rose=rose, tipo_partita_target=tipo_target,
        livelli=_mappa_livelli(comp_df), variazioni=_variazioni_da_row(row))
    return a, home, away, odds


def _barra(pct, col):
    return (f'<div style="height:8px;border-radius:99px;background:{COL["ink"]};overflow:hidden;">'
            f'<div style="height:100%;width:{max(0,min(100,pct)):.0f}%;background:{col};"></div></div>')


def _col_conf(v):  # v in 0..100
    return COL["win"] if v >= 65 else (COL["draw"] if v >= 45 else COL["loss"])


def _peso_partita(match_comp, comp_df, target_livello, target_categoria, target_comp_key):
    """Peso di una partita storica in base al contesto, relativo alla partita da
    pronosticare: amichevoli e categorie inferiori pesano meno; stessa competizione
    pesa di più. Ritorna (peso, motivo|None)."""
    peso, motivo = 1.0, None
    cat = categoria_di(match_comp, comp_df)          # Amichevole/Coppa.../Campionato/None
    liv = _livello_di(match_comp, comp_df)           # int o None
    ckey = _key(match_comp) if match_comp else None

    # amichevole = rumore
    if cat == "Amichevole":
        return 0.35, "amichevole"
    # stessa identica competizione del match target -> peso pieno/rinforzato
    if target_comp_key and ckey == target_comp_key:
        return 1.25, "stessa competizione"
    # categoria inferiore rispetto al target (es. LL2 quando il target è in Liga)
    if liv is not None and target_livello is not None and liv > target_livello:
        diff = liv - target_livello
        peso = max(0.35, 1.0 - 0.35 * diff)
        motivo = f"categoria inferiore (liv. {liv} vs {target_livello})"
        return peso, motivo
    # coppa (avversari di livello misto) -> leggermente ridotta
    if cat in ("Coppa nazionale", "Coppa internazionale"):
        return 0.85, "coppa (livello avversari misto)"
    return peso, motivo


def _livello_di(codice, comp_df):
    if comp_df is None or comp_df.empty or "livello" not in comp_df.columns:
        return None
    k = _key(codice)
    for _, c in comp_df.iterrows():
        if k in _chiavi_competizione(c):
            liv = c.get("livello")
            try:
                return int(liv) if not pd.isna(liv) else None
            except (TypeError, ValueError):
                return None
    return None


def _giorni_tra(d1, d2):
    """Giorni tra due date (stringhe ISO o simili). None se non calcolabile."""
    try:
        a = pd.to_datetime(d1); b = pd.to_datetime(d2)
        return abs((a - b).days)
    except Exception:
        return None


def _partite_squadra_evidenze(df, team, prima_di=None, escludi_id=None,
                              comp_df=None, target_livello=None, target_categoria=None,
                              target_comp_key=None, recency_decay=None):
    """Estrae le partite giocate di una squadra nel formato del nuovo motore
    (gf/gs dal suo punto di vista, casa True/False, peso), dalla più recente. Filtra per
    data < prima_di e ESCLUDE la fixture stessa. Assegna a ogni partita un PESO in base
    al contesto (competizione) e, se recency_decay è dato, anche al DECADIMENTO temporale
    rispetto alla data della fixture: peso_recency = exp(-giorni/decay)."""
    if df.empty:
        return []
    d = df[(df["squadra_casa"] == team) | (df["squadra_trasferta"] == team)]
    d = d[d["gol_casa"].notna() & d["gol_trasferta"].notna()]
    if escludi_id is not None and "id" in d.columns:
        d = d[d["id"] != escludi_id]
    if prima_di is not None and "data" in d.columns:
        try:
            d = d[d["data"] < prima_di]
        except Exception:
            pass
    if "data" in d.columns:
        d = d.sort_values("data", ascending=False)
    import math as _math
    out = []
    for _, m in d.iterrows():
        peso, motivo = 1.0, None
        if comp_df is not None:
            peso, motivo = _peso_partita(m.get("competizione"), comp_df,
                                         target_livello, target_categoria, target_comp_key)
        # decadimento temporale (opzionale): pesa di più le partite recenti
        if recency_decay and prima_di is not None and "data" in m:
            gg = _giorni_tra(m.get("data"), prima_di)
            if gg is not None:
                peso = peso * _math.exp(-gg / float(recency_decay))
        rec = {"peso": peso, "motivo_peso": motivo, "data": m.get("data"),
               "livello": _livello_di(m.get("competizione"), comp_df) if comp_df is not None else None}
        if m["squadra_casa"] == team:
            rec.update({"gf": int(m["gol_casa"]), "gs": int(m["gol_trasferta"]), "casa": True})
        else:
            rec.update({"gf": int(m["gol_trasferta"]), "gs": int(m["gol_casa"]), "casa": False})
        out.append(rec)
    return out


def _handicap_livello(partite, target_livello):
    """Handicap di forza per una squadra che ha giocato in una categoria diversa dal
    target. Se il livello medio giocato è INFERIORE (numero più alto) al target, la
    squadra parte svantaggiata (es. sale dalla seconda divisione)."""
    if target_livello is None:
        return 1.0
    num = den = 0.0
    for p in partite:
        liv = p.get("livello")
        if liv is None:
            continue
        w = p.get("peso", 1.0)
        num += liv * w
        den += w
    if den == 0:
        return 1.0
    avg = num / den
    if avg <= target_livello + 0.15:      # gioca al livello target o superiore
        return 1.0
    diff = avg - target_livello
    return max(0.55, 1.0 - 0.28 * diff)   # ogni livello sotto -> ~28% di handicap


def analisi_ragionata(df, home, away, data_partita=None, odds=None, variazioni=None, escludi_id=None, competizione=None, recency_decay=None):
    """Ponte verso il nuovo motore: evidenze -> signal score -> racconto.
    Ritorna il dict del racconto, oppure None se manca lo storico."""
    comp_df = carica_competizioni()
    t_liv = _livello_di(competizione, comp_df) if competizione else None
    t_cat = categoria_di(competizione, comp_df) if competizione else None
    t_key = _key(competizione) if competizione else None
    ph = _partite_squadra_evidenze(df, home, data_partita, escludi_id, comp_df, t_liv, t_cat, t_key, recency_decay)
    pa = _partite_squadra_evidenze(df, away, data_partita, escludi_id, comp_df, t_liv, t_cat, t_key, recency_decay)
    if not ph or not pa:
        return None
    hcap_h = _handicap_livello(ph, t_liv)
    hcap_a = _handicap_livello(pa, t_liv)
    ev = evidenze.costruisci_evidenze(ph, pa, odds=odds, variazioni=variazioni,
                                      hcap_home=hcap_h, hcap_away=hcap_a)
    ev["peso_info"] = {"home": _riepilogo_pesi(ph, hcap_h),
                       "away": _riepilogo_pesi(pa, hcap_a)}
    sig = segnali.calcola_signal(ev)
    stat = statistico.analizza(ph, pa)     # motore statistico (conteggi grezzi)
    return racconto.racconta(home, away, ev, sig, competizione=competizione, statistico=stat)


def _riepilogo_pesi(partite, hcap):
    """Riassume quante partite sono state pesate meno e perché, per la trasparenza."""
    motivi = {}
    for p in partite:
        mo = p.get("motivo_peso")
        if mo:
            # normalizza "categoria inferiore (liv. 2 vs 1)" -> "categoria inferiore"
            base = mo.split(" (")[0]
            motivi[base] = motivi.get(base, 0) + 1
    return {"motivi": motivi, "handicap": round(hcap, 2)}


def render_racconto_st(racc):
    """Rende l'analisi ragionata (nuovo motore) in Streamlit."""
    if not racc:
        st.info("Analisi ragionata non disponibile (storico insufficiente).")
        return
    pron = racc["pronostico"]
    colore = COL["win"] if pron["score"] >= 55 else (COL["draw"] if pron["score"] >= 40 else COL["loss"])
    st.html(
        f'<div style="{FONT}max-width:560px;background:{COL["panel"]};border:1px solid {COL["line"]};'
        f'border-left:4px solid {colore};border-radius:14px;padding:16px;margin-bottom:6px;">'
        f'<div style="color:{COL["lo"]};font-size:11px;text-transform:uppercase;letter-spacing:.15em;">'
        f'Pronostico ragionato</div>'
        f'<div style="color:{COL["hi"]};font-size:20px;font-weight:700;margin-top:6px;">'
        f'{_esc(pron["testo"])}</div></div>')
    for sez in racc["sezioni"]:
        with st.expander(sez["titolo"], expanded=sez["titolo"] in
                         ("💎 Giocate di valore (EV → edge → signal)",
                          "⭐ Top 5 pronostici più sicuri (confidence combinata)",
                          "🔀 Motore vs Statistico (validazione)")):
            for r in sez["righe"]:
                st.markdown(f"- {r}")


def _ultime_partite_testo(df, team, n=15, escludi_id=None):
    """Righe di testo con le ultime partite giocate di una squadra."""
    if df.empty:
        return []
    d = df[((df["squadra_casa"] == team) | (df["squadra_trasferta"] == team))]
    d = d[d["gol_casa"].notna() & d["gol_trasferta"].notna()]
    if escludi_id is not None and "id" in d.columns:
        d = d[d["id"] != escludi_id]
    if "data" in d.columns:
        d = d.sort_values("data", ascending=False)
    out = []
    for _, m in d.head(n).iterrows():
        data = m["data"]
        try:
            data = pd.to_datetime(str(data)).strftime("%d.%m.%y")
        except Exception:
            pass
        gc, gt = int(m["gol_casa"]), int(m["gol_trasferta"])
        comp = _txt(m.get("competizione"))
        comp = f"[{comp}] " if comp else ""
        out.append(f"{data} {comp}{m['squadra_casa']} {gc}-{gt} {m['squadra_trasferta']}")
    return out


def riepilogo_testo(a, df, home, away, odds, row=None):
    """Blocco di testo (da copiare) con pronostico + quote + ultime partite."""
    p = a["prob"]
    L = []
    data = ""
    if row is not None and "data" in row:
        try:
            data = pd.to_datetime(str(row["data"])).strftime("%d.%m.%Y")
        except Exception:
            data = str(row.get("data") or "")
    L.append(f"{home.upper()} - {away.upper()}" + (f"  ({data})" if data else ""))
    comp = _txt(row.get("competizione")) if row is not None else ""
    if comp:
        L.append(f"Competizione: {comp}")
    L.append("")
    L.append(f"PRONOSTICO: {a['best']['mercato']}  (confidence {a['best']['confidence']:.0f}/100)")
    if a["best"].get("alert"):
        L.append(f"  ⚠️ alert quota {a['best']['alert']} — stat {a['best']['prob']*100:.0f}% "
                 f"vs quota {a['best']['market_prob']*100:.0f}%")
    L.append(f"1: {p['1']*100:.0f}%   X: {p['X']*100:.0f}%   2: {p['2']*100:.0f}%")
    L.append(f"Over 2.5: {a['over_prob']*100:.0f}%   Goal: {a['btts_prob']*100:.0f}%")
    L.append(f"Gol attesi: {home} {p['lambda_home']:.2f} - {away} {p['lambda_away']:.2f}")
    rpe = a.get("risultati_per_esito", {})
    if a.get("prob", {}).get("risultati"):
        top = ", ".join(f"{r['risultato']} ({r['p']*100:.0f}%)" for r in p["risultati"][:4])
        L.append(f"Risultati probabili: {top}")
    if odds:
        qs = []
        for k, et in [("1", "1"), ("X", "X"), ("2", "2"), ("over25", "Over"),
                      ("under25", "Under"), ("goal", "Goal"), ("nogoal", "NoGoal")]:
            if odds.get(k):
                qs.append(f"{et} {odds[k]}")
        if qs:
            L.append("Quote: " + " / ".join(qs))
    if a.get("alerts"):
        L.append("")
        L.append("ALERT QUOTA:")
        for al in a["alerts"]:
            L.append(f"  - {al['mercato']}: {al['livello']} (stat {al['prob']*100:.0f}% "
                     f"vs quota {al['market_prob']*100:.0f}%)")
    for team in (home, away):
        righe = _ultime_partite_testo(df, team,
                                      escludi_id=(row.get("id") if row is not None else None))
        if righe:
            L.append("")
            L.append(f"ULTIME PARTITE {team.upper()}:")
            L.extend(righe)
    return "\n".join(L)


def pagina_analisi(user):
    st.header("🔮 Analisi & Pronostico")
    st.caption("Motore multi-fattore: forma pesata (5/10/15 + recency), split casa/trasferta, "
               "qualità avversari (Elo), Expected Goals (Poisson), score per mercato con "
               "conflitto dei segnali. Le quote sono usate solo come consenso, non per decidere.")

    if not supabase_pronto():
        st.warning("Supabase non configurato.")
        return

    df = carica_partite()
    if df.empty:
        st.info("Database vuoto.")
        return

    pend = df[df["gol_casa"].isna()] if "gol_casa" in df else df.iloc[0:0]
    if not pend.empty:
        opz = {f'{r["squadra_casa"]} - {r["squadra_trasferta"]}  ({r["data"]})': r
               for _, r in pend.iterrows()}
        scelta = st.selectbox("Partita da analizzare (in attesa di risultato)", list(opz.keys()))
        row = opz[scelta]
        home, away = row["squadra_casa"], row["squadra_trasferta"]
        odds = _eff_odds(row)
        data_partita = row["data"]
    else:
        st.caption("Nessuna partita in attesa: scegli due squadre dal database.")
        squadre = sorted(set(df["squadra_casa"]) | set(df["squadra_trasferta"]))
        c = st.columns(2)
        home = c[0].selectbox("Casa", squadre, key="an_home")
        away = c[1].selectbox("Trasferta", squadre, key="an_away", index=min(1, len(squadre) - 1))
        odds, data_partita = {}, None

    if home == away:
        st.warning("Seleziona due squadre diverse.")
        return

    # --- pesi configurabili (nuovo motore) ---
    with st.expander("⚙️ Pesi del modello"):
        usa_rec = st.checkbox("Decadimento temporale (pesa di più le partite recenti)",
                              value=False, key="an_usa_recency")
        decay_gg = st.slider("Decadimento (giorni)", 20, 180, 60, step=10,
                             disabled=not usa_rec, key="an_recency",
                             help="Più basso = più peso alle partite recentissime. "
                                  "peso = exp(-giorni/decadimento). Testa l'effetto nel Backtest.")
        an_recency = decay_gg if usa_rec else None
        st.caption("Il tipo di competizione pesa già le partite (campionato conta più di "
                   "un'amichevole). Il decadimento aggiunge il peso della recency.")
        st.caption("Le quote NON mediano la probabilità: servono solo a segnalare le "
                   "discrepanze (alert) e a modulare la confidence.")
    # il vecchio motore usa ancora alcuni campi: config con i default consolidati
    config = _config_default()

    calibratori = carica_calibrazione()
    # livelli di lega dalle competizioni (per Elo consapevole della divisione)
    comp_df = carica_competizioni()
    livelli = {}
    if not comp_df.empty and "livello" in comp_df:
        for _, cc in comp_df.iterrows():
            liv = cc.get("livello")
            if pd.isna(liv) if hasattr(pd, "isna") else (liv is None):
                continue
            if liv is None or (isinstance(liv, float) and liv != liv):
                continue
            for kk in _chiavi_competizione(cc):
                livelli[kk] = int(liv)
    rose = None
    tipo_target = None
    variazioni = None
    if not pend.empty:
        rose = (row.get("val_casa"), row.get("val_trasferta"))
        tp = row.get("tipo_partita")
        tipo_target = tp if (tp and str(tp) not in ("ND", "Non assegnata", "None")) else None
        variazioni = {}
        for kk in ("1", "x", "2", "over", "under", "goal", "nogoal"):
            v = row.get(f"variazione_quota_{kk}")
            if v is not None and not (pd.isna(v) if hasattr(pd, "isna") else False):
                mapk = {"x": "X", "over": "over25", "under": "under25"}.get(kk, kk)
                variazioni[mapk] = float(v)
    a = analisi.analizza_partita(home, away, df, odds=odds,
                                 data_partita=data_partita, config=config,
                                 calibratori=calibratori, rose=rose,
                                 tipo_partita_target=tipo_target, livelli=livelli,
                                 variazioni=variazioni)

    if a.get("errore"):
        st.warning(f"Dati storici insufficienti ({a.get('n_home',0)} / {a.get('n_away',0)} partite).")
        return

    # === ANALISI & PRONOSTICO (motore unico: frequenze + Poisson/Elo fusi) ===
    st.subheader("🎯 Analisi & Pronostico")
    # campo per il pronostico personale (combo con '+', es. "1X + Under 2.5")
    pron_cristiano = ""
    if not pend.empty:
        _pid_c = str(row.get("id"))
        _pron_esist = ""
        try:
            _pr = carica_pronostici()
            if not _pr.empty and "pron_cristiano" in _pr.columns:
                _match = _pr[_pr["partita_id"].astype(str) == _pid_c]
                if not _match.empty:
                    _pron_esist = _txt(_match.iloc[0].get("pron_cristiano"))
        except Exception:
            pass
        pron_cristiano = st.text_input(
            "✍️ Pronostico Cristiano", value=_pron_esist, key=f"pron_cri_{_pid_c}",
            placeholder="Es. 1X + Under 2.5  ·  Over 0.5 casa  ·  Goal + Over 2.5",
            help="Il tuo pronostico. Combo con '+' (tutte da vincere). Mercati: 1/X/2, 1X/X2/12, "
                 "Over/Under 1.5/2.5/3.5, Goal/No Goal, e Over/Under di squadra (es. 'Over 1.5 casa').")
        if st.button("💾 Salva Pronostico Cristiano", key=f"salva_pron_cri_{_pid_c}"):
            try:
                cli = get_client()
                val = pron_cristiano.strip() if pron_cristiano else None
                # assicura che esista la riga pronostico, poi aggiorna la colonna
                ex = cli.table("pronostici").select("id").eq("partita_id", _pid_c).execute()
                if ex.data:
                    cli.table("pronostici").update({"pron_cristiano": val}).eq(
                        "partita_id", _pid_c).execute()
                else:
                    cli.table("pronostici").insert({
                        "partita_id": _pid_c, "data": str(row.get("data")),
                        "squadra_casa": home, "squadra_trasferta": away,
                        "pron_cristiano": val}).execute()
                st.cache_data.clear()
                st.success("Pronostico Cristiano salvato.")
            except Exception as e:
                st.warning(f"Salvataggio non riuscito: {e}")
    st.caption("Motore unico: probabilità a frequenze e Poisson/Elo fuse in un'unica stima. "
               "Signal = robustezza statistica; value (EV) = convenienza vs quota, separati.")
    comp_target = (_label_da_comp(row.get("competizione"), carica_competizioni())
                   if not pend.empty else None)
    racc = analisi_ragionata(df, home, away, data_partita=data_partita,
                             odds=odds, variazioni=variazioni,
                             escludi_id=(row.get("id") if not pend.empty else None),
                             competizione=comp_target, recency_decay=an_recency)
    render_racconto_st(racc)

    # (motore fuso: la visualizzazione è la sola "Analisi ragionata" sopra;
    #  qui teniamo solo i valori del vecchio motore per il salvataggio storico e la calibrazione)
    p = a["prob"]
    best = a["best"]

    # --- riepilogo (una volta) usato sia per la copia sia per lo storico ---
    testo = riepilogo_testo(a, df, home, away, odds,
                            row=row if not pend.empty else None)

    # --- salvataggio AUTOMATICO del pronostico (fixture pre-partita) ---
    if not pend.empty and "id" in row:
        sig = f'{row["id"]}|{best["mercato"]}|{best["confidence"]:.0f}'
        salvati = st.session_state.setdefault("auto_pron", set())
        if sig not in salvati:
            try:
                merc_rag = None
                score_rag = None
                if racc and racc.get("pronostico"):
                    merc_rag = racc["pronostico"].get("mercato")
                    score_rag = racc["pronostico"].get("score")
                # tre motori: motore (probabilistico), statistico, fusione
                m_mot = merc_rag
                c_mot = int(score_rag) if score_rag is not None else None
                m_stat = c_stat = None
                if racc and racc.get("pronostico_statistico"):
                    ps = racc["pronostico_statistico"]
                    m_stat = ps.get("pronostico")
                    c_stat = {"alta": 90, "media": 70, "bassa": 50}.get(ps.get("confidence"))
                m_fus = c_fus = None
                if racc and racc.get("fusione_media"):
                    fm = racc["fusione_media"]
                    m_fus = fm.get("mercato")
                    c_fus = fm.get("confidence")
                m_sstat = c_sstat = None
                if racc and racc.get("solo_statistico"):
                    ss = racc["solo_statistico"]
                    m_sstat = ss.get("mercato")
                    c_sstat = ss.get("confidence")
                upsert_pronostico({
                    "partita_id": str(row["id"]),
                    "data": str(row["data"]) if "data" in row else None,
                    "squadra_casa": home, "squadra_trasferta": away,
                    "mercato": best["mercato"], "prob": float(best["prob"]),
                    "confidence": float(best["confidence"]),
                    "quota": float(best["quota"]) if best.get("quota") else None,
                    "prob_over25": float(a["over_prob"]), "prob_goal": float(a["btts_prob"]),
                    "prob_1": float(p["1"]), "prob_x": float(p["X"]), "prob_2": float(p["2"]),
                    "riepilogo": testo,
                    "scheda_json": json.dumps(_snapshot_analisi(a, odds)),
                    "mercato_ragionato": merc_rag,
                    "score_ragionato": int(score_rag) if score_rag is not None else None,
                    "merc_motore": m_mot, "conf_motore": c_mot,
                    "merc_statistico": m_stat, "conf_statistico": c_stat,
                    "merc_fusione": m_fus, "conf_fusione": c_fus,
                    "merc_solo_stat": m_sstat, "conf_solo_stat": c_sstat,
                    "merc_ev": (racc.get("miglior_ev") or {}).get("mercato") if racc else None,
                    "val_ev": (racc.get("miglior_ev") or {}).get("ev") if racc else None,
                    "pron_cristiano": pron_cristiano.strip() if pron_cristiano else None,
                })
                salvati.add(sig)
                st.caption("💾 Pronostico salvato automaticamente in 📈 Storico pronostici.")
            except Exception as e:
                st.warning(f"Salvataggio automatico non riuscito: {e}")
        else:
            st.caption("💾 Pronostico salvato automaticamente in 📈 Storico pronostici.")

    # --- riepilogo da copiare (per studio / note) ---
    with st.expander("📋 Riepilogo da copiare (note)"):
        st.caption("Blocco di testo con pronostico, quote e ultime partite: "
                   "usa l'icona 📋 in alto a destra del riquadro per copiarlo tutto.")
        st.code(testo, language=None)

    # --- CALIBRAZIONE ---
    st.divider()
    st.subheader("📐 Calibrazione del modello")
    st.caption("Costruisce un dataset dalle partite già giocate (usando solo i dati precedenti a "
               "ciascuna), verifica se le probabilità sono realistiche e le corregge con una "
               "regressione isotonica. Una confidence del 80% dovrebbe valere ~80% di successi.")

    cal_att = carica_calibrazione()
    if cal_att:
        st.caption("Calibratori attivi: " + ", ".join(cal_att.keys()))

    cca = st.columns(3)
    if cca[0].button("🔄 Aggiorna risultati pronostici"):
        n = completa_risultati_pronostici()
        st.success(f"Risultati abbinati a {n} pronostici salvati." if n
                   else "Nessun pronostico da aggiornare.")
    if cca[1].button("🎛️ Ottimizza parametri"):
        with st.spinner("Ricerca dei parametri migliori sul database…"):
            opt = analisi.ottimizza_parametri(df)
        if not opt or opt.get("errore"):
            need = opt.get("min", 80) if opt else 80
            st.info(f"Servono almeno ~{need} partite giocate per ottimizzare senza overfitting.")
        else:
            mig = opt["migliore"]
            st.success(f"Migliori parametri: rho Dixon-Coles **{mig['rho']}**, "
                       f"peso mercato **{mig['peso_mercato']}** (Brier {mig['brier']:.3f}, "
                       f"su {mig['n']} partite).")
            st.caption("Puoi impostarli come default nel motore quando sei soddisfatto.")
    if cca[2].button("📊 Come ragiona il bookmaker"):
        st.session_state["odds_study"] = analisi_quote_bookmaker()
    study = st.session_state.get("odds_study")
    if study:
        st.markdown("**Quote per fascia — quanto spesso si verificano davvero:**")
        st.dataframe(pd.DataFrame([{
            "Fascia quota": s["fascia"], "N": s["n"], "Vinte": s["vinte"],
            "Tasso reale": f'{s["tasso"]*100:.0f}%',
            "Implicita nella quota": f'{s["implicita"]*100:.0f}%'} for s in study]),
            use_container_width=True, hide_index=True)
        st.caption("Se il 'tasso reale' supera l'implicita, in quella fascia il bookmaker "
                   "tende a sottostimare; il contrario se è più basso.")

    st.markdown("**Validazione della logica quote↔statistiche** (tesi confidence):")
    st.caption("Verifica sui pronostici salvati se davvero conviene fidarsi quando il mercato "
               "conferma le statistiche e diffidare quando diverge molto.")
    if st.button("🔬 Analizza discrepanze quote/statistiche"):
        st.session_state["discrep"] = analisi_discrepanze()
    disc = st.session_state.get("discrep")
    if disc:
        if disc.get("insufficiente"):
            st.info(f"Servono più pronostici salvati con risultato (ora {disc['n']}). "
                    "Salva i pronostici e aggiorna i risultati per popolare l'analisi.")
        else:
            st.caption(f"{disc['n']} pronostici analizzati. Tasso = quante volte la giocata è "
                       "riuscita davvero.")
            st.markdown("*Per scarto tra quota implicita e statistica:*")
            st.dataframe(pd.DataFrame([{
                "Situazione": r["gruppo"], "N": r["n"],
                "Tasso reale": f'{r["tasso"]*100:.0f}%',
                "Stat media": f'{r["stat_media"]*100:.0f}%',
                "Quota media": f'{r["implicita_media"]*100:.0f}%'} for r in disc["per_scarto"]]),
                use_container_width=True, hide_index=True)
            st.markdown("*Per livello di alert:*")
            st.dataframe(pd.DataFrame([{
                "Alert": r["gruppo"], "N": r["n"], "Tasso reale": f'{r["tasso"]*100:.0f}%'}
                for r in disc["per_alert"]]),
                use_container_width=True, hide_index=True)
            st.caption("Lettura: se 'quota più alta (conferma)' ha un tasso alto e 'quota molto "
                       "più alta (discrepanza)' un tasso basso, la tesi regge. Se gli alto/medio "
                       "hanno tassi bassi, gli alert stanno segnalando bene le giocate rischiose.")

    if st.button("Calcola calibrazione (dal backtest)"):
        with st.spinner("Backtest e fit in corso… (può richiedere qualche secondo)"):
            bt = analisi.backtest(df, config=config)
        if not bt or not bt.get("pairs_over"):
            st.info("Servono più partite giocate con storico sufficiente.")
        else:
            st.session_state["cal_bt"] = {
                "n": bt["n"], "pairs_over": bt["pairs_over"], "pairs_goal": bt["pairs_goal"],
                "acc_1x2": bt["acc_1x2"], "acc_over": bt["acc_over"], "acc_goal": bt["acc_goal"],
            }

    bt = st.session_state.get("cal_bt")
    if bt:
        st.markdown(f"**{bt['n']}** partite testate · accuratezza 1X2 {bt['acc_1x2']*100:.0f}% · "
                    f"Over {bt['acc_over']*100:.0f}% · Goal {bt['acc_goal']*100:.0f}%")
        for mercato, etichetta, pairs in [("over25", "Over 2.5", bt["pairs_over"]),
                                          ("goal", "Goal", bt["pairs_goal"])]:
            iso = analisi.fit_isotonic(pairs)
            b_pre = analisi.brier(pairs)
            b_post = analisi.brier_calibrato(pairs, iso) if iso else None
            st.markdown(f"**{etichetta}** — Brier grezzo {b_pre:.3f}"
                        + (f" → calibrato **{b_post:.3f}**" if b_post is not None else
                           " (dati insufficienti per calibrare)"))
            rel = analisi.reliability(pairs, 10)
            if rel:
                rel_df = pd.DataFrame([{
                    "Bin": r["bin"], "Previsto": round(r["previsto"], 2),
                    "Reale": round(r["reale"], 2),
                    "Calibrato": round(analisi.applica_iso(r["previsto"], iso), 2) if iso else None,
                    "n": r["n"]} for r in rel])
                st.dataframe(rel_df, use_container_width=True, hide_index=True)
                st.line_chart(rel_df.set_index("Bin")[["Previsto", "Reale"]])
            if iso and b_post is not None and b_post <= b_pre:
                if st.button(f"✅ Attiva calibrazione {etichetta}", key=f"savecal_{mercato}"):
                    try:
                        salva_calibrazione(mercato, iso, len(pairs), b_pre, b_post)
                        st.success(f"Calibrazione {etichetta} attivata. Verrà applicata alle analisi.")
                        st.cache_data.clear()
                    except Exception as e:
                        st.error(f"Errore: {e}")
            elif iso:
                st.caption(f"La calibrazione non migliora il Brier: meglio non attivarla per {etichetta}.")


# =============================================================================
#  MAIN
# =============================================================================
def backfill_tre_motori(pron, df_tutte, comp_df, progress=None, forza=False, limite=None,
                        solo_da_aggiornare=False):
    """Ricalcola e SALVA motore + fusione per i pronostici.
    - forza=True: ricalcola tutti; altrimenti solo quelli senza motore/fusione.
    - limite: elabora al massimo N pronostici (per lavorare A LOTTI ed evitare timeout).
    - solo_da_aggiornare: salta quelli già aggiornati con la logica nuova (che hanno merc_ev
      valorizzato o dove il motore non è più '12'), utile per riprendere.
    Ritorna (aggiornati, rimasti_da_fare)."""
    cli = get_client()
    if not cli:
        return 0, 0
    n = 0
    da_fare = 0
    righe = list(pron.iterrows())
    for i, (_, r) in enumerate(righe):
        if progress and i % 3 == 0:
            progress.progress(min(i / max(1, len(righe)), 1.0),
                              text=f"Pronostico {i+1}/{len(righe)}…")
        if not forza and _txt(r.get("merc_motore")) and _txt(r.get("merc_fusione")):
            continue    # già completo (motore + fusione)
        # modalità ripresa: salta quelli il cui motore NON è più '12' (già ricalcolati)
        if solo_da_aggiornare and _txt(r.get("merc_motore")) not in ("12", ""):
            continue
        if limite is not None and n >= limite:
            da_fare += 1
            continue
        pid = _txt(r.get("partita_id"))
        if not pid:
            continue
        try:
            comp = None
            odds = None
            if not df_tutte.empty and "id" in df_tutte.columns:
                mrow = df_tutte[df_tutte["id"].astype(str) == pid]
                if not mrow.empty:
                    comp = _label_da_comp(mrow.iloc[0].get("competizione"), comp_df)
                    odds = _eff_odds(mrow.iloc[0])   # quote salvate -> abilita l'EV
            racc = analisi_ragionata(df_tutte, r.get("squadra_casa"), r.get("squadra_trasferta"),
                                     data_partita=r.get("data"), escludi_id=pid,
                                     competizione=comp, odds=odds)
            if not racc:
                continue
            pm = racc.get("pronostico") or {}
            ps = racc.get("pronostico_statistico") or {}
            fm = racc.get("fusione_media") or {}
            ss = racc.get("solo_statistico") or {}
            upd = {
                "merc_motore": pm.get("mercato"),
                "conf_motore": int(pm["score"]) if pm.get("score") is not None else None,
                "merc_statistico": ps.get("pronostico"),
                "conf_statistico": {"alta": 90, "media": 70, "bassa": 50}.get(ps.get("confidence")),
                "merc_fusione": fm.get("mercato"),
                "conf_fusione": fm.get("confidence"),
                "merc_solo_stat": ss.get("mercato"),
                "conf_solo_stat": ss.get("confidence"),
                "merc_ev": (racc.get("miglior_ev") or {}).get("mercato"),
                "val_ev": (racc.get("miglior_ev") or {}).get("ev"),
            }
            cli.table("pronostici").update(upd).eq("id", r["id"]).execute()
            n += 1
        except Exception:
            continue
    st.cache_data.clear()
    return n, da_fare


def _config_default():
    """Pesi di default del modello (gli stessi valori iniziali degli slider in Analisi)."""
    return {
        "recency_decay": 45, "home_adv_goals": 1.08, "over": {"h2h": 0.10},
        "pesi_competizione": {
            "Campionato": 1.00, "Playoff": 1.00,
            "Coppa nazionale": 0.85, "Coppa internazionale": 0.85,
            "Coppa/torneo secondario": 0.75, "Altro": 0.75, "Amichevole": 0.35,
        },
        "blending_mercato": False, "peso_mercato": 0.0,
    }


def _livelli_da_comp(comp_df):
    livelli = {}
    if comp_df is not None and not comp_df.empty and "livello" in comp_df:
        for _, cc in comp_df.iterrows():
            liv = cc.get("livello")
            if liv is None or (isinstance(liv, float) and liv != liv):
                continue
            for kk in _chiavi_competizione(cc):
                livelli[kk] = int(liv)
    return livelli


def _record_pronostico_da_fixture(row, df, comp_df, calibratori, livelli, config):
    """Calcola il record completo del pronostico per una fixture (senza salvarlo).
    Replica esattamente ciò che fa il salvataggio automatico nella pagina Analisi.
    Ritorna il dict pronto per upsert_pronostico, o None se dati insufficienti."""
    home, away = row["squadra_casa"], row["squadra_trasferta"]
    odds = _eff_odds(row)
    data_partita = row.get("data")
    rose = (row.get("val_casa"), row.get("val_trasferta"))
    tp = row.get("tipo_partita")
    tipo_target = tp if (tp and str(tp) not in ("ND", "Non assegnata", "None")) else None
    variazioni = {}
    for kk in ("1", "x", "2", "over", "under", "goal", "nogoal"):
        v = row.get(f"variazione_quota_{kk}")
        if v is not None and not (pd.isna(v) if hasattr(pd, "isna") else False):
            mapk = {"x": "X", "over": "over25", "under": "under25"}.get(kk, kk)
            variazioni[mapk] = float(v)

    a = analisi.analizza_partita(home, away, df, odds=odds, data_partita=data_partita,
                                 config=config, calibratori=calibratori, rose=rose,
                                 tipo_partita_target=tipo_target, livelli=livelli,
                                 variazioni=variazioni)
    if a.get("errore"):
        return None

    comp_target = _label_da_comp(row.get("competizione"), comp_df)
    racc = analisi_ragionata(df, home, away, data_partita=data_partita, odds=odds,
                             variazioni=variazioni, escludi_id=row.get("id"),
                             competizione=comp_target)
    p = a["prob"]
    best = a["best"]
    testo = riepilogo_testo(a, df, home, away, odds, row=row)

    merc_rag = score_rag = None
    if racc and racc.get("pronostico"):
        merc_rag = racc["pronostico"].get("mercato")
        score_rag = racc["pronostico"].get("score")
    m_stat = c_stat = None
    if racc and racc.get("pronostico_statistico"):
        ps = racc["pronostico_statistico"]
        m_stat = ps.get("pronostico")
        c_stat = {"alta": 90, "media": 70, "bassa": 50}.get(ps.get("confidence"))
    m_fus = c_fus = None
    if racc and racc.get("fusione_media"):
        m_fus = racc["fusione_media"].get("mercato")
        c_fus = racc["fusione_media"].get("confidence")
    m_sstat = c_sstat = None
    if racc and racc.get("solo_statistico"):
        m_sstat = racc["solo_statistico"].get("mercato")
        c_sstat = racc["solo_statistico"].get("confidence")

    return {
        "partita_id": str(row["id"]),
        "data": str(data_partita) if data_partita is not None else None,
        "squadra_casa": home, "squadra_trasferta": away,
        "mercato": best["mercato"], "prob": float(best["prob"]),
        "confidence": float(best["confidence"]),
        "quota": float(best["quota"]) if best.get("quota") else None,
        "prob_over25": float(a["over_prob"]), "prob_goal": float(a["btts_prob"]),
        "prob_1": float(p["1"]), "prob_x": float(p["X"]), "prob_2": float(p["2"]),
        "riepilogo": testo,
        "scheda_json": json.dumps(_snapshot_analisi(a, odds)),
        "mercato_ragionato": merc_rag,
        "score_ragionato": int(score_rag) if score_rag is not None else None,
        "merc_motore": merc_rag,
        "conf_motore": int(score_rag) if score_rag is not None else None,
        "merc_statistico": m_stat, "conf_statistico": c_stat,
        "merc_fusione": m_fus, "conf_fusione": c_fus,
        "merc_solo_stat": m_sstat, "conf_solo_stat": c_sstat,
    }


def genera_pronostici_mancanti(df, comp_df, pron, progress=None):
    """Genera e salva il pronostico per le fixture (is_target) senza pronostico salvato.
    Ritorna (creati, saltati_gia, saltati_dati, primo_errore) per piena trasparenza."""
    if "is_target" not in df.columns:
        return 0, 0, 0, "colonna is_target assente"
    fixtures = df[df["is_target"] == True]
    gia = set()
    if pron is not None and not pron.empty and "partita_id" in pron.columns:
        gia = set(pron["partita_id"].astype(str))
    calibratori = carica_calibrazione()
    livelli = _livelli_da_comp(comp_df)
    config = _config_default()
    creati = saltati_gia = saltati_dati = 0
    primo_errore = None
    righe = list(fixtures.iterrows())
    for i, (_, row) in enumerate(righe):
        if progress and i % 3 == 0:
            progress.progress(i / max(1, len(righe)), text=f"Partita {i+1}/{len(righe)}…")
        pid = str(row.get("id"))
        if pid in gia:
            saltati_gia += 1
            continue
        try:
            rec = _record_pronostico_da_fixture(row, df, comp_df, calibratori, livelli, config)
            if rec is None:
                saltati_dati += 1
                continue
            upsert_pronostico(rec)
            creati += 1
        except Exception as e:
            saltati_dati += 1
            if primo_errore is None:
                primo_errore = str(e)
    return creati, saltati_gia, saltati_dati, primo_errore


def pagina_storico_pronostici(user):
    st.header("📈 Storico pronostici")
    st.caption("I pronostici salvati prima della partita, confrontati col risultato reale. "
               "Solo le partite per cui hai salvato il pronostico dall'Analisi.")

    if not supabase_pronto():
        st.warning("Supabase non configurato.")
        return

    cc = st.columns(3)
    if cc[0].button("🔄 Ricarica"):
        st.cache_data.clear()
    if cc[1].button("✅ Aggiorna risultati"):
        n = completa_risultati_pronostici()
        st.success(f"Risultati abbinati a {n} pronostici." if n else "Nessun nuovo risultato.")
    if cc[2].button("⚙️ Popola motore+fusione (mancanti)"):
        _pr = st.progress(0.0, text="Ricalcolo in corso…")
        _pron = carica_pronostici()
        _n, _ = backfill_tre_motori(_pron, carica_partite(), carica_competizioni(), _pr)
        _pr.progress(1.0, text="Completato.")
        st.success(f"Popolati {_n} pronostici. Ora restano salvati.")
        st.rerun()
    if st.button("🔄 Ricalcola fusione per TUTTI (a lotti di 40)"):
        _pron = carica_pronostici().reset_index(drop=True)
        _cur = st.session_state.get("_ricalc_cursore", 0)
        _lotto = _pron.iloc[_cur:_cur + 40]
        _pr = st.progress(0.0, text=f"Ricalcolo lotto (da {_cur+1})…")
        _n, _ = backfill_tre_motori(_lotto, carica_partite(), carica_competizioni(),
                                    _pr, forza=True)
        _pr.progress(1.0, text="Lotto completato.")
        _nuovo_cur = _cur + 40
        if _nuovo_cur < len(_pron):
            st.session_state["_ricalc_cursore"] = _nuovo_cur
            st.warning(f"Ricalcolati {_n} in questo lotto ({_nuovo_cur}/{len(_pron)}). "
                       "Premi di nuovo per continuare col prossimo lotto.")
        else:
            st.session_state["_ricalc_cursore"] = 0
            st.success(f"Completato! Tutti i {len(_pron)} pronostici ricalcolati.")
        st.rerun()
    _cur_now = st.session_state.get("_ricalc_cursore", 0)
    if _cur_now > 0:
        st.caption(f"↩️ Ricalcolo a metà: {_cur_now} già fatti. Premi il pulsante per il "
                   "prossimo lotto, oppure azzera qui sotto.")
        if st.button("🔁 Azzera cursore ricalcolo"):
            st.session_state["_ricalc_cursore"] = 0
            st.rerun()
    if st.button("✨ Genera pronostici per le partite in attesa"):
        _pr = st.progress(0.0, text="Generazione in corso…")
        _creati, _s_gia, _s_dati, _err = genera_pronostici_mancanti(
            carica_partite(), carica_competizioni(), carica_pronostici(), _pr)
        _pr.progress(1.0, text="Completato.")
        st.success(f"Creati e SALVATI {_creati} nuovi pronostici. "
                   f"Saltati {_s_gia} (già salvati) · {_s_dati} (storico insufficiente).")
        if _err:
            st.error(f"⚠️ Alcuni salvataggi non riusciti (primo errore): {_err}")
        st.cache_data.clear()
        st.rerun()

    pron = carica_pronostici()
    if pron.empty:
        st.info("Nessun pronostico salvato. Vai in 🔮 Analisi & Pronostico e usa "
                "'💾 Salva questo pronostico'.")
        return

    # dataframe partite (serve solo se chiedi il ricalcolo dei vecchi)
    df_tutte = carica_partite()
    comp_df_st = carica_competizioni()
    # mappa partita_id -> codice competizione (una volta, per velocità)
    comp_per_id = {}
    if not df_tutte.empty and "id" in df_tutte.columns:
        for _, mm_ in df_tutte.iterrows():
            comp_per_id[str(mm_.get("id"))] = mm_.get("competizione")
    ricalcola_vecchi = st.checkbox(
        "Ricalcola i pronostici vecchi mancanti (più lento)", value=False,
        help="Se attivo, per i pronostici salvati prima dei tre motori ricostruisce "
             "l'analisi pre-partita. Lascialo spento per un caricamento istantaneo.")

    def _tre_motori_di(r):
        """{motore:(merc,conf), statistico:(merc,conf), fusione:(merc,conf)} dai valori
        SALVATI (istantaneo). Ricalcola solo se richiesto e se mancano."""
        def _get(col_m, col_c):
            m = _txt(r.get(col_m)) if col_m in pron.columns else ""
            c = r.get(col_c) if col_c in pron.columns else None
            c = int(c) if c is not None and not pd.isna(c) else None
            return (m, c)
        mot = _get("merc_motore", "conf_motore")
        sta = _get("merc_statistico", "conf_statistico")
        fus = _get("merc_fusione", "conf_fusione")
        sstat = _get("merc_solo_stat", "conf_solo_stat")
        # fallback leggero: se manca merc_motore usa il mercato_ragionato salvato
        if not mot[0] and "mercato_ragionato" in pron.columns:
            mr = _txt(r.get("mercato_ragionato"))
            sc = r.get("score_ragionato")
            mot = (mr, int(sc) if sc is not None and not pd.isna(sc) else None)
        if mot[0] or not ricalcola_vecchi:
            return {"motore": mot, "statistico": sta, "fusione": fus, "solostat": sstat}
        # ricalcolo al volo SOLO se richiesto esplicitamente
        try:
            pid = _txt(r.get("partita_id"))
            comp = None
            if pid and not df_tutte.empty and "id" in df_tutte.columns:
                mrow = df_tutte[df_tutte["id"].astype(str) == pid]
                if not mrow.empty:
                    comp = _label_da_comp(mrow.iloc[0].get("competizione"), carica_competizioni())
            racc_r = analisi_ragionata(df_tutte, r.get("squadra_casa"), r.get("squadra_trasferta"),
                                       data_partita=r.get("data"), escludi_id=(pid or None),
                                       competizione=comp)
            if racc_r:
                pm = racc_r.get("pronostico") or {}
                m_mot = (pm.get("mercato") or "", pm.get("score"))
                ps = racc_r.get("pronostico_statistico") or {}
                m_sta = (ps.get("pronostico") or "",
                         {"alta": 90, "media": 70, "bassa": 50}.get(ps.get("confidence")))
                fmm = racc_r.get("fusione_media") or {}
                m_fus = (fmm.get("mercato") or "", fmm.get("confidence"))
                ssm = racc_r.get("solo_statistico") or {}
                m_ss = (ssm.get("mercato") or "", ssm.get("confidence"))
                return {"motore": m_mot, "statistico": m_sta, "fusione": m_fus, "solostat": m_ss}
        except Exception:
            pass
        return {"motore": mot, "statistico": sta, "fusione": fus, "solostat": sstat}

    def _norm_merc(m):
        """Normalizza i nomi lunghi dello statistico ai mercati verificabili."""
        if not m:
            return ""
        base = m.split(" (")[0].replace(" totali", "")
        base = base.replace("Over 2.5 gol squadra di casa", "").strip()
        return base

    righe = []
    conteggi = {"motore": [0, 0], "fusione": [0, 0], "solostat": [0, 0],
                "ev": [0, 0], "cristiano": [0, 0]}
    aperti = 0
    for _, r in pron.iterrows():
        gc, gt = r.get("gol_casa"), r.get("gol_trasferta")
        tre = _tre_motori_di(r)
        pron_cri = _txt(r.get("pron_cristiano")) if "pron_cristiano" in pron.columns else ""
        merc_ev = _txt(r.get("merc_ev")) if "merc_ev" in pron.columns else ""
        val_ev = r.get("val_ev") if "val_ev" in pron.columns else None
        cell = {}
        if gc is not None and gt is not None and not (pd.isna(gc) or pd.isna(gt)):
            ris = f"{int(gc)}-{int(gt)}"
            for eng in ("motore", "fusione", "solostat"):
                merc, conf = tre[eng]
                won = _pronostico_vinto(_norm_merc(merc), gc, gt) if merc else None
                if won is True:
                    cell[eng] = "✅"; conteggi[eng][0] += 1
                elif won is False:
                    cell[eng] = "❌"; conteggi[eng][1] += 1
                else:
                    cell[eng] = "—"
            # pronostico EV (mercato singolo col massimo valore atteso)
            we = _pronostico_vinto(_norm_merc(merc_ev), gc, gt) if merc_ev else None
            if we is True:
                cell["ev"] = "✅"; conteggi["ev"][0] += 1
            elif we is False:
                cell["ev"] = "❌"; conteggi["ev"][1] += 1
            else:
                cell["ev"] = "—" if merc_ev else ""
            # pronostico Cristiano (combo)
            wc = _combo_vinta(pron_cri, gc, gt) if pron_cri else None
            if wc is True:
                cell["cristiano"] = "✅"; conteggi["cristiano"][0] += 1
            elif wc is False:
                cell["cristiano"] = "❌"; conteggi["cristiano"][1] += 1
            else:
                cell["cristiano"] = "—" if pron_cri else ""
        else:
            ris = "in attesa"; aperti += 1
            for eng in ("motore", "fusione", "solostat", "ev", "cristiano"):
                cell[eng] = ""
        righe.append({
            "Data": r.get("data"), "Casa": r.get("squadra_casa"),
            "Trasferta": r.get("squadra_trasferta"), "Risultato": ris,
            "Competizione": _label_da_comp(comp_per_id.get(_txt(r.get("partita_id"))),
                                           comp_df_st) or "",
            "🎯 Motore": tre["motore"][0] or "—", "Conf. M": tre["motore"][1],
            "✓M": cell["motore"],
            "🔀 Fusione": tre["fusione"][0] or "—", "Conf. F": tre["fusione"][1],
            "✓F": cell["fusione"],
            "📊 Statistico": tre["solostat"][0] or "—", "Conf. S": tre["solostat"][1],
            "✓S": cell["solostat"],
            "💰 EV": (f"{merc_ev} ({'+' if (val_ev or 0) >= 0 else ''}{val_ev}%)"
                      if merc_ev and val_ev is not None else (merc_ev or "—")),
            "✓EV": cell["ev"],
            "✍️ Cristiano": pron_cri or "—", "✓C": cell["cristiano"],
        })

    st.markdown("**Riuscita dei pronostici**")
    cols = st.columns(5)
    nomi = {"motore": "🎯 Motore", "fusione": "🔀 Fusione", "solostat": "📊 Statistico",
            "ev": "💰 EV", "cristiano": "✍️ Cristiano"}
    for i, eng in enumerate(("motore", "fusione", "solostat", "ev", "cristiano")):
        v, p = conteggi[eng]
        tot = v + p
        with cols[i]:
            st.caption(nomi[eng])
            if tot:
                cc = st.columns(2)
                cc[0].metric("Giocati", tot)
                cc[1].metric("Riuscita", f"{v/tot*100:.0f}%")
            else:
                st.caption("Nessuno concluso.")
    if aperti:
        st.caption(f"{aperti} in attesa di risultato.")

    # opzioni competizione per il menu a tendina: anagrafica + TUTTE quelle nei dati
    comp_opts = [""]
    comp_code_by_label = {}
    # 1) dall'anagrafica (etichetta leggibile -> codice)
    if comp_df_st is not None and not comp_df_st.empty:
        for _, cc_ in comp_df_st.iterrows():
            lab = _label_da_comp(cc_.get("nome_corto"), comp_df_st)
            if lab and lab not in comp_code_by_label:
                comp_code_by_label[lab] = cc_.get("nome_corto")
                comp_opts.append(lab)
    # 2) da tutte le competizioni presenti nelle partite (anche non mappate in anagrafica)
    if not df_tutte.empty and "competizione" in df_tutte.columns:
        for code in df_tutte["competizione"].dropna().unique():
            lab = _label_da_comp(code, comp_df_st)
            if lab and lab not in comp_code_by_label:
                comp_code_by_label[lab] = code
                comp_opts.append(lab)
    comp_opts = [comp_opts[0]] + sorted(comp_opts[1:], key=lambda x: x.lower())

    tab = pd.DataFrame(righe)

    # --- filtro per data (calendario) ---
    if "Data" in tab.columns and not tab.empty:
        _date = pd.to_datetime(tab["Data"], errors="coerce")
        dmin = _date.min()
        dmax = _date.max()
        if pd.notna(dmin) and pd.notna(dmax):
            with st.expander("📅 Filtra per data"):
                usa_filtro = st.checkbox("Attiva filtro per data", value=False,
                                         key="storico_usa_data")
                intervallo = st.date_input(
                    "Intervallo (da – a)", value=(dmin.date(), dmax.date()),
                    min_value=dmin.date(), max_value=dmax.date(), key="storico_range_data")
                if usa_filtro and isinstance(intervallo, (list, tuple)) and len(intervallo) == 2:
                    d_da, d_a = intervallo
                    mask = (_date.dt.date >= d_da) & (_date.dt.date <= d_a)
                    tab = tab[mask].reset_index(drop=True)
                    st.caption(f"Mostro {len(tab)} pronostici dal {d_da:%d/%m/%Y} al {d_a:%d/%m/%Y}.")

    # --- filtro per nome squadra (testo) ---
    if not tab.empty and {"Casa", "Trasferta"}.issubset(tab.columns):
        cerca = st.text_input("🔎 Filtra per squadra", key="storico_cerca_squadra",
                              placeholder="Scrivi parte del nome (casa o trasferta)…")
        if cerca and cerca.strip():
            q = cerca.strip().lower()
            m = (tab["Casa"].astype(str).str.lower().str.contains(q, na=False) |
                 tab["Trasferta"].astype(str).str.lower().str.contains(q, na=False))
            tab = tab[m].reset_index(drop=True)
            st.caption(f"Mostro {len(tab)} pronostici che contengono «{cerca.strip()}».")

    st.markdown("**Inserisci risultati e competizioni** direttamente qui (formato risultato: "
                "`1-1`, `2-0`…). Poi premi Salva. Le colonne dei pronostici non sono modificabili.")
    edit = st.data_editor(
        tab, use_container_width=True, hide_index=True, key="editor_storico",
        column_config={
            "Risultato": st.column_config.TextColumn("Risultato", help="Es. 1-1, 2-0"),
            "Competizione": st.column_config.SelectboxColumn(
                "Competizione", options=comp_opts,
                help="Seleziona la competizione dal menu. Include tutte quelle presenti "
                     "nei dati e in anagrafica."),
            "Data": st.column_config.Column(disabled=True),
            "Casa": st.column_config.Column(disabled=True),
            "Trasferta": st.column_config.Column(disabled=True),
            "🎯 Motore": st.column_config.Column(disabled=True),
            "Conf. M": st.column_config.Column(disabled=True),
            "✓M": st.column_config.Column(disabled=True),
            "🔀 Fusione": st.column_config.Column(disabled=True),
            "Conf. F": st.column_config.Column(disabled=True),
            "✓F": st.column_config.Column(disabled=True),
            "📊 Statistico": st.column_config.Column(disabled=True),
            "Conf. S": st.column_config.Column(disabled=True),
            "✓S": st.column_config.Column(disabled=True),
            "💰 EV": st.column_config.Column("💰 EV", disabled=True,
                help="Mercato col valore atteso (EV) più alto e la sua percentuale."),
            "✓EV": st.column_config.Column(disabled=True),
            "✍️ Cristiano": st.column_config.TextColumn(
                "✍️ Cristiano", help="Il tuo pronostico (combo con '+'). Modificabile qui."),
            "✓C": st.column_config.Column(disabled=True),
        })

    if st.button("💾 Salva risultati e competizioni", type="primary"):
        import re as _re
        recs = []
        for i, r in pron.reset_index(drop=True).iterrows():
            pid = _txt(r.get("partita_id"))
            if not pid:
                continue
            rec = {"id": pid}
            cambia = False
            # risultato
            ris_txt = _txt(edit.iloc[i]["Risultato"])
            mm = _re.match(r"^\s*(\d+)\s*[-:]\s*(\d+)\s*$", ris_txt)
            if mm:
                gc_new, gt_new = int(mm.group(1)), int(mm.group(2))
                old_gc, old_gt = r.get("gol_casa"), r.get("gol_trasferta")
                if (pd.isna(old_gc) or pd.isna(old_gt) or int(old_gc) != gc_new
                        or int(old_gt) != gt_new):
                    rec["gol_casa"] = gc_new
                    rec["gol_trasferta"] = gt_new
                    cambia = True
            # competizione: risolvi al codice tramite la mappa etichetta->codice
            comp_lab = _txt(edit.iloc[i].get("Competizione"))
            comp_old_lab = _label_da_comp(comp_per_id.get(pid), comp_df_st) or ""
            if comp_lab and comp_lab != comp_old_lab:
                nuovo = comp_code_by_label.get(comp_lab)
                if nuovo:
                    rec["competizione"] = nuovo
                    cambia = True
            if cambia:
                recs.append(rec)
            # pronostico Cristiano: salva se modificato (aggiornamento diretto sui pronostici)
            if "✍️ Cristiano" in edit.columns:
                cri_new = _txt(edit.iloc[i].get("✍️ Cristiano"))
                cri_new = "" if cri_new == "—" else cri_new
                cri_old = _txt(r.get("pron_cristiano")) if "pron_cristiano" in pron.columns else ""
                if cri_new != cri_old:
                    try:
                        cli = get_client()
                        if cli:
                            cli.table("pronostici").update(
                                {"pron_cristiano": cri_new or None}).eq(
                                "partita_id", pid).execute()
                    except Exception:
                        pass
        if recs:
            try:
                aggiorna_partite(recs)
                # aggiorna anche i gol nella tabella pronostici (esito immediato)
                cli = get_client()
                if cli:
                    for rc in recs:
                        if "gol_casa" in rc and "gol_trasferta" in rc:
                            try:
                                cli.table("pronostici").update(
                                    {"gol_casa": rc["gol_casa"], "gol_trasferta": rc["gol_trasferta"]}
                                ).eq("partita_id", rc["id"]).execute()
                            except Exception:
                                pass
                st.cache_data.clear()
                st.success(f"Aggiornate {len(recs)} partite (risultati/competizioni). "
                           "Gli esiti dei pronostici sono aggiornati.")
                st.rerun()
            except Exception as e:
                st.error(f"Errore nel salvataggio: {e}")
        else:
            st.info("Nessuna modifica valida da salvare.")

    # riepilogo pre-partita tramite selettore (non più click, per compatibilità con l'editor)
    st.caption("Per aprire il 📋 riepilogo pre-partita di un pronostico, selezionalo qui:")
    opzioni_riep = {f"{r.get('data')} · {r.get('squadra_casa')} - {r.get('squadra_trasferta')}": idx
                    for idx, (_, r) in enumerate(pron.iterrows())}
    scelta_riep = st.selectbox("Pronostico", ["—"] + list(opzioni_riep.keys()), key="sel_riep")
    idxs = [opzioni_riep[scelta_riep]] if scelta_riep in opzioni_riep else []
    if idxs:
        r = pron.iloc[idxs[0]]
        home = _txt(r.get("squadra_casa"))
        away = _txt(r.get("squadra_trasferta"))
        st.divider()
        st.subheader(f"🔎 {home} - {away}")
        # scheda completa salvata (snapshot), altrimenti il testo del riepilogo
        snap = None
        if "scheda_json" in pron.columns and _txt(r.get("scheda_json")):
            try:
                snap = json.loads(r["scheda_json"])
            except Exception:
                snap = None
        if snap:
            render_scheda_st(snap, home, away)
        else:
            txt = _txt(r.get("riepilogo")) if "riepilogo" in pron.columns else ""
            st.markdown("**📋 Riepilogo da copiare (note)**")
            if txt:
                st.code(txt, language=None)
            else:
                st.info("Per questo pronostico non è stata salvata la scheda completa. "
                        "Verrà salvata riaprendo la partita in 🔮 Analisi prima che si giochi, "
                        "oppure per i pronostici generati d'ora in avanti.")
    st.caption("Nota: la percentuale di riuscita è indicativa finché i numeri sono piccoli. "
               "Serve tempo e volume per trarne conclusioni.")


def main():
    user = login_gate()

    with st.sidebar:
        st.markdown(f"**Utente:** {user['username']}  \n_ruolo: {user['ruolo']}_")
        pagina = st.radio("Menu", ["📥 Ultimi risultati e quote", "📊 Estrattore risultati",
                                   "🗓️ Estrattore pianificazione", "🔮 Analisi & Pronostico",
                                   "📈 Storico pronostici", "🧪 Backtest",
                                   "🗄️ Database", "⚙️ Configurazione"])
        if st.button("Esci"):
            st.session_state.pop("user", None)
            st.rerun()

    if pagina.startswith("📥"):
        pagina_estrattore(user)
    elif pagina.startswith("📊"):
        pagina_estrattore_risultati(user)
    elif pagina.startswith("🗓️"):
        pagina_estrattore_pianificazione(user)
    elif pagina.startswith("🔮"):
        pagina_analisi(user)
    elif pagina.startswith("📈"):
        pagina_storico_pronostici(user)
    elif pagina.startswith("🧪"):
        pagina_backtest(user)
    elif pagina.startswith("🗄️"):
        pagina_database(user)
    else:
        pagina_configurazione(user)


if __name__ == "__main__":
    main()
